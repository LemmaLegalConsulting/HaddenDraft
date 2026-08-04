#!/usr/bin/env python
import os
import sys
import json
import io
import requests
import django
from PIL import Image, ImageDraw
import docx

# Set up Django environment
backend_dir = os.path.dirname(os.path.abspath(__file__))
if backend_dir not in sys.path:
    sys.path.insert(0, backend_dir)

os.environ.setdefault("DJANGO_SETTINGS_MODULE", "config.settings")
django.setup()

from django.conf import settings
from apps.matters.models import Matter
from apps.sources.connectors.legalserver import LegalServerClient

# Helper functions to convert text into PNG images, JPG images, PDFs (image-based, non-searchable), DOCX, and TXT
def generate_image_bytes(filename, text_content, fmt="PNG"):
    width, height = 1200, 1600
    img = Image.new("RGB", (width, height), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)

    # Draw header bar
    draw.rectangle([(0, 0), (width, 90)], fill=(235, 238, 242))
    draw.text((40, 30), f"DOCUMENT ATTACHMENT: {filename}", fill=(30, 40, 50))

    # Render document body lines
    y = 130
    for paragraph in text_content.split("\n"):
        while len(paragraph) > 75:
            chunk = paragraph[:75]
            draw.text((50, y), chunk, fill=(10, 10, 10))
            y += 32
            paragraph = paragraph[75:]
        if paragraph:
            draw.text((50, y), paragraph, fill=(10, 10, 10))
            y += 32
        else:
            y += 18
        if y > height - 80:
            break

    buf = io.BytesIO()
    img.save(buf, format=fmt)
    return buf.getvalue()


def generate_docx_bytes(title, text_content):
    doc = docx.Document()
    doc.add_heading(title, level=1)
    for line in text_content.split("\n"):
        if line.strip():
            doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# Document Definitions with specific multi-format file extensions
MULTI_FORMAT_CASES = {
    "Eleanor Vance": [
        ("Summons_and_Complaint_Eleanor_Vance.pdf", "pdf", (
            "COURT OF COMMON PLEAS\nCUYAHOGA COUNTY, OHIO\nHOUSING DIVISION\n\n"
            "APEX PROPERTIES LLC, Plaintiff, v. ELEANOR VANCE, Defendant.\n"
            "Case No. 2026-CVG-008912\n\n"
            "COMPLAINT IN EVICTION AND FOR MONEY DAMAGES\n"
            "1. Plaintiff Apex Properties LLC is landlord of 1428 Elm Street, Apt 3B, Cleveland, OH 44113.\n"
            "2. Defendant failed to pay monthly rent ($950.00) for April, May, and June 2026.\n"
            "3. 3-day notice served June 1, 2026. Back rent owed: $2,850.00 plus $150 late fees."
        )),
        ("Lease_Agreement_Eleanor_Vance.pdf", "pdf", (
            "RESIDENTIAL LEASE AGREEMENT\n"
            "Landlord: Apex Properties LLC | Tenant: Eleanor Vance\n"
            "Premises: 1428 Elm Street, Apt 3B, Cleveland, OH 44113\n"
            "Rent: $950.00/month. Term: August 1, 2025 - July 31, 2026.\n"
            "Maintenance: Landlord shall maintain plumbing and structural components in good repair."
        )),
        ("Notice_to_Quit_Photo.png", "png", (
            "3-DAY NOTICE TO LEAVE PREMISES (O.R.C. 1923.04)\n"
            "Date: June 1, 2026\nTo: Eleanor Vance, 1428 Elm Street, Apt 3B, Cleveland, OH 44113\n"
            "Reason: Unpaid rent for April and May 2026 totaling $1,900.00."
        )),
        ("Intake_Notes_Eleanor_Vance.docx", "docx", (
            "LEGAL AID INTAKE NOTES\nClient: Eleanor Vance | Date: June 15, 2026\n"
            "Bathroom ceiling leak reported Jan 18. Black mold 3x4ft developed. Child ER visit April 28 for asthma.\n"
            "Rent withheld saved in escrow account ($2,850 total). Defense under ORC 5321.04 & 5321.02 retaliation."
        ))
    ],

    "Marcus Vance": [
        ("Summons_and_Complaint_Marcus_Vance.pdf", "pdf", (
            "CLEVELAND MUNICIPAL COURT - HOUSING DIVISION\n"
            "METRO HOUSING RENTALS LLC v. MARCUS VANCE\nCase No. 2026-CVG-010422\n\n"
            "COMPLAINT FOR EVICTION AND MONEY DAMAGES\n"
            "1. Landlord owns 3204 Superior Ave, Apt 12, Cleveland, OH 44114.\n"
            "2. Defendant Marcus Vance is in default of monthly tenant portion of rent ($150/mo) and claims $1,800 total balance.\n"
            "3. 3-day notice served May 28, 2026."
        )),
        ("HAP_Contract_and_Voucher_CMHA.pdf", "pdf", (
            "HOUSING ASSISTANCE PAYMENTS (HAP) CONTRACT - SECTION 8 VOUCHER PROGRAM\n"
            "Housing Authority: Cuyahoga Metropolitan Housing Authority (CMHA)\n"
            "Tenant: Marcus Vance | Landlord: Metro Housing Rentals LLC\n"
            "Contract Rent: $900.00 | Housing Assistance Payment (CMHA): $750.00 | Tenant Rent Share: $150.00."
        )),
        ("Tenant_Rent_Ledger_Extract.png", "png", (
            "METRO HOUSING RENTALS LLC - TENANT ACCOUNT STATEMENT\n"
            "Account: Marcus Vance - 3204 Superior Ave #12\n"
            "Jan 2026: Rent $900 | CMHA paid $750 | Tenant paid $150 | Bal $0\n"
            "Feb-Apr 2026: CMHA electronic payment failed due to vendor bank update; landlord charged full $900/mo to tenant."
        )),
        ("Legal_Aid_Intake_Marcus_Vance.txt", "txt", (
            "INTAKE SUMMARY - MARCUS VANCE\n"
            "Client paid his $150 share every month via money order. CMHA missed 2 payments during portal transition.\n"
            "Landlord cannot evict tenant for CMHA's delayed subsidy under Federal HAP contract regulations & HUD rules."
        ))
    ],

    "Linda Thompson": [
        ("Summons_and_Complaint_Linda_Thompson.pdf", "pdf", (
            "CLEVELAND HEIGHTS MUNICIPAL COURT\n"
            "EUCLID HEIGHTS APARTMENTS LLC v. LINDA THOMPSON\nCase No. 2026-CVG-004119\n\n"
            "COMPLAINT FOR FORCIBLE ENTRY AND DETAINER\n"
            "Premises: 2195 Coventry Rd, Apt 4C, Cleveland Heights, OH 44118.\n"
            "Claim: Unpaid rent for April and May 2026 ($1,700 balance)."
        )),
        ("Partial_Rent_Payment_Receipt.jpg", "jpg", (
            "EUCLID HEIGHTS APARTMENTS - RECEIPT OF PAYMENT\n"
            "Date: June 5, 2026 (4 days after 3-day notice served on June 1, 2026)\n"
            "Received from Linda Thompson: $400.00 via cashier's check #88412.\n"
            "Accepted and deposited by Property Manager without reservation."
        )),
        ("Hospital_ER_Discharge_Report.pdf", "pdf", (
            "UNIVERSITY HOSPITALS RAINBOW BABIES & CHILDREN'S HOSPITAL\n"
            "Patient: Jayden Thompson (Age 5) | Mother: Linda Thompson\n"
            "Admission Date: May 12, 2026 | Diagnosis: Acute severe asthma exacerbation.\n"
            "Environmental Notes: Triggered by severe indoor mold spore exposure in bathroom."
        )),
        ("Intake_Notes_Linda_Thompson.docx", "docx", (
            "LEGAL AID INTAKE NOTES - LINDA THOMPSON\n"
            "Defenses: 1. Waiver of 3-Day Notice under Ohio law by accepting $400 partial rent after notice.\n"
            "2. Breach of Warranty of Habitability (O.R.C. 5321.04) due to toxic mold & ceiling collapse."
        ))
    ],

    "Robert Garcia": [
        ("Summons_and_Complaint_Robert_Garcia.pdf", "pdf", (
            "LAKEWOOD MUNICIPAL COURT\nLAKEWOOD PROPERTY GROUP LLC v. ROBERT GARCIA\n"
            "Case No. 2026-CVG-002180\n"
            "Filed: June 12, 2026 for nonpayment of $1,100 June rent at 14500 Detroit Ave #602."
        )),
        ("Three_Day_Notice_Scan.png", "png", (
            "3-DAY NOTICE TO LEAVE PREMISES\nServed: Wednesday, June 10, 2026 at 4:00 PM.\n"
            "Note: Landlord filed court action Friday, June 12, violating 3 full business days rule."
        )),
        ("Physician_Verification_Mobility_Impairment.jpg", "jpg", (
            "METROHEALTH MEDICAL CENTER - PHYSICIAN STATEMENT\n"
            "Patient: Robert Garcia (DOB: 03/14/1954)\n"
            "Diagnosis: Severe osteoarthritis, wheelchair dependent. Requires elevator access and home care."
        )),
        ("Intake_Notes_Robert_Garcia.txt", "txt", (
            "LEGAL AID INTAKE NOTES - ROBERT GARCIA\n"
            "Motion to Dismiss for lack of subject matter jurisdiction: 3-day notice served June 10, suit filed June 12.\n"
            "Under Ohio law (ORC 1923.04), 3 full business days must elapse before complaint filing."
        ))
    ],

    "James Miller": [
        ("Summons_and_Complaint_James_Miller.pdf", "pdf", (
            "PARMA MUNICIPAL COURT\nPARMA LANDLORDS LLC v. JAMES MILLER\nCase No. 2026-CVG-005120\n"
            "Premises: 5812 Ridge Rd #104, Parma, OH 44129. Amount claimed: $1,900.00."
        )),
        ("Rental_Assistance_Portal_Screenshot.png", "png", (
            "CUYAHOGA COUNTY HOUSING STABILITY FUND\n"
            "Applicant: James Miller | Status: APPROVED for $2,400.00.\n"
            "Disbursement check scheduled to Parma Landlords LLC within 10 business days."
        )),
        ("Intake_Notes_James_Miller.docx", "docx", (
            "INTAKE NOTES - JAMES MILLER\n"
            "File Motion for Continuance to allow rental assistance check to process. Landlord open to dismissal upon receipt."
        ))
    ],

    "Sarah Jenkins": [
        ("Summons_and_Complaint_Sarah_Jenkins.pdf", "pdf", (
            "CLEVELAND HEIGHTS MUNICIPAL COURT\nHEIGHTS REALTY LLC v. SARAH JENKINS\nCase No. 2026-CVG-003910\n"
            "Action for holdover possession of 3421 Mayfield Rd #2B."
        )),
        ("City_Housing_Inspection_Report.pdf", "pdf", (
            "CLEVELAND HEIGHTS BUILDING DEPARTMENT\n"
            "Inspection Date: February 14, 2026 | Premises: 3421 Mayfield Rd #2B\n"
            "Violation Cited: Heating unit failing to maintain 68 deg F minimum temperature requirement (ORC 5321.04)."
        )),
        ("Heat_Violation_Photo.png", "png", (
            "PHOTOGRAPH OF THERMOMETER - LIVING ROOM\n"
            "Date: Feb 12, 2026 | Indoor Temperature Reading: 48 deg F during freezing outdoor weather."
        ))
    ],

    "Charles Davis": [
        ("Summons_and_Complaint_Charles_Davis.pdf", "pdf", (
            "SHAKER HEIGHTS MUNICIPAL COURT\nSHAKER SQUARE APARTMENTS LLC v. CHARLES DAVIS\n"
            "Case No. 2026-CVG-001920\nUnpaid rent for May/June 2026 ($1,650 balance)."
        )),
        ("Notice_to_Quit_Photo.jpg", "jpg", (
            "PHOTO OF 3-DAY NOTICE POSTED ON DOOR\n"
            "Address: 16800 Chagrin Blvd #3A, Shaker Heights, OH."
        ))
    ],

    "Donna Evans": [
        ("Summons_and_Complaint_Donna_Evans.png", "png", (
            "EAST CLEVELAND MUNICIPAL COURT\nEAST CLEVELAND PROPERTIES LLC v. DONNA EVANS\n"
            "Case No. 2026-CVG-007140\nAlleged lease violation: Unauthorized guest over 14 days."
        )),
        ("Guest_Caregiver_Statement.docx", "docx", (
            "STATEMENT OF MARY EVANS\n"
            "I am Donna Evans' sister. I stayed at 13405 Euclid Ave #510 for 5 days following her knee surgery to assist her."
        ))
    ],

    "Thomas Wilson": [
        ("Summons_and_Complaint_Thomas_Wilson.pdf", "pdf", (
            "GARFIELD HEIGHTS MUNICIPAL COURT\nTURNEY ROAD MANAGEMENT LLC v. THOMAS WILSON\n"
            "Case No. 2026-CVG-003110\nEviction for alleged waste and property damage."
        )),
        ("Police_Incident_Report_Garfield_Heights.png", "png", (
            "GARFIELD HEIGHTS POLICE DEPARTMENT - INCIDENT REPORT\n"
            "Report No. 26-09812 | Incident: Attempted Burglary / Vandalism\n"
            "Location: 4810 Turney Rd #14 | Reporting Party: Thomas Wilson\n"
            "Officer Notes: Unknown suspect attempted forced entry causing outer window crack and lock deformation."
        ))
    ],

    "Patricia Taylor": [
        ("Summons_and_Complaint_Patricia_Taylor.pdf", "pdf", (
            "BEREA MUNICIPAL COURT\nSTRONGSVILLE APARTMENTS LLC v. PATRICIA TAYLOR\n"
            "Case No. 2026-CVG-004810\nComplaint signed: /s/ Brenda Vance, Property Manager (non-attorney)."
        )),
        ("Notice_Photo_Patricia_Taylor.jpg", "jpg", (
            "PHOTO OF NOTICE TO VACATE POSTED BY PROPERTY MANAGER"
        ))
    ],

    "Christopher Anderson": [
        ("Summons_and_Complaint_Christopher_Anderson.pdf", "pdf", (
            "ROCKY RIVER MUNICIPAL COURT\nWESTLAKE VILLAGE APARTMENTS LLC v. CHRISTOPHER ANDERSON\n"
            "Case No. 2026-CVG-001840\nClaim: Lease termination for alleged noise disturbance."
        )),
        ("Gas_Company_Inspection_Report.png", "png", (
            "DOMINION ENERGY OHIO - EMERGENCY SERVICE CALL LOG\n"
            "Date: May 19, 2026 | Address: 26900 Detroit Rd #115\n"
            "Technician Findings: Active gas leak detected at stove supply line; gas shut off pending landlord repair."
        ))
    ],

    "Samuel Green": [
        ("Medicaid_Notice_of_Action.pdf", "pdf", (
            "OHIO DEPARTMENT OF MEDICAID - NOTICE OF ACTION\n"
            "To: Samuel Green | Date: May 10, 2026\n"
            "Notice: Termination of PASSPORT Waiver Personal Care Aide Hours effective June 1, 2026."
        )),
        ("Treating_Physician_Level_of_Care_Assessment.jpg", "jpg", (
            "CLEVELAND CLINIC - GERIATRIC MEDICINE EVALUATION\n"
            "Patient: Samuel Green (Age 78)\n"
            "Assessment: Patient requires assistance with 4 ADLs (bathing, dressing, mobility, medication management)."
        ))
    ],

    "Maria Santos": [
        ("Divorce_Complaint_Santos.pdf", "pdf", (
            "CUYAHOGA COUNTY COMMON PLEAS COURT - DOMESTIC RELATIONS DIVISION\n"
            "MARIA SANTOS v. CARLOS SANTOS\nCase No. DR-26-381902\n"
            "Complaint for Divorce, Custody, and Equitable Division of Marital Property."
        )),
        ("Marriage_Certificate_Scan.png", "png", (
            "STATE OF OHIO - CERTIFICATE OF MARRIAGE\n"
            "Spouse A: Maria Santos | Spouse B: Carlos Santos | Date: October 14, 2018."
        )),
        ("Proposed_Parenting_Plan.docx", "docx", (
            "PROPOSED SHARED PARENTING PLAN\nMother Maria Santos proposes primary residential parent designation."
        ))
    ],

    "David Kowalski": [
        ("Creditor_Balance_Sheet.pdf", "pdf", (
            "SCHEDULE OF UNSECURED CREDITORS - DAVID KOWALSKI\n"
            "1. Cleveland Clinic Health System: $32,450.00 (Medical)\n"
            "2. Capital One Bank: $12,800.00 (Credit Card)"
        )),
        ("Medical_Bills_Scan.png", "png", (
            "CLEVELAND CLINIC HEALTH SYSTEM - STATEMENT OF ACCOUNT\n"
            "Patient: David Kowalski | Total Outstanding Balance: $32,450.00."
        ))
    ],

    "Aisha Jackson": [
        ("DUI_Arrest_Record_Excerpt.png", "png", (
            "CLEVELAND MUNICIPAL COURT - CRIMINAL DIVISION\n"
            "State of Ohio v. Marcus Jackson | Case No. 2026-TRC-011204\n"
            "Charge: OVI / Operating Vehicle Impaired."
        )),
        ("Motion_for_Supervised_Visitation.pdf", "pdf", (
            "CUYAHOGA COUNTY JUVENILE COURT\nAISHA JACKSON v. MARCUS JACKSON\n"
            "Motion to Modify Allocation of Parental Rights and Order Supervised Visitation."
        ))
    ],

    "Carlos Mendez": [
        ("Overtime_Hours_Log_Photo.jpg", "jpg", (
            "CARLOS MENDEZ - HANDWRITTEN OVERTIME HOURS LOG PHOTO\n"
            "Avg hours worked: 54 hours/week. Paid flat cash rate without 1.5x overtime multiplier."
        )),
        ("FLSA_Wage_Demand_Letter.pdf", "pdf", (
            "NOTICE OF WAGE CLAIM AND DEMAND FOR PAYMENT\n"
            "To: El Sol Restaurant & Grill LLC / Ricardo Lopez\n"
            "Demand for unpaid overtime compensation and liquidated damages under 29 U.S.C. 207."
        ))
    ],

    "Brenda Taylor": [
        ("SSA_Reconsideration_Denial.pdf", "pdf", (
            "SOCIAL SECURITY ADMINISTRATION - NOTICE OF RECONSIDERATION DENIAL\n"
            "Claimant: Brenda Taylor | SSN: XXX-XX-4819\nClaim for Supplemental Security Income."
        )),
        ("MRI_Spine_Medical_Report.png", "png", (
            "METROHEALTH NEUROLOGY - MRI LUMBAR SPINE SCAN REPORT\n"
            "Findings: Severe spinal stenosis L4-L5 with nerve root compression."
        ))
    ],

    "Karen Novak": [
        ("Resignation_Email_Screenshot.png", "png", (
            "SCREENSHOT OF RESIGNATION EMAIL\n"
            "From: Karen Novak To: HR Midwest Logistics\n"
            "Subject: Resignation due to unaddressed safety hazards on forklift loading dock."
        )),
        ("ODJFS_Unemployment_Appeal.pdf", "pdf", (
            "OHIO DEPT OF JOB & FAMILY SERVICES - APPEAL TO REVIEW COMMISSION\n"
            "Claimant Karen Novak appeals initial disallowance of benefits."
        ))
    ],

    "Olivia Martinez": [
        ("IEP_Speech_Evaluation_Excerpt.pdf", "pdf", (
            "CLEVELAND METROPOLITAN SCHOOL DISTRICT - IEP EVALUATION REPORT\n"
            "Student: Mateo Martinez (Age 8) | Diagnosis: ASD Level 2."
        )),
        ("Speech_Therapy_Assessment_Scan.jpg", "jpg", (
            "CLEVELAND SPEECH & HEARING CENTER - INDEPENDENT ASSESSMENT SCAN\n"
            "Recommends minimum 3 hours/week direct speech-language pathology services."
        ))
    ],

    "Evelyn Carter": [
        ("Property_Deed_Scan.png", "png", (
            "CUYAHOGA COUNTY RECORDER - DEED OF RECORD SCAN\n"
            "Grantor: Evelyn Carter | Location: 14208 Kinsman Rd, Cleveland OH."
        )),
        ("Last_Will_and_Testament_Draft.docx", "docx", (
            "LAST WILL AND TESTAMENT OF EVELYN CARTER\n"
            "I, Evelyn Carter, declare this to be my Last Will and Testament..."
        ))
    ],

    "Daniel O'Connor": [
        ("Medicare_Part_D_Denial_Notice.pdf", "pdf", (
            "HUMANA MEDICARE PART D - NOTICE OF DENIAL OF MEDICARE PRESCRIPTION DRUG COVERAGE\n"
            "Enrollee: Daniel O'Connor | Drug: Entresto 97/103mg | Reason: Non-formulary."
        )),
        ("Cardiologist_Prior_Auth_Letter.jpg", "jpg", (
            "CLEVELAND CLINIC HEART & VASCULAR INSTITUTE - DOCTOR LETTER PHOTO\n"
            "Re: Daniel O'Connor | Physician statement confirming formulary alternatives caused severe angioedema."
        ))
    ]
}


def run_multi_format_uploads():
    base_url = settings.LEGALSERVER_BASE_URL.rstrip('/')
    token = settings.LEGALSERVER_API_TOKEN
    headers = {
        'Authorization': f'Bearer {token}',
        'Accept': 'application/json'
    }
    v2_url = f"{base_url}/api/v2/documents"

    client = LegalServerClient()
    matters = client.search_matters(limit=100)
    print(f"Loaded {len(matters)} matters from LegalServer.")

    for m in matters:
        muuid = m.get('matter_uuid')
        db_id = m.get('case_id') or m.get('database_id') or m.get('id')
        mno = m.get('matter_identification_number') or m.get('case_number')
        first = m.get('first', '')
        last = m.get('last', '')
        name = f"{first} {last}".strip() or m.get('case_title', '')

        matching_docs = None
        for key, docs_list in MULTI_FORMAT_CASES.items():
            if key.casefold() in name.casefold() or name.casefold() in key.casefold():
                matching_docs = docs_list
                break

        if not matching_docs:
            continue

        existing_docs = client.get_matter_documents(muuid or db_id)
        existing_names = {d.get('name') or d.get('title') for d in existing_docs}

        print(f"\nProcessing [{mno}] {name} (DB ID {db_id}):")

        for filename, file_kind, content_text in matching_docs:
            if filename in existing_names:
                print(f"  -> Already present: {filename}")
                continue

            # Generate binary bytes based on target format (PDF, PNG, JPG, DOCX, TXT)
            if file_kind == "pdf":
                mime_type = "application/pdf"
                # Image-based PDF without searchable text layer
                file_bytes = generate_image_bytes(filename, content_text, fmt="PDF")
            elif file_kind == "png":
                mime_type = "image/png"
                # PNG image without searchable text layer
                file_bytes = generate_image_bytes(filename, content_text, fmt="PNG")
            elif file_kind in ("jpg", "jpeg"):
                mime_type = "image/jpeg"
                # JPEG image without searchable text layer
                file_bytes = generate_image_bytes(filename, content_text, fmt="JPEG")
            elif file_kind == "docx":
                mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                file_bytes = generate_docx_bytes(filename, content_text)
            else:
                mime_type = "text/plain"
                file_bytes = content_text.encode("utf-8")

            files = {
                'file': (filename, file_bytes, mime_type)
            }
            data = {
                'name': filename,
                'matter_uuid': muuid
            }

            res = requests.post(v2_url, headers=headers, files=files, data=data)
            if res.status_code in (200, 201):
                doc_uuid = res.json().get('data', {}).get('uuid')
                print(f"  -> SUCCESS [{file_kind.upper()}] uploaded {filename} via v2 API! (UUID: {doc_uuid})")
            else:
                print(f"  -> FAILED to upload {filename} via v2 API: {res.status_code} - {res.text}")

    print("\nFinished multi-format v2 document upload execution!")

if __name__ == "__main__":
    run_multi_format_uploads()
