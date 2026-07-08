"""Extract plain text/metadata from a rendered DOCX so validation can inspect it."""

import io
import re

from docx import Document

DOCKET_RE = re.compile(
    r"\b(?:No\.|Case No\.?|Docket No\.?)\s*[:#]?\s*([A-Za-z0-9][A-Za-z0-9\-/.]{2,40})",
    re.IGNORECASE,
)
PLACEHOLDER_RE = re.compile(r"\[[^\]\n]{2,80}\]")


def extract_docx_text(docx_bytes):
    document = Document(io.BytesIO(docx_bytes))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    tables = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    tables.append(cell.text)
    text = "\n".join(paragraphs + tables)
    return {
        "text": text,
        "paragraphs": paragraphs,
        "tables": tables,
        "wordCount": len(text.split()),
        "docketNumbers": DOCKET_RE.findall(text),
        "placeholders": PLACEHOLDER_RE.findall(text),
    }
