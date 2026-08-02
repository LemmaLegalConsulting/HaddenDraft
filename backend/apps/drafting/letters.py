"""Drafting a letter, as its own workflow rather than a motion without a caption.

A letter differs from a filing in three ways that matter here. Its stationery
carries the advocate's identity, so the body must not restate it. It is written
to a person rather than to a court, so the recipient and the action they need to
take drive the text. And it has no blocks to select, so the review step is over
one body of prose.

Composition renders the letter body into the letterhead, which keeps the
masthead, margins, and continuation header intact.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docxtpl import DocxTemplate

from apps.ai.prompt_catalog import render_prompt
from apps.ai.services import GenerationContext
from apps.templates_app.jinja_filters import template_environment
from apps.templates_app.letterhead_library import letterhead_for_author, letterhead_path
from apps.templates_app.letterheads import letterhead_context
from apps.templates_app.template_variables import normalize_docxtpl_blocks


LETTER_KINDS = [
    ("advice", "Advice letter"),
    ("engagement", "Engagement or non-engagement letter"),
    ("closing", "Closing letter"),
    ("demand", "Demand letter to a landlord"),
    ("service", "Service letter to opposing counsel"),
    ("appointment", "Appointment or contact letter"),
]

RECIPIENT_ROLES = [
    ("client", "Client"),
    ("opposing_counsel", "Opposing counsel"),
    ("landlord", "Landlord"),
    ("agency", "Agency or housing authority"),
    ("court", "Court"),
]


@dataclass
class LetterRequest:
    """Everything the advocate decides before the model writes anything."""

    letter_kind: str = "advice"
    recipient_name: str = ""
    recipient_role: str = "client"
    recipient_address: str = ""
    purpose: str = ""
    deadline: str = ""
    delivery: list[str] = field(default_factory=list)
    subject: str = ""

    def recipient_description(self):
        role = dict(RECIPIENT_ROLES).get(self.recipient_role, self.recipient_role)
        return f"{self.recipient_name or 'the recipient'} ({role.lower()})"


def _facts_block(context):
    lines = []
    for fact in context.selected_facts or []:
        lines.append(f"- {getattr(fact, 'text', '')} [{getattr(fact, 'source_label', 'case record')}]")
    for fact in context.selected_curated_facts or []:
        source = fact.get("citation") or fact.get("source") or "curated source"
        lines.append(f"- {fact.get('text', '')} [{source}]")
    return "\n".join(lines) or "- None"


def _sources_block(context):
    lines = [
        f"- {source.get('title', '')} [{source.get('citation') or source.get('sourceLabel', '')}]"
        for source in context.selected_sources or []
    ]
    return "\n".join(lines) or "- None"


def _case_reference(context):
    matter = context.matter
    parts = [getattr(matter, "summary", "") or getattr(matter, "matter_type", "")]
    number = (context.template_data or {}).get("court_case_number", "")
    if number:
        parts.append(f"Case No. {number}")
    court = getattr(matter, "jurisdiction", "")
    if court:
        parts.append(court)
    return ", ".join(part for part in parts if part) or "not supplied"


def draft_letter_body(request: LetterRequest, context: GenerationContext, *, llm_client=None):
    """Ask the model for the letter body, falling back to a reviewable stub."""
    from apps.ai.openai_client import OpenAIBackendError, OpenAIClient
    from apps.ai.services import ConstrainedDraftingService

    author = context.author_profile or {}
    prompt = render_prompt(
        "drafting.letter",
        letter_kind=dict(LETTER_KINDS).get(request.letter_kind, request.letter_kind),
        recipient_description=request.recipient_description(),
        matter_summary=getattr(context.matter, "summary", "") or "not supplied",
        jurisdiction=getattr(context.matter, "jurisdiction", "") or "not supplied",
        client_name=getattr(context.matter, "client_name", "") or "not supplied",
        client_pronouns=(context.template_data or {}).get("client_pronouns") or "not supplied",
        recipient_name=request.recipient_name or "not supplied",
        recipient_role=dict(RECIPIENT_ROLES).get(request.recipient_role, request.recipient_role),
        case_reference=_case_reference(context),
        advocate_name=author.get("displayName") or "the advocate",
        advocate_title=author.get("title") or "advocate",
        purpose=request.purpose or "not supplied",
        deadline=request.deadline or "not supplied",
        facts=_facts_block(context),
        sources=_sources_block(context),
    )
    service = ConstrainedDraftingService(llm_client=llm_client)
    client = llm_client or OpenAIClient()
    try:
        return service.normalize_generated_text(
            client.complete(
                system=prompt.system,
                user=prompt.user,
                model=prompt.default_model,
                reasoning_level=prompt.default_reasoning_level,
            )
        )
    except OpenAIBackendError:
        return letter_fallback(request, context)


def letter_fallback(request: LetterRequest, context: GenerationContext):
    """A reviewable skeleton when the model is unavailable.

    Every case-specific value the advocate has not supplied is left as a visible
    bracketed marker rather than guessed at, so an unattended failure cannot
    produce a letter that reads as finished.
    """
    author = context.author_profile or {}
    lines = [
        request.deadline or "[Date]",
        "",
        request.recipient_name or "[Recipient name]",
    ]
    if request.recipient_address:
        lines.extend(request.recipient_address.splitlines())
    else:
        lines.append("[Recipient address]")
    lines.append("")
    for method in request.delivery:
        lines.append(f"Sent via {method}")
    if request.delivery:
        lines.append("")
    lines.append(f"Re: {request.subject or _case_reference(context)}")
    lines.append("")
    lines.append(f"Dear {request.recipient_name or '[Recipient]'}:")
    lines.append("")
    lines.append(f"[Attorney review required: state the purpose - {request.purpose or 'not supplied'}]")
    lines.append("")
    lines.append(author.get("signoff") or "Sincerely,")
    lines.append("")
    lines.append(author.get("displayName") or "[Advocate name]")
    if author.get("title"):
        lines.append(author["title"])
    return "\n".join(lines)


def compose_letter_docx(body: str, *, author_profile=None, request: LetterRequest = None, output_path: Path):
    """Render the letter body onto the organization's letterhead."""
    request = request or LetterRequest()
    author_profile = author_profile or {}
    letterhead = letterhead_for_author(author_profile)
    path = letterhead_path(letterhead)
    context = letterhead_context(
        author_profile,
        subject=request.subject or request.recipient_name,
        date=request.deadline,
    )

    if not path or not path.is_file():
        # Without stationery the letter still has to be deliverable, so it is
        # written as a plain document rather than failing the draft.
        document = Document()
        for paragraph in body.split("\n"):
            document.add_paragraph(paragraph)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(output_path)
        return output_path, None

    class NormalizingDocxTemplate(DocxTemplate):
        def patch_xml(self, source_xml):
            return normalize_docxtpl_blocks(super().patch_xml(source_xml))

    template = NormalizingDocxTemplate(path)
    template.render(context, jinja_env=template_environment())
    rendered = template.docx

    _append_body(rendered, body)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    rendered.save(output_path)
    return output_path, letterhead


def _append_body(document, body: str):
    """Put the letter text in the stationery's body, keeping its empty layout."""
    existing = [paragraph for paragraph in document.paragraphs if paragraph.text.strip()]
    for paragraph in existing:
        # A maintained letterhead sometimes carries a sample letter; the drafted
        # body replaces it rather than appending underneath it.
        parent = paragraph._p.getparent()
        if parent is not None:
            parent.remove(paragraph._p)
    for line in re.split(r"\n", body):
        document.add_paragraph(line)
