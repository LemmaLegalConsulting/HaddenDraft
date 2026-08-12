import tempfile
import zipfile
from pathlib import Path

from django.test import TestCase, override_settings
from docx import Document

from apps.ai.services import GenerationContext
from apps.drafting.letters import LetterRequest, compose_letter_docx, letter_fallback
from apps.matters.models import Matter
from apps.templates_app.models import Letterhead
from apps.templates_app.test_letterheads import make_letterhead
from apps.templates_app.letterheads import prepare_letterhead


AUTHOR = {
    "displayName": "Dana Ruiz",
    "title": "Staff Attorney",
    "phone": "216.555.0142",
    "fax": "440.352.0015",
    "email": "druiz@example.org",
    "signoff": "Sincerely,",
    "officeName": "Cleveland",
}


class LetterDraftingTests(TestCase):
    def setUp(self):
        self.temp = tempfile.TemporaryDirectory()
        self.addCleanup(self.temp.cleanup)
        self.root = Path(self.temp.name)
        self.media = self.root / "media"
        self.media.mkdir()
        media = override_settings(MEDIA_ROOT=str(self.media))
        media.enable()
        self.addCleanup(media.disable)
        self.matter = Matter.objects.create(
            external_id="2026-CVG-9",
            client_name="Jane Tenant",
            matter_type="Eviction",
            jurisdiction="Cleveland Municipal Court",
            summary="Nonpayment eviction",
        )
        self.context = GenerationContext(
            matter=self.matter,
            selected_facts=[],
            selected_curated_facts=[],
            selected_sources=[],
            template=None,
            mode="letter",
            instructions="",
            author_profile=AUTHOR,
            template_data={"court_case_number": "2026 CVG 011123"},
        )
        self.request = LetterRequest(
            letter_kind="advice",
            recipient_name="Mr. Charles Mosby",
            recipient_role="client",
            recipient_address="970 Stevenson Rd.\nCleveland, Ohio 44110",
            purpose="Urge attendance at the eviction hearing",
            deadline="August 2, 2026",
            delivery=["Email"],
            subject="Advice for Your Eviction Hearing",
        )

    def install_letterhead(self):
        source = self.root / "advocate.docx"
        make_letterhead(source)
        prepared = self.media / "letterhead.docx"
        prepare_letterhead(source, prepared)
        return Letterhead.objects.create(
            slug="example",
            title="Example Legal Aid",
            docx="letterhead.docx",
            source_kind="database",
            is_default=True,
        )

    def test_fallback_marks_missing_values_instead_of_inventing_them(self):
        body = letter_fallback(self.request, self.context)

        self.assertIn("August 2, 2026", body)
        self.assertIn("Mr. Charles Mosby", body)
        self.assertIn("Re: Advice for Your Eviction Hearing", body)
        self.assertIn("Attorney review required", body)
        self.assertIn("Dana Ruiz", body)

    def test_letter_composes_onto_the_letterhead(self):
        self.install_letterhead()
        output = self.root / "letter.docx"

        path, letterhead = compose_letter_docx(
            letter_fallback(self.request, self.context),
            author_profile=AUTHOR,
            request=self.request,
            output_path=output,
        )

        document = Document(path)
        header = "\n".join(p.text for p in document.sections[0].header.paragraphs)
        body = "\n".join(p.text for p in document.paragraphs)
        self.assertIsNotNone(letterhead)
        # The stationery supplies the advocate; the body must not restate it.
        self.assertIn("Dana Ruiz", header)
        self.assertIn("216.555.0142", header)
        self.assertIn("Advice for Your Eviction Hearing", header)
        self.assertIn("Mr. Charles Mosby", body)
        self.assertNotIn("{{", header)

    def test_letter_is_still_produced_without_any_letterhead(self):
        """A missing letterhead must not cost the advocate the draft."""
        output = self.root / "plain.docx"

        path, letterhead = compose_letter_docx(
            letter_fallback(self.request, self.context),
            author_profile=AUTHOR,
            request=self.request,
            output_path=output,
        )

        self.assertIsNone(letterhead)
        body = "\n".join(p.text for p in Document(path).paragraphs)
        self.assertIn("Mr. Charles Mosby", body)

    def test_sample_letter_in_the_stationery_is_replaced_not_appended(self):
        letterhead = self.install_letterhead()
        document = Document(letterhead.docx.path)
        document.add_paragraph("Sample letter text from whoever saved this file.")
        document.save(letterhead.docx.path)
        output = self.root / "letter.docx"

        path, _letterhead = compose_letter_docx(
            letter_fallback(self.request, self.context),
            author_profile=AUTHOR,
            request=self.request,
            output_path=output,
        )

        body = "\n".join(p.text for p in Document(path).paragraphs)
        self.assertNotIn("Sample letter text", body)
        self.assertIn("Mr. Charles Mosby", body)

    def test_export_sanitizes_a_legacy_letterhead_without_mutating_it(self):
        source = self.root / "legacy.docx"
        make_letterhead(source)
        target = self.media / "legacy.docx"
        target.write_bytes(source.read_bytes())
        letterhead = Letterhead.objects.create(
            slug="legacy",
            title="Legacy Legal Aid",
            docx="legacy.docx",
            source_kind="database",
            is_default=True,
        )
        original = target.read_bytes()
        output = self.root / "legacy-export.docx"

        path, selected = compose_letter_docx(
            letter_fallback(self.request, self.context),
            author_profile=AUTHOR,
            request=self.request,
            output_path=output,
        )

        self.assertEqual(selected.id, letterhead.id)
        self.assertEqual(target.read_bytes(), original)
        with zipfile.ZipFile(path) as archive:
            core = archive.read("docProps/core.xml").decode("utf-8")
        self.assertNotIn("lastPrinted", core)
        self.assertNotIn("<dc:title", core)
        self.assertIn("Mr. Charles Mosby", "\n".join(p.text for p in Document(path).paragraphs))


class LetterheadWithoutBulletStyleTests(TestCase):
    """A letterhead that cannot render a list must not lose the letter.

    Stationery prepared from an organization's own Word file carries only the
    styles that file used, so "List Bullet" is frequently absent. python-docx
    raises KeyError for an unknown style name, and that took the whole export
    down with a 500.
    """

    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory()
        self.addCleanup(self.tmp.cleanup)
        self.root = Path(self.tmp.name)
        self.media = self.root / "media"
        self.media.mkdir()

    def install_letterhead_without_bullets(self):
        source = self.root / "advocate.docx"
        make_letterhead(source)
        document = Document(source)
        for style in list(document.styles):
            if style.name in ("List Bullet", "ListBullet"):
                style.delete()
        document.save(source)
        prepared = self.media / "letterhead.docx"
        prepare_letterhead(source, prepared)
        return Letterhead.objects.create(
            slug="no-bullets",
            title="Example Legal Aid",
            docx="letterhead.docx",
            source_kind="database",
            is_default=True,
        )

    def test_a_bulleted_letter_still_exports(self):
        with override_settings(MEDIA_ROOT=self.media):
            self.install_letterhead_without_bullets()
            output = self.root / "letter.docx"

            path, _ = compose_letter_docx(
                "- First point\n- Second point",
                author_profile=AUTHOR,
                request=LetterRequest(letter_kind="brief_advice", recipient_role="client"),
                output_path=output,
                formatted_body=[
                    {"runs": [{"text": "- First point", "format": 0}]},
                    {"runs": [{"text": "- Second point", "format": 0}]},
                ],
            )

        text = "\n".join(p.text for p in Document(path).paragraphs)
        # The dash survives, so the reader still sees a list.
        self.assertIn("- First point", text)
        self.assertIn("- Second point", text)
