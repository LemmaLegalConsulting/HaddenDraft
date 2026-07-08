import zipfile
import tempfile
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace
from xml.etree import ElementTree

from django.contrib.auth import get_user_model
from django.core.files import File
from django.urls import reverse
from django.test import TestCase
from django.test.utils import override_settings
from docx import Document

from apps.ai.services import ConstrainedDraftingService
from apps.drafting.models import DraftDocument, DraftingSession
from apps.drafting.services import (
    _normalize_goal_candidates,
    apply_plan_edits,
    create_drafts_from_plan,
    create_or_update_plan,
    recommend_goal_candidates,
    unanswered_missing_information,
)
from apps.exporting.services import _docx_render_context
from apps.exporting.services import export_docx
from apps.matters.models import Matter, MatterFact
from apps.templates_app.models import DocumentTemplate, TemplateBlock
from apps.templates_app.serializers import template_to_dict
from apps.templates_app.template_variables import (
    block_variable_metadata,
    extract_template_variables_from_text,
)


class DraftRenderingTests(TestCase):
    def test_goal_recommendation_fallback_suggests_continuance_from_rental_assistance(self):
        matter = Matter.objects.create(
            external_id="CASE-GOAL-CONTINUANCE",
            client_name="Jane Tenant",
            matter_type="Eviction",
            jurisdiction="Cleveland Municipal Court - Housing Division",
            summary="Tenant has pending rental assistance and needs more time before hearing.",
        )
        assistance_fact = MatterFact.objects.create(
            matter=matter,
            slug="pending-rental-assistance",
            title="Pending rental assistance",
            text="Tenant has an ERAP rental assistance application pending.",
            source_label="Intake",
        )
        MatterFact.objects.create(
            matter=matter,
            slug="hearing-deadline",
            title="Hearing deadline",
            text="The eviction hearing is scheduled soon and tenant needs more time.",
            source_label="Intake",
        )
        template, _created = DocumentTemplate.objects.update_or_create(
            slug="motion-continuance-cleveland",
            defaults={
                "title": "Motion for Continuance",
                "kind": "motion",
                "goal": "Ask the court to continue or postpone a hearing.",
                "aliases": ["continue hearing", "more time"],
                "jurisdiction": "Cleveland Municipal Court - Housing Division",
            },
        )
        TemplateBlock.objects.create(
            template=template,
            key="body",
            label="Body",
            block_type="argument",
            body="Continuance body.",
        )
        session = DraftingSession.objects.create(
            mode="draft_from_template",
            matter=matter,
        )

        with override_settings(AI_DRAFTING_ENABLED=False):
            payload = recommend_goal_candidates(session)

        self.assertTrue(payload["goals"])
        goal = payload["goals"][0]
        self.assertRegex(goal["goal"].casefold(), r"continue|more time")
        self.assertIn(assistance_fact.id, goal["supportingFactIds"])
        self.assertIn("motion-continuance-cleveland", goal["templateSlugs"])

    def test_goal_recommendation_endpoint_requires_login_and_keeps_session_goal(self):
        user = get_user_model().objects.create_user(username="reviewer", password="pass", is_superuser=True)
        matter = Matter.objects.create(
            external_id="CASE-GOAL-ENDPOINT",
            client_name="Jane Tenant",
            matter_type="Eviction",
            jurisdiction="Housing Court",
            summary="Tenant has a hearing and needs more time.",
        )
        DocumentTemplate.objects.create(
            slug="endpoint-continuance-template",
            title="Motion for Continuance",
            kind="motion",
            goal="Ask the court for more time.",
            aliases=["more time"],
        )
        session = DraftingSession.objects.create(
            mode="draft_from_template",
            matter=matter,
            goal="Original user goal",
        )
        url = reverse("api_session_recommend_goals", args=[session.id])

        unauthenticated = self.client.post(url, data='{"limit": 5}', content_type="application/json")
        self.assertEqual(unauthenticated.status_code, 401)

        self.client.login(username="reviewer", password="pass")
        with override_settings(AI_DRAFTING_ENABLED=False):
            response = self.client.post(url, data='{"limit": 5}', content_type="application/json")

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertIn("goals", payload)
        self.assertEqual(payload["session"]["goal"], "Original user goal")
        session.refresh_from_db()
        self.assertEqual(session.goal, "Original user goal")

    def test_goal_recommendation_normalizer_removes_invalid_ai_references(self):
        matter = Matter.objects.create(
            external_id="CASE-GOAL-NORMALIZE",
            client_name="Jane Tenant",
            matter_type="Eviction",
            jurisdiction="Housing Court",
            summary="Tenant has a hearing.",
        )
        valid_fact = MatterFact.objects.create(
            matter=matter,
            slug="valid-fact",
            title="Valid fact",
            text="Tenant has a pending hearing.",
            source_label="Intake",
        )
        valid_template = DocumentTemplate.objects.create(
            slug="valid-goal-template",
            title="Valid Goal Template",
            kind="motion",
            is_active=True,
        )
        inactive_template = DocumentTemplate.objects.create(
            slug="inactive-goal-template",
            title="Inactive Goal Template",
            kind="motion",
            is_active=False,
        )

        normalized = _normalize_goal_candidates(
            [
                {
                    "title": "Hearing continuance",
                    "goal": "Ask the court for more time before the hearing.",
                    "supporting_fact_ids": [valid_fact.id, 99999],
                    "template_slugs": [valid_template.slug, inactive_template.slug, "missing-template"],
                }
            ],
            [valid_fact],
            [valid_template, inactive_template],
        )

        self.assertEqual(normalized[0]["id"], "hearing-continuance")
        self.assertEqual(normalized[0]["confidence"], "medium")
        self.assertEqual(normalized[0]["supportingFactIds"], [valid_fact.id])
        self.assertEqual(normalized[0]["templateSlugs"], [valid_template.slug])
        self.assertEqual(normalized[0]["templateIds"], [valid_template.id])

    def test_template_variable_parser_resolves_dotted_paths_and_loop_aliases(self):
        variables = extract_template_variables_from_text(
            "{% for fact in selected_facts %}{{ fact.text }} {{ client.name.first }}{% endfor %}"
        )

        self.assertIn("selected_facts[i].text", variables)
        self.assertIn("client.name.first", variables)

    def test_template_variable_parser_normalizes_numeric_field_keys(self):
        variables = extract_template_variables_from_text(
            "Hearing under {{ fields.24_cfr_982554_if_hcvp_applicant }}."
        )

        self.assertIn('fields["24_cfr_982554_if_hcvp_applicant"]', variables)

    def test_body_variable_metadata_reports_parse_errors_without_raising(self):
        template = DocumentTemplate.objects.create(
            slug="malformed-body-metadata-test",
            title="Malformed body metadata",
            kind="motion",
        )
        block = TemplateBlock.objects.create(
            template=template,
            key="body",
            label="Body",
            block_type="argument",
            order=10,
            body="Broken {{ fields.unclosed",
        )

        metadata = block_variable_metadata(template, block)

        self.assertEqual(metadata["variables"]["all"], [])
        self.assertTrue(metadata["parseError"])

    def test_template_renderer_accepts_numeric_field_keys(self):
        service = ConstrainedDraftingService()
        context = SimpleNamespace(
            author_profile={},
            template_data={"24_cfr_982554_if_hcvp_applicant": "24 C.F.R. 982.554"},
            matter=SimpleNamespace(
                jurisdiction="Housing Court",
                client_name="Jane Tenant",
                external_id="CASE-1",
            ),
        )

        rendered = service.render_template_body(
            "Hearing under {{ fields.24_cfr_982554_if_hcvp_applicant }}.",
            context,
        )

        self.assertEqual(rendered, "Hearing under 24 C.F.R. 982.554.")

    def test_repository_docx_template_variables_are_classified(self):
        template, _created = DocumentTemplate.objects.update_or_create(
            slug="answer-counterclaims-cleveland",
            defaults={
                "title": "Answer and Counterclaims",
                "kind": "answer_counterclaims",
            },
        )
        block, _created = TemplateBlock.objects.update_or_create(
            template=template,
            key="caption",
            defaults={
                "label": "Court caption",
                "block_type": "caption",
                "order": 10,
                "body": "{{ court }}",
            },
        )

        metadata = block_variable_metadata(template, block)

        self.assertEqual(metadata["source"], "repository")
        self.assertIn("defendant", metadata["variables"]["providedBySystem"])
        self.assertIn("document.title", metadata["variables"]["providedBySystem"])
        self.assertEqual(metadata["variables"]["externalData"], [])

    def test_template_serializer_includes_word_template_variable_metadata(self):
        template, _created = DocumentTemplate.objects.update_or_create(
            slug="answer-counterclaims-cleveland",
            defaults={
                "title": "Answer and Counterclaims",
                "kind": "answer_counterclaims",
            },
        )
        template.blocks.all().delete()
        TemplateBlock.objects.create(
            template=template,
            key="custom",
            label="Custom",
            block_type="optional_clause",
            order=10,
            body="{{ client_preferred_name }} {{ matter.client_name }}",
        )

        data = template_to_dict(template, include_blocks=True)

        self.assertIn("wordTemplateVariables", data)
        self.assertIn("client_preferred_name", data["wordTemplateVariables"]["variables"]["externalData"])
        self.assertIn("matter.client_name", data["wordTemplateVariables"]["variables"]["providedBySystem"])

    def test_generated_text_normalizes_html_breaks(self):
        service = ConstrainedDraftingService()

        self.assertEqual(service.normalize_generated_text("Line one<br/>Line two<br>Line three"), "Line one\nLine two\nLine three")

    def test_generated_text_removes_markdown_artifacts_and_duplicate_heading(self):
        service = ConstrainedDraftingService()
        normalized = service.normalize_generated_text("## LAW AND ARGUMENT\n\n### I. Standard\nUse *Smith v. Jones*.")

        self.assertEqual(
            service._without_duplicate_heading(normalized, "Law and Argument"),
            "I. Standard\nUse Smith v. Jones.",
        )

    def test_missing_template_fields_render_as_visible_placeholders(self):
        service = ConstrainedDraftingService()
        context = SimpleNamespace(
            author_profile={},
            template_data={},
            matter=SimpleNamespace(
                jurisdiction="Housing Court",
                client_name="Jane Tenant",
                external_id="CASE-1",
            ),
        )

        rendered = service.render_template_body(
            "Defendant NAME requests a hearing on {{ fields.hearing_date }}; statutory text {{ fields.s }}upply.",
            context,
        )

        self.assertEqual(
            rendered,
            "Defendant Jane Tenant requests a hearing on [Hearing Date]; statutory text [s]upply.",
        )

    def test_export_docx_renders_sections_and_numbering_part(self):
        draft = SimpleNamespace(
            id=42,
            title="Test Draft",
            sections=[
                {
                    "label": "Facts",
                    "body": "First fact.\nSecond fact.",
                    "format": {"style": "numbered", "headingNumbering": "roman"},
                }
            ],
        )

        response = export_docx(draft)

        self.assertEqual(response["Content-Disposition"], 'attachment; filename="draft-42.docx"')
        self.assertEqual(response["Content-Type"], "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        with zipfile.ZipFile(BytesIO(response.content)) as archive:
            names = set(archive.namelist())
            self.assertIn("word/document.xml", names)
            self.assertIn("word/numbering.xml", names)
            self.assertIn("word/styles.xml", names)
            document_xml = archive.read("word/document.xml").decode("utf-8")
        self.assertIn("I. FACTS", document_xml)
        self.assertIn("First fact.", document_xml)
        self.assertIn("<w:numPr>", document_xml)

    def test_export_docx_can_restart_numbering(self):
        draft = SimpleNamespace(
            id=43,
            title="Restart Draft",
            sections=[
                {"label": "Facts", "body": "One", "format": {"style": "numbered"}},
                {"label": "Argument", "body": "Fresh one", "format": {"style": "numbered", "restartNumbering": True}},
            ],
        )

        response = export_docx(draft)

        with zipfile.ZipFile(BytesIO(response.content)) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8")
            numbering_xml = archive.read("word/numbering.xml").decode("utf-8")
        self.assertIn('<w:numId w:val="1"/>', document_xml)
        self.assertIn('<w:numId w:val="2"/>', document_xml)
        self.assertIn('<w:num w:numId="2">', numbering_xml)

    def test_docx_render_context_exposes_numbered_items_for_templates(self):
        matter = Matter.objects.create(
            external_id="CASE-NUMBERED",
            client_name="Jane Tenant",
            matter_type="Eviction",
            jurisdiction="Housing Court",
        )
        template = DocumentTemplate.objects.create(
            slug="numbered-context-template",
            title="Numbered Context",
            kind="motion",
        )
        TemplateBlock.objects.create(
            template=template,
            key="argument",
            label="Argument",
            block_type="argument",
            order=10,
            body="{{ blocks.argument.numbered_items }}",
        )
        session = DraftingSession.objects.create(
            mode="draft_from_template",
            matter=matter,
            template=template,
        )
        draft = DraftDocument.objects.create(
            session=session,
            title="Numbered Draft",
            sections=[
                {
                    "key": "argument",
                    "label": "Argument",
                    "body": "First point\nSecond point",
                    "format": {"style": "numbered"},
                }
            ],
            plain_text="Argument\nFirst point\nSecond point",
        )

        context = _docx_render_context(draft, draft.sections[0])

        self.assertTrue(context["blocks"]["argument"]["numbered"])
        self.assertEqual(context["blocks"]["argument"]["items"], ["First point", "Second point"])
        self.assertEqual(context["blocks"]["argument"]["numbered_items"], ["1. First point", "2. Second point"])

    def test_plan_selects_continuance_template_and_generates_without_author(self):
        matter = Matter.objects.create(
            external_id="CASE-CONTINUE",
            client_name="Jane Tenant",
            matter_type="Eviction",
            jurisdiction="Cleveland Municipal Court - Housing Division",
            summary="Tenant needs more time for rental assistance.",
        )
        dismissal, _created = DocumentTemplate.objects.update_or_create(
            slug="motion-dismissal-cleveland",
            defaults={
                "title": "Motion to Dismiss",
                "kind": "motion",
                "goal": "Ask the court to dismiss the case.",
                "negative_goal": "Do not use when the user asks only to continue or postpone a hearing.",
                "aliases": ["dismiss case"],
                "jurisdiction": "Cleveland Municipal Court - Housing Division",
            },
        )
        TemplateBlock.objects.update_or_create(
            template=dismissal,
            key="body",
            defaults={
                "label": "Body",
                "block_type": "argument",
                "order": 10,
                "body": "Dismiss this case.",
            },
        )
        continuance, _created = DocumentTemplate.objects.update_or_create(
            slug="motion-continuance-cleveland",
            defaults={
                "title": "Motion for Continuance",
                "kind": "motion",
                "goal": "Ask the court to postpone or continue a scheduled hearing, deadline, or proceeding.",
                "negative_goal": "Do not seek dismissal, judgment, or merits resolution unless separately requested.",
                "aliases": ["continue hearing", "postpone hearing", "more time", "adjournment"],
                "jurisdiction": "Cleveland Municipal Court - Housing Division",
            },
        )
        TemplateBlock.objects.update_or_create(
            template=continuance,
            key="motion-body",
            defaults={
                "label": "Motion body",
                "block_type": "argument",
                "order": 10,
                "body": "Defendant asks for a continuance.",
                "ai_fill_mode": "none",
            },
        )
        TemplateBlock.objects.update_or_create(
            template=continuance,
            key="motion-signature",
            defaults={
                "label": "Signature",
                "block_type": "signature",
                "order": 20,
                "body": "{{ advocate_name }}",
            },
        )
        session = DraftingSession.objects.create(
            mode="draft_from_template",
            matter=matter,
            goal="continue hearing because rental assistance is pending",
            instructions="Do not request dismissal.",
            author_profile={},
        )

        session = create_or_update_plan(session, {"goal": session.goal})
        drafts = create_drafts_from_plan(session)

        self.assertEqual(session.draft_plan["document_items"][0]["template_slug"], "motion-continuance-cleveland")
        self.assertNotEqual(session.draft_plan["document_items"][0]["template_slug"], "motion-dismissal-cleveland")
        self.assertEqual(len(drafts), 1)
        self.assertEqual(drafts[0].title, "Motion for Continuance")
        self.assertIn("motion-body", [section["key"] for section in drafts[0].sections])

    def test_plan_can_generate_multiple_documents(self):
        matter = Matter.objects.create(
            external_id="CASE-MULTI",
            client_name="Jane Tenant",
            matter_type="Eviction",
            jurisdiction="Housing Court",
            summary="Tenant needs an answer and more time.",
        )
        templates = []
        for slug, title, alias in [
            ("answer-counterclaims-cleveland", "Answer and Counterclaims", "answer eviction"),
            ("motion-continuance-cleveland", "Motion for Continuance", "more time"),
        ]:
            template, _created = DocumentTemplate.objects.update_or_create(
                slug=slug,
                defaults={
                    "title": title,
                    "kind": "motion",
                    "goal": title,
                    "aliases": [alias],
                },
            )
            TemplateBlock.objects.update_or_create(
                template=template,
                key=f"{slug}-body",
                defaults={
                    "label": "Body",
                    "block_type": "argument",
                    "order": 10,
                    "body": f"{title} body.",
                },
            )
            templates.append(template)
        session = DraftingSession.objects.create(
            mode="draft_from_template",
            matter=matter,
            selected_template_ids=[template.id for template in templates],
            goal="answer eviction and ask for more time",
        )

        session = create_or_update_plan(session, {"allowMultipleDocuments": True})
        drafts = create_drafts_from_plan(session)

        self.assertEqual(len(session.draft_plan["document_items"]), 2)
        self.assertEqual(len(drafts), 2)

    def test_known_template_can_create_plan_without_typed_goal(self):
        matter = Matter.objects.create(
            external_id="CASE-KNOWN-TEMPLATE",
            client_name="Jane Tenant",
            matter_type="Eviction",
            jurisdiction="Housing Court",
            summary="Tenant needs a prepared filing.",
        )
        template = DocumentTemplate.objects.create(
            slug="known-template-no-goal",
            title="Known Template",
            kind="motion",
            goal="Make the known template filing.",
            description="Prepared motion shell.",
        )
        TemplateBlock.objects.create(
            template=template,
            key="body",
            label="Body",
            block_type="argument",
            order=10,
            body="Known template body.",
        )
        session = DraftingSession.objects.create(
            mode="draft_from_template",
            matter=matter,
            selected_template_ids=[template.id],
            goal="",
            instructions="",
        )

        session = create_or_update_plan(session, {"selectedTemplateIds": [template.id]})

        self.assertEqual(session.draft_plan["summary"], "Make the known template filing.")
        self.assertEqual(session.draft_plan["document_items"][0]["template_slug"], "known-template-no-goal")
        self.assertEqual(session.draft_plan["document_items"][0]["drafting_instructions"], "Make the known template filing.")

    def test_generic_template_description_is_not_used_as_drafting_instructions(self):
        matter = Matter.objects.create(
            external_id="CASE-GENERIC-NOTE",
            client_name="Jane Tenant",
            matter_type="Eviction",
            jurisdiction="Housing Court",
            summary="Tenant needs a prepared filing.",
        )
        template = DocumentTemplate.objects.create(
            slug="legacy-content-library-template",
            title="Legacy Motion",
            kind="motion",
            description="Prepared from the maintained original Word template.",
            goal="",
        )
        TemplateBlock.objects.create(
            template=template,
            key="body",
            label="Body",
            block_type="argument",
            order=10,
            body="Legacy motion body.",
        )
        session = DraftingSession.objects.create(
            mode="draft_from_template",
            matter=matter,
            selected_template_ids=[template.id],
            goal="",
            instructions="",
        )

        session = create_or_update_plan(session, {"selectedTemplateIds": [template.id]})

        item = session.draft_plan["document_items"][0]
        self.assertEqual(item["goal"], "Draft Legacy Motion.")
        self.assertEqual(
            item["drafting_instructions"],
            "Use the selected Legacy Motion template structure and draft the active blocks with case-specific facts, requested relief, and reviewer-approved sources.",
        )
        self.assertNotIn("maintained original Word template", item["drafting_instructions"])

    def test_optional_missing_information_can_be_promoted_to_predraft_questions(self):
        plan = {
            "document_items": [
                {
                    "missing_information": [
                        {
                            "field": "hearing_date",
                            "question": "What is the current hearing date?",
                            "required_for_generation": False,
                        }
                    ]
                }
            ]
        }

        self.assertEqual(unanswered_missing_information(plan), [])
        self.assertEqual(len(unanswered_missing_information(plan, require_all=True)), 1)

    def test_plan_missing_information_answers_populate_template_fields(self):
        matter = Matter.objects.create(
            external_id="CASE-QUESTION-ANSWER",
            client_name="Jane Tenant",
            matter_type="Eviction",
            jurisdiction="Housing Court",
            summary="Tenant needs a continuance.",
        )
        template = DocumentTemplate.objects.create(
            slug="question-answer-template",
            title="Question Answer Template",
            kind="motion",
        )
        TemplateBlock.objects.create(
            template=template,
            key="body",
            label="Body",
            block_type="argument",
            order=10,
            body="The current hearing date is {{ fields.hearing_date }}.",
        )
        session = DraftingSession.objects.create(
            mode="draft_from_template",
            matter=matter,
            template=template,
            selected_template_ids=[template.id],
        )
        plan = {
            "summary": "Ask for more time.",
            "document_items": [
                {
                    "id": template.slug,
                    "template_slug": template.slug,
                    "template_id": template.id,
                    "title": template.title,
                    "selected_block_keys": ["body"],
                    "drafting_instructions": "Use the answered hearing date.",
                    "missing_information": [
                        {
                            "field": "hearing_date",
                            "question": "What is the current hearing date?",
                            "answer": "August 1, 2026",
                            "required_for_generation": False,
                        }
                    ],
                }
            ],
        }

        session = apply_plan_edits(session, {"draftPlan": plan})
        drafts = create_drafts_from_plan(session)

        self.assertEqual(session.template_data["hearing_date"], "August 1, 2026")
        self.assertIn("August 1, 2026", drafts[0].sections[0]["body"])

    def test_plan_missing_information_includes_unset_template_field_question(self):
        matter = Matter.objects.create(
            external_id="CASE-FIELD-QUESTION",
            client_name="Jane Tenant",
            matter_type="Eviction",
            jurisdiction="Housing Court",
        )
        template = DocumentTemplate.objects.create(
            slug="answer-plaintiff-field-test",
            title="Answer and Counterclaims",
            kind="answer_counterclaims",
        )
        TemplateBlock.objects.create(
            template=template,
            key="caption",
            label="Caption",
            block_type="caption",
            order=10,
            body="{{ plaintiff }} v. {{ defendant }}",
        )
        session = DraftingSession.objects.create(mode="draft_from_template", matter=matter, template=template)

        session = create_or_update_plan(session, {"selectedTemplateIds": [template.id]})

        questions = [item["question"] for item in session.missing_information]
        self.assertTrue(any("plaintiff name" in question.casefold() for question in questions), questions)

    def test_plan_missing_information_omits_already_answered_template_field(self):
        matter = Matter.objects.create(
            external_id="CASE-FIELD-ANSWERED",
            client_name="Jane Tenant",
            matter_type="Eviction",
            jurisdiction="Housing Court",
        )
        template = DocumentTemplate.objects.create(
            slug="answer-plaintiff-answered-test",
            title="Answer and Counterclaims",
            kind="answer_counterclaims",
        )
        TemplateBlock.objects.create(
            template=template,
            key="caption",
            label="Caption",
            block_type="caption",
            order=10,
            body="{{ plaintiff }} v. {{ defendant }}",
        )
        session = DraftingSession.objects.create(
            mode="draft_from_template",
            matter=matter,
            template=template,
            template_data={"plaintiff_name": "Acme Realty LLC"},
        )

        session = create_or_update_plan(session, {"selectedTemplateIds": [template.id]})

        questions = [item["question"] for item in session.missing_information]
        self.assertFalse(any("plaintiff name" in question.casefold() for question in questions), questions)


class SessionTemplateDataEndpointTests(TestCase):
    def test_update_session_template_data_merges_values(self):
        user = get_user_model().objects.create_user(username="filler", password="pass", is_superuser=True)
        matter = Matter.objects.create(
            external_id="CASE-TEMPLATE-DATA",
            client_name="Jane Tenant",
            matter_type="Eviction",
            jurisdiction="Housing Court",
        )
        session = DraftingSession.objects.create(
            mode="draft_from_template",
            matter=matter,
            template_data={"existing": "value"},
        )

        self.client.login(username="filler", password="pass")
        url = reverse("api_session_template_data", args=[session.id])
        response = self.client.post(
            url,
            data='{"templateData": {"plaintiff_name": "Acme Realty LLC"}}',
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 200)
        session.refresh_from_db()
        self.assertEqual(session.template_data["plaintiff_name"], "Acme Realty LLC")
        self.assertEqual(session.template_data["existing"], "value")

    def test_export_docx_removes_xml_forbidden_characters(self):
        draft = SimpleNamespace(
            id=44,
            title="Bad\x0bTitle",
            sections=[
                {
                    "label": "Facts",
                    "body": "Text with invalid XML control \x0b and safe <xml> characters.",
                    "format": {},
                }
            ],
        )

        response = export_docx(draft)

        with zipfile.ZipFile(BytesIO(response.content)) as archive:
            document_xml = archive.read("word/document.xml")
            core_xml = archive.read("docProps/core.xml")
        ElementTree.fromstring(document_xml)
        ElementTree.fromstring(core_xml)
        self.assertNotIn(b"\x0b", document_xml)
        self.assertNotIn(b"\x0b", core_xml)

    def test_export_docx_renders_and_composes_uploaded_block_templates(self):
        temp_dir = tempfile.TemporaryDirectory()
        media_dir = tempfile.TemporaryDirectory()
        self.addCleanup(temp_dir.cleanup)
        self.addCleanup(media_dir.cleanup)
        source_dir = Path(temp_dir.name)

        style_source = source_dir / "styles.docx"
        style_doc = Document()
        style_doc.add_paragraph("Style source body should be cleared.")
        style_doc.save(style_source)

        caption_source = source_dir / "caption.docx"
        caption_doc = Document()
        caption_doc.add_paragraph("Caption for {{ defendant }}")
        caption_doc.add_paragraph("Reviewed body: {{ section.body }}")
        caption_doc.add_paragraph("Numbered values: {{ blocks.caption.numbered_items | join('; ') }}")
        caption_doc.save(caption_source)

        with override_settings(MEDIA_ROOT=media_dir.name):
            matter = Matter.objects.create(
                external_id="24-CV-100",
                client_name="Jane Tenant",
                matter_type="Eviction",
                jurisdiction="Housing Court",
            )
            template = DocumentTemplate.objects.create(
                title="Answer",
                slug="answer-docx-template-test",
                kind="answer_counterclaims",
                style_template=File(style_source.open("rb"), name="styles.docx"),
            )
            TemplateBlock.objects.create(
                template=template,
                key="caption",
                label="Caption",
                block_type="caption",
                order=10,
                body="{{ court }}",
                docx_template=File(caption_source.open("rb"), name="caption.docx"),
            )
            session = DraftingSession.objects.create(
                mode="draft_from_template",
                matter=matter,
                template=template,
                author_profile={"displayName": "Ada Advocate"},
            )
            draft = DraftDocument.objects.create(
                session=session,
                title="Answer",
                sections=[{"key": "caption", "label": "Caption", "body": "Edited caption text.\nSecond caption point.", "format": {"style": "numbered"}}],
                plain_text="CAPTION\nEdited caption text.",
            )

            response = export_docx(draft)

        self.assertEqual(response["Content-Type"], "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        with zipfile.ZipFile(BytesIO(response.content)) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8")
        ElementTree.fromstring(document_xml)
        self.assertIn("Jane Tenant", document_xml)
        self.assertIn("Edited caption text.", document_xml)
        self.assertIn("1. Edited caption text.", document_xml)
        self.assertIn("2. Second caption point.", document_xml)
        self.assertNotIn("Style source body should be cleared.", document_xml)

    def test_export_docx_uses_repository_default_block_templates(self):
        matter = Matter.objects.create(
            external_id="24-CV-101",
            client_name="John Tenant",
            matter_type="Eviction",
            jurisdiction="Housing Court",
        )
        template, _created = DocumentTemplate.objects.update_or_create(
            slug="answer-counterclaims-cleveland",
            defaults={
                "title": "Answer and Counterclaims",
                "kind": "answer_counterclaims",
            },
        )
        template.blocks.all().delete()
        TemplateBlock.objects.create(
            template=template,
            key="caption",
            label="Court caption",
            block_type="caption",
            order=10,
            body="{{ court }}",
        )
        TemplateBlock.objects.create(
            template=template,
            key="relief",
            label="Prayer for relief",
            block_type="relief",
            order=90,
            body="{{ section.body }}",
        )
        session = DraftingSession.objects.create(
            mode="draft_from_template",
            matter=matter,
            template=template,
            author_profile={"displayName": "Ada Advocate"},
        )
        draft = DraftDocument.objects.create(
            session=session,
            title="Answer and Counterclaims",
            sections=[
                {"key": "caption", "label": "Court caption", "body": "Caption text."},
                {"key": "relief", "label": "Prayer for relief", "body": "dismiss the complaint"},
            ],
            plain_text="CAPTION\nCaption text.\n\nRELIEF\ndismiss the complaint",
        )

        response = export_docx(draft)

        with zipfile.ZipFile(BytesIO(response.content)) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8")
        ElementTree.fromstring(document_xml)
        self.assertIn("John Tenant", document_xml)
        self.assertIn("Prayer for Relief", document_xml)
        self.assertIn("dismiss the complaint", document_xml)

