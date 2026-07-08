import io
import re
import tempfile
import zipfile
from html import escape

from django.http import HttpResponse
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docxcompose.composer import Composer
from docxtpl import DocxTemplate

from apps.templates_app.word_templates import (
    block_template_path,
    has_word_template_assets,
    style_template_path,
)
from apps.templates_app.content_library import full_template_path
from apps.templates_app.template_variables import normalize_docxtpl_blocks, template_field_values
from apps.templates_app.jinja_filters import listify, template_environment


CONTENT_TYPES = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/word/numbering.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml"/>
  <Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>
  <Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>
  <Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/>
</Types>
"""

ROOT_RELS = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
  <Relationship Id="rId2" Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties" Target="docProps/core.xml"/>
  <Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties" Target="docProps/app.xml"/>
</Relationships>
"""

APP_PROPS = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties" xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">
  <Application>Legal Drafting Tool</Application>
</Properties>
"""

CORE_PROPS = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:dcterms="http://purl.org/dc/terms/" xmlns:dcmitype="http://purl.org/dc/dcmitype/" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">
  <dc:title>{title}</dc:title>
  <dc:creator>Legal Drafting Tool</dc:creator>
</cp:coreProperties>
"""

STYLES = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:style w:type="paragraph" w:default="1" w:styleId="Normal">
    <w:name w:val="Normal"/>
    <w:qFormat/>
  </w:style>
  <w:style w:type="paragraph" w:styleId="Heading1">
    <w:name w:val="heading 1"/>
    <w:basedOn w:val="Normal"/>
    <w:next w:val="Normal"/>
    <w:uiPriority w:val="9"/>
    <w:qFormat/>
    <w:pPr><w:spacing w:before="240" w:after="120"/></w:pPr>
    <w:rPr><w:b/><w:sz w:val="28"/></w:rPr>
  </w:style>
</w:styles>
"""


def _xml_text(value):
    text = "" if value is None else str(value)
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)
    return escape(text)


def _paragraph(text, *, style=None, num_id=None, ilvl=0):
    props = []
    if style:
        props.append(f'<w:pStyle w:val="{style}"/>')
    if num_id:
        props.append(f'<w:numPr><w:ilvl w:val="{ilvl}"/><w:numId w:val="{num_id}"/></w:numPr>')
    ppr = f"<w:pPr>{''.join(props)}</w:pPr>" if props else ""
    runs = "".join(
        f'<w:r><w:t xml:space="preserve">{_xml_text(part)}</w:t></w:r><w:br/>'
        for part in text.split("\n")[:-1]
    )
    runs += f'<w:r><w:t xml:space="preserve">{_xml_text(text.split(chr(10))[-1])}</w:t></w:r>'
    return f"<w:p>{ppr}{runs}</w:p>"


def _numbering(num_count=1):
    nums = "\n".join(
        f'  <w:num w:numId="{num_id}"><w:abstractNumId w:val="1"/></w:num>'
        for num_id in range(1, num_count + 1)
    )
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:abstractNum w:abstractNumId="1"><w:lvl w:ilvl="0"><w:start w:val="1"/><w:numFmt w:val="decimal"/><w:lvlText w:val="%1."/></w:lvl></w:abstractNum>
{nums}
</w:numbering>
"""


def _document_rels():
    return """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rIdNumbering" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering" Target="numbering.xml"/>
  <Relationship Id="rIdStyles" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/>
</Relationships>
"""


def _doc_xml(draft):
    body = []
    paragraph_numbering = False
    num_id = 1
    max_num_id = 1
    for section in draft.sections or []:
        section_format = section.get("format") or {}
        heading = section.get("label", "")
        heading_numbering = section_format.get("headingNumbering")
        if heading_numbering == "roman":
            roman_index = _roman(len([item for item in body if 'Heading1' in item]) + 1)
            heading = f"{roman_index}. {heading}"
        body.append(_paragraph(heading.upper(), style="Heading1"))
        style = section_format.get("style")
        if style == "numbered" and section_format.get("restartNumbering") and paragraph_numbering:
            num_id += 1
            max_num_id = max(max_num_id, num_id)
        paragraphs = [part.strip() for part in re.split(r"\n{2,}|\n", section.get("body", "")) if part.strip()]
        for paragraph in paragraphs:
            cleaned = re.sub(r"^\d+\.\s*", "", paragraph)
            if style == "numbered":
                paragraph_numbering = True
                body.append(_paragraph(cleaned, num_id=num_id))
            else:
                body.append(_paragraph(cleaned))
    sect_props = '<w:sectPr><w:pgSz w:w="12240" w:h="15840"/><w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440"/></w:sectPr>'
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>{''.join(body)}{sect_props}</w:body>
</w:document>
""", max_num_id


def _roman(number):
    values = [(10, "X"), (9, "IX"), (5, "V"), (4, "IV"), (1, "I")]
    result = []
    for value, numeral in values:
        while number >= value:
            result.append(numeral)
            number -= value
    return "".join(result)


def _draft_template(draft):
    session = getattr(draft, "session", None)
    if not session:
        return None
    editor_state = getattr(draft, "editor_state", None) or {}
    template_id = editor_state.get("templateId")
    template_slug = editor_state.get("templateSlug")
    if template_id or template_slug:
        from apps.templates_app.models import DocumentTemplate

        query = DocumentTemplate.objects.prefetch_related("blocks")
        if template_id:
            template = query.filter(id=template_id).first()
            if template:
                return template
        if template_slug:
            template = query.filter(slug=template_slug).first()
            if template:
                return template
    return getattr(session, "template", None)


def _has_word_template_assets(draft):
    return has_word_template_assets(_draft_template(draft))


def _matter_context(matter):
    return {
        "external_id": getattr(matter, "external_id", ""),
        "client_name": getattr(matter, "client_name", ""),
        "matter_type": getattr(matter, "matter_type", ""),
        "jurisdiction": getattr(matter, "jurisdiction", ""),
        "posture": getattr(matter, "posture", ""),
        "risk": getattr(matter, "risk", ""),
        "summary": getattr(matter, "summary", ""),
        "source_system": getattr(matter, "source_system", ""),
    }


def _author_context(author):
    author = author or {}
    contact = "\n".join(
        item
        for item in [
            author.get("organization", ""),
            author.get("address", ""),
            author.get("phone", ""),
            author.get("email", ""),
        ]
        if item
    )
    return {
        "display_name": author.get("displayName") or "Advocate",
        "signoff": author.get("signoff") or "Respectfully submitted,",
        "salutation": author.get("salutation") or "",
        "organization": author.get("organization") or "",
        "email": author.get("email") or "",
        "phone": author.get("phone") or "",
        "address": author.get("address") or "",
        "contact": contact,
        "signature_image": author.get("signatureImage") or "",
    }


def _selected_facts(session):
    from apps.matters.models import MatterFact

    facts = MatterFact.objects.filter(id__in=session.selected_fact_ids)
    return [
        {
            "slug": fact.slug,
            "title": fact.title,
            "text": fact.text,
            "source_label": fact.source_label,
            "confidence": fact.confidence,
        }
        for fact in facts
    ]


def _docx_render_context(draft, section):
    session = draft.session
    matter = session.matter
    template = _draft_template(draft)
    author = _author_context(session.author_profile)
    matter_data = _matter_context(matter)
    sections = draft.sections or []
    block_context = {}
    if template:
        block_context = {
            block.key: {"body": "", "paragraphs": [], "items": [], "numbered": False, "format": {}, "numbered_items": []}
            for block in template.blocks.all()
        }
    for draft_section in sections:
        section_format = draft_section.get("format") or {}
        parts = [
            part.strip()
            for part in re.split(r"\n{2,}|\n", draft_section.get("body", ""))
            if part.strip()
        ]
        items = [re.sub(r"^\d+[.)]\s*", "", part) for part in parts]
        numbered = section_format.get("style") == "numbered"
        block_context[draft_section.get("key", "")] = {
            "body": draft_section.get("body", ""),
            "paragraphs": parts,
            "items": items,
            "numbered": numbered,
            "format": section_format,
            "numbered_items": [f"{index}. {item}" for index, item in enumerate(items, start=1)] if numbered else items,
        }
    fields = template_field_values(session.template_data)
    client = {
        "name": matter_data["client_name"],
        "pronouns": (session.template_data or {}).get("client_pronouns", ""),
        "title": (session.template_data or {}).get("client_title", ""),
    }
    household = listify((session.template_data or {}).get("other_occupants", []))
    context = {
        "document": {
            "title": draft.title,
            "plain_text": draft.plain_text,
            "sections": sections,
        },
        "section": section,
        "matter": matter_data,
        "client": client,
        "household": household,
        "author": author,
        "selected_facts": _selected_facts(session),
        "selected_curated_facts": session.selected_curated_facts,
        "selected_sources": session.selected_source_results,
        "instructions": session.instructions,
        "fields": fields,
        "blocks": block_context,
        "court": matter_data["jurisdiction"] or getattr(template, "jurisdiction", ""),
        "plaintiff": fields["plaintiff_name"],
        "defendant": matter_data["client_name"],
        "case_number": (session.template_data or {}).get("court_case_number") or "[Court Case Number]",
        "advocate_name": author["display_name"],
        "advocate_signoff": author["signoff"],
        "advocate_salutation": author["salutation"],
        "advocate_organization": author["organization"],
        "advocate_email": author["email"],
        "advocate_phone": author["phone"],
        "advocate_address": author["address"],
        "advocate_contact": author["contact"],
        "advocate_signature_image": author["signature_image"],
    }
    return context


def _render_docx_template(template_path, context, output_path):
    class NormalizingDocxTemplate(DocxTemplate):
        def patch_xml(self, source_xml):
            return normalize_docxtpl_blocks(super().patch_xml(source_xml))

    doc = NormalizingDocxTemplate(template_path)
    doc.render(context, jinja_env=template_environment())
    doc.save(output_path)


def _write_text_section_doc(section, output_path):
    doc = Document()
    label = section.get("label", "")
    if label:
        doc.add_heading(label, level=1)
    for paragraph in re.split(r"\n{2,}|\n", section.get("body", "")):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
    doc.save(output_path)


def _remove_editorial_caption_label(path, block, context=None):
    """Replace an ingestion-only caption marker with a minimal filing caption."""
    if not block or not block.label.casefold().startswith("case caption"):
        return
    document = Document(path)
    first = next((paragraph for paragraph in document.paragraphs if paragraph.text.strip()), None)
    first_text = re.sub(r"\s+", " ", first.text).strip().casefold() if first else ""
    if not first or not first_text.startswith("case caption"):
        return
    context = context or {}
    fields = context.get("fields") or template_field_values()
    court = str(context.get("court") or "").strip()
    if "court" not in court.casefold():
        rendered_text = " ".join(paragraph.text for paragraph in document.paragraphs)
        inferred = re.search(r"\b([A-Z][A-Za-z]+ Municipal (?:Housing )?Court)\b", rendered_text)
        court = inferred.group(1) if inferred else "[Court]"
    plaintiff = str(context.get("plaintiff") or fields["plaintiff_name"])
    defendant = str(context.get("defendant") or "[Defendant Name]")
    case_number = str(context.get("case_number") or "[Case Number]")
    title = re.sub(r"^case caption\s*[-–—:]\s*", "", block.label, flags=re.IGNORECASE)

    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_run = heading.add_run(court.upper())
    heading_run.bold = True
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(0, 0).text = f"{plaintiff}\nPlaintiff / Landlord\n\nv.\n\n{defendant}\nDefendant / Tenant"
    table.cell(0, 1).text = f"Case No. {case_number}\n\n{title}"
    first._element.addprevious(heading._element)
    first._element.addprevious(table._element)
    first._element.getparent().remove(first._element)
    document.save(path)


def _clear_document_body(document):
    body = document._body._element
    section_properties = None
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            section_properties = child
        body.remove(child)
    if section_properties is not None:
        body.append(section_properties)


def _composed_docx(draft):
    template = _draft_template(draft)
    block_by_key = {block.key: block for block in template.blocks.all()}
    sections = draft.sections or []
    output = io.BytesIO()
    with tempfile.TemporaryDirectory() as tmpdir:
        section_paths = []
        for index, section in enumerate(sections):
            block = block_by_key.get(section.get("key"))
            section_path = f"{tmpdir}/section-{index}.docx"
            selected_block_template_path = None
            if block and block.docx_template:
                selected_block_template_path = block.docx_template.path
            elif block:
                selected_block_template_path = block_template_path(template, block)
            if selected_block_template_path:
                context = _docx_render_context(draft, section)
                _render_docx_template(selected_block_template_path, context, section_path)
                _remove_editorial_caption_label(section_path, block, context)
            else:
                _write_text_section_doc(section, section_path)
            section_paths.append(section_path)

        master_style_template_path = style_template_path(template)
        if master_style_template_path and template.replace_child_styles:
            master = Document(master_style_template_path)
            _clear_document_body(master)
        elif section_paths:
            master = Document(section_paths.pop(0))
        else:
            master = Document()

        composer = Composer(master)
        for section_path in section_paths:
            composer.append(Document(section_path))
        composer.save(output)
    return output.getvalue()


def _full_template_docx(draft, template_path):
    output = io.BytesIO()
    with tempfile.TemporaryDirectory() as tmpdir:
        rendered_path = f"{tmpdir}/rendered.docx"
        _render_docx_template(template_path, _docx_render_context(draft, {}), rendered_path)
        with open(rendered_path, "rb") as stream:
            output.write(stream.read())
    return output.getvalue()


def export_docx(draft):
    selected_full_template = full_template_path(_draft_template(draft))
    if selected_full_template:
        response = HttpResponse(
            _full_template_docx(draft, selected_full_template),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = f'attachment; filename="draft-{draft.id}.docx"'
        return response
    if _has_word_template_assets(draft):
        response = HttpResponse(
            _composed_docx(draft),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = f'attachment; filename="draft-{draft.id}.docx"'
        return response

    document_xml, num_count = _doc_xml(draft)
    output = io.BytesIO()
    with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as docx:
        docx.writestr("[Content_Types].xml", CONTENT_TYPES)
        docx.writestr("_rels/.rels", ROOT_RELS)
        docx.writestr("docProps/app.xml", APP_PROPS)
        docx.writestr("docProps/core.xml", CORE_PROPS.format(title=_xml_text(draft.title)))
        docx.writestr("word/document.xml", document_xml)
        docx.writestr("word/_rels/document.xml.rels", _document_rels())
        docx.writestr("word/numbering.xml", _numbering(num_count))
        docx.writestr("word/styles.xml", STYLES)
    response = HttpResponse(
        output.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = f'attachment; filename="draft-{draft.id}.docx"'
    return response


def export_plain_text(draft):
    response = HttpResponse(draft.plain_text, content_type="text/plain; charset=utf-8")
    filename = f"draft-{draft.id}.txt"
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response
