"""Preparing the client advice-letter catalog.

Advice letters are built differently from filings. A filing is one maintained
document with fill-ins; an advice letter is a fixed wrapper plus however many
short sections the client's situation calls for, chosen per tenant. The
Cleveland working group maintains them that way deliberately -- a spreadsheet
catalog, a Model Letter, and a folder of interchangeable sub-sections -- so the
catalog is modeled as its own thing rather than squeezed into the litigation
template list.

Three things about the maintained sources need handling before the text is
usable:

- Nine sub-sections carry unresolved tracked changes. Reading them naively
  yields shredded prose ("The Notice landlord MUST specific conspicuous"),
  because a run inside `w:ins` is not a direct child of the paragraph.
- Five sub-sections repeat the whole Model Letter wrapper. Ingesting those as
  written would print the greeting and the closing several times in an
  assembled letter.
- Sections point at each other with authoring notes such as
  "[Insert next defense/advice]". Those are composition slots, not text.
"""

from __future__ import annotations

import hashlib
import re
import zipfile
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from xml.etree import ElementTree as ET

import yaml
from docx import Document
from django.utils.text import slugify
from lxml import etree

from apps.templates_app.placeholders import convert_editor_state, convert_paragraph, convert_text
from apps.validation.copyedit import copyedit_lines
from apps.validation.readability import check_readability


W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
SHEET_NS = "http://schemas.openxmlformats.org/spreadsheetml/2006/main"

CATALOG_VERSION = 1

ROLE_INTRO = "intro"
ROLE_BODY = "body"
ROLE_CLOSING = "closing"

STATUS_READY = "ready"
STATUS_NEEDS_REVIEW = "needs_review"
STATUS_AI_DRAFTED = "ai_drafted"
STATUS_STUB = "stub"

# Lines that belong to the Model Letter wrapper rather than to a section.
WRAPPER_OPENERS = (
    "thank you for asking legal aid",
    "we’re sorry, but legal aid does not have enough lawyers",
    "we're sorry, but legal aid does not have enough lawyers",
)
WRAPPER_CLOSERS = (
    "closing your file",
    "closing your intake",
    "i hope this advice is helpful",
    "if you need legal help in the future",
)
WRAPPER_ADDRESS = (
    "client name",
    "street address",
    "city, oh zip",
    "email address",
    "case caption",
    "signature",
    "attorney name",
    "attorney",
    "dear mx.",
    "sincerely",
    "sincerely,",
    "re:",
    "delivered by email",
    "attached:",
)
DATE_LINE_RE = re.compile(r"^[A-Z][a-z]+ \d{1,2}, \d{4}$")

# "[Insert next defense/advice]", "[INSERT - Attend Hearing by Zoom]"
SLOT_RE = re.compile(r"\[\s*(?:insert|section to add)\b[^\]]*\]", re.I)
# "[If notice is proper]" opens a variant that runs to the next variant marker.
VARIANT_RE = re.compile(r"^\[\s*if\s+([^\]]+)\]\s*", re.I)

STUB_WORD_COUNT = 120

# "(216) xxx-xxxx" and "216-xxx-xxxx" stand in for the advocate's own number.
ADVOCATE_PHONE_RE = re.compile(r"\(?\d{3}\)?[\s.-]?x{3}[\s.-]?x{4}", re.I)
# "call paralegal Name Name at ..." is an authoring note for a real name.
PLACEHOLDER_NAME_RE = re.compile(r"\bName Name\b")


@dataclass
class SectionDraft:
    slug: str
    title: str
    role: str = ROLE_BODY
    topic: str = ""
    letter_type: str = "brief_advice"
    region: str = ""
    cleveland_specific: bool = False
    status: str = STATUS_READY
    paragraphs: list = field(default_factory=list)
    slots: list = field(default_factory=list)
    variants: list = field(default_factory=list)
    fields: list = field(default_factory=list)
    notes: list = field(default_factory=list)
    source_file: str = ""
    source_sha256: str = ""
    tracked_changes: int = 0
    comments: int = 0
    copyedit: dict = field(default_factory=dict)
    editor_state: dict = field(default_factory=dict)

    @property
    def body(self):
        return "\n".join(self.paragraphs)

    @property
    def word_count(self):
        return len(re.findall(r"[A-Za-z']+", self.body))


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


# ---------------------------------------------------------------- tracked changes


def tracked_change_counts(path: Path) -> tuple[int, int]:
    with zipfile.ZipFile(path) as archive:
        xml = archive.read("word/document.xml").decode("utf-8", errors="ignore")
    return xml.count("<w:ins ") + xml.count("<w:del "), xml.count("<w:commentRangeStart")


def accept_tracked_changes(document):
    """Apply the editor's marks so the prepared text is the accepted text.

    An insertion's runs live inside `w:ins`, which python-docx does not treat as
    paragraph content, so leaving them wrapped drops the words entirely. A
    deletion is removed outright rather than kept as `w:delText`.

    Returns the indexes of the paragraphs that carried an edit. Accepting is
    faithful, not corrective: where an editor deleted more than they replaced,
    the accepted text is grammatically broken. Those paragraphs are exactly the
    ones a copy-editor has to read, so their positions are reported rather than
    discarded.
    """
    body = document._body._element
    paragraphs = list(body.iter(f"{W}p"))
    touched = {
        index
        for index, paragraph in enumerate(paragraphs)
        if paragraph.find(f".//{W}ins") is not None or paragraph.find(f".//{W}del") is not None
    }
    for deletion in list(body.iter(f"{W}del")):
        parent = deletion.getparent()
        if parent is not None:
            parent.remove(deletion)
    for insertion in list(body.iter(f"{W}ins")):
        parent = insertion.getparent()
        if parent is None:
            continue
        index = list(parent).index(insertion)
        for offset, child in enumerate(list(insertion)):
            parent.insert(index + offset, child)
        parent.remove(insertion)
    # Review artifacts should not survive into a prepared section.
    for tag in ("commentRangeStart", "commentRangeEnd", "commentReference"):
        for node in list(body.iter(f"{W}{tag}")):
            parent = node.getparent()
            if parent is not None:
                parent.remove(node)
    return document, touched


# ---------------------------------------------------------------- wrapper


def _normalized(text):
    return " ".join((text or "").split()).casefold()


def _title_key(text):
    """Compare titles ignoring punctuation, so "3 Day:" matches "3 Day -"."""
    return re.sub(r"[^a-z0-9]+", " ", (text or "").casefold()).strip()


def substitute_known_placeholders(text: str) -> str:
    text = ADVOCATE_PHONE_RE.sub("{{ advocate_phone }}", text)
    return PLACEHOLDER_NAME_RE.sub("{{ fields.paralegal_name }}", text)


def is_wrapper_line(text: str) -> bool:
    normalized = _normalized(text)
    if not normalized:
        return False
    if DATE_LINE_RE.match(text.strip()):
        return True
    if any(normalized.startswith(opener) for opener in WRAPPER_OPENERS):
        return True
    if any(normalized.startswith(closer) for closer in WRAPPER_CLOSERS):
        return True
    return any(normalized.startswith(entry) for entry in WRAPPER_ADDRESS)


def strip_wrapper(paragraphs: list[str]) -> tuple[list[str], bool]:
    """Drop Model Letter scaffolding, keeping only the section's own advice."""
    kept = [line for line in paragraphs if not is_wrapper_line(line)]
    return kept, len(kept) != len(paragraphs)


# ---------------------------------------------------------------- extraction


def strip_slots(document) -> list[str]:
    """Remove composition notes before anything can mistake them for fill-ins.

    "[INSERT - Attend Hearing by Zoom]" names another section. Left in place it
    is indistinguishable from "[DATE]" to the placeholder converter, which
    turned it into `fields.insert_attend_hearing_zoom` -- a variable nobody can
    fill, standing where a whole section belongs.
    """
    found = []
    for paragraph in list(document.paragraphs):
        text = paragraph.text
        if not SLOT_RE.search(text):
            continue
        found.extend(match.group(0) for match in SLOT_RE.finditer(text))
        if SLOT_RE.fullmatch(text.strip()):
            parent = paragraph._p.getparent()
            if parent is not None:
                parent.remove(paragraph._p)
            continue
        # Partial match: clear the slot from the runs it spans, keeping the rest.
        cursor = 0
        spans = [(match.start(), match.end()) for match in SLOT_RE.finditer(text)]
        for run in paragraph.runs:
            start, end = cursor, cursor + len(run.text)
            cursor = end
            kept = "".join(
                character
                for offset, character in enumerate(run.text, start=start)
                if not any(span_start <= offset < span_end for span_start, span_end in spans)
            )
            if kept != run.text:
                run.text = kept
    return found


# Lexical's text-format bitmask: bold, italic, and underline. Keeping this
# conversion here means the maintained DOCX remains the formatting authority;
# the browser receives a normal Lexical JSON state rather than a second hand-
# authored representation of the same section.
LEXICAL_BOLD = 1
LEXICAL_ITALIC = 1 << 1
LEXICAL_UNDERLINE = 1 << 3


def _lexical_format(run):
    value = 0
    if run.bold:
        value |= LEXICAL_BOLD
    if run.italic:
        value |= LEXICAL_ITALIC
    if run.underline:
        value |= LEXICAL_UNDERLINE
    return value


def _lexical_text_node(text, format_value=0):
    return {
        "detail": 0,
        "format": format_value,
        "mode": "normal",
        "style": "",
        "text": text,
        "type": "text",
        "version": 1,
    }


def _reflow_formatted_runs(runs, text):
    """Fit a changed paragraph onto its source formatting boundaries.

    The normal ingest path keeps the source text byte-for-byte (apart from
    known placeholders), but copy-editing or a generated replacement can alter
    its length. A proportional fallback is preferable to flattening a whole
    paragraph to the first run's style.
    """
    if not text:
        return []
    source_length = sum(len(run_text) for run_text, _format in runs)
    if not source_length:
        return [_lexical_text_node(text)]
    children = []
    source_cursor = 0
    target_length = len(text)
    for run_text, format_value in runs:
        start = round(source_cursor * target_length / source_length)
        source_cursor += len(run_text)
        end = round(source_cursor * target_length / source_length)
        if end > start:
            children.append(_lexical_text_node(text[start:end], format_value))
    return children or [_lexical_text_node(text)]


def _lexical_paragraph(paragraph, text_override=""):
    """Serialize one DOCX paragraph, retaining inline run formatting."""
    source_runs = [
        (run.text, _lexical_format(run))
        for run in paragraph.runs
        if run.text
    ]
    source_text = "".join(run_text for run_text, _format in source_runs)
    cleaned_source_text = substitute_known_placeholders(SLOT_RE.sub("", source_text))
    desired = text_override
    prefix = ""
    if desired.startswith("- ") and not cleaned_source_text.startswith("- "):
        prefix = "- "
        desired = desired[2:]

    if not desired and not cleaned_source_text.strip():
        children = []
    elif desired == source_text:
        children = [_lexical_text_node(text, format_value) for text, format_value in source_runs]
    elif desired == cleaned_source_text:
        cleaned_runs = []
        for run_text, format_value in source_runs:
            run_text = substitute_known_placeholders(SLOT_RE.sub("", run_text))
            if run_text:
                cleaned_runs.append((run_text, format_value))
        children = [_lexical_text_node(text, format_value) for text, format_value in cleaned_runs]
    elif desired == source_text.strip():
        leading = len(source_text) - len(source_text.lstrip())
        trailing = len(source_text) - len(source_text.rstrip())
        trimmed = source_text[leading : len(source_text) - trailing if trailing else None]
        children = _reflow_formatted_runs(source_runs, trimmed)
    elif desired:
        children = _reflow_formatted_runs(source_runs, desired)
    else:
        children = []

    if prefix:
        children.insert(0, _lexical_text_node(prefix))
    return {
        "children": children,
        "direction": "ltr",
        "format": "",
        "indent": 0,
        "type": "paragraph",
        "version": 1,
    }


def lexical_editor_state(entries):
    """Create the initial Lexical state for ``[(paragraph, text), ...]``."""
    children = [_lexical_paragraph(paragraph, text) for paragraph, text in entries]
    return {
        "root": {
            "children": children,
            "direction": "ltr",
            "format": "",
            "indent": 0,
            "type": "root",
            "version": 1,
        }
    }


def _section_paragraph_entries(document, title):
    """Keep the section's paragraphs, including intentional blank paragraphs."""
    paragraphs = [paragraph for paragraph in document.paragraphs if not is_wrapper_line(paragraph.text)]
    nonempty = [paragraph for paragraph in paragraphs if paragraph.text.strip()]
    if nonempty and _title_key(nonempty[0].text.strip()) == _title_key(title):
        paragraphs.remove(nonempty[0])

    entries = []
    for paragraph in paragraphs:
        text = paragraph.text
        if not text.strip():
            entries.append((paragraph, ""))
            continue
        cleaned = substitute_known_placeholders(SLOT_RE.sub("", text))
        if paragraph._p.find(f"{W}pPr/{W}numPr") is not None:
            cleaned = f"- {cleaned.lstrip()}"
        if cleaned:
            entries.append((paragraph, cleaned))

    # Wrapper removal can leave the blank paragraph that separated the wrapper
    # from the section at either edge. It is not the section's own spacing.
    while entries and not entries[0][1]:
        entries.pop(0)
    while entries and not entries[-1][1]:
        entries.pop()
    return entries


def _paragraph_lines(document):
    lines = []
    for paragraph in document.paragraphs:
        conversion = convert_paragraph(paragraph, f"section_{len(lines) + 1}")
        text = paragraph.text.strip()
        if not text:
            continue
        style = (paragraph.style.name if paragraph.style else "") or ""
        numbered = paragraph._p.find(f"{W}pPr/{W}numPr") is not None
        lines.append((text, style, numbered, sorted(conversion.fields)))
    return lines


def extract_section(path: Path, *, slug="", title="") -> SectionDraft:
    """Read one maintained sub-section into a prepared draft."""
    tracked, comments = tracked_change_counts(path)
    document, touched = accept_tracked_changes(Document(path))
    slots = strip_slots(document)
    rows = _paragraph_lines(document)

    fields: set[str] = set()
    paragraphs = []
    edited_positions = set()
    for index, (text, _style, numbered, found) in enumerate(rows):
        fields.update(found)
        paragraphs.append(f"- {text}" if numbered else text)
        if index in touched:
            edited_positions.add(index)

    body_lines, had_wrapper = strip_wrapper(paragraphs)

    draft = SectionDraft(
        slug=slug or slugify(path.stem)[:100],
        title=title or path.stem,
        source_file=path.name,
        source_sha256=sha256_file(path),
        tracked_changes=tracked,
        comments=comments,
        fields=sorted(fields),
    )

    # A section whose first line repeats its own title reads as a duplicate
    # heading once the letter is assembled.
    if body_lines and _title_key(body_lines[0].lstrip("- ")) == _title_key(draft.title):
        body_lines = body_lines[1:]

    draft.slots = slots
    cleaned = [
        substitute_known_placeholders(SLOT_RE.sub("", line)).strip() for line in body_lines
    ]
    cleaned = [line for line in cleaned if line]

    # Copy-edit last, so it sees the text exactly as it will be sent.
    draft.paragraphs, report = copyedit_lines(
        cleaned, touched={index for index in edited_positions if index < len(cleaned)}
    )
    draft.copyedit = report.as_dict()
    entries = _section_paragraph_entries(document, draft.title)
    if entries:
        # The source DOCX is the formatting authority. Keep its run boundaries
        # and intentional blank paragraphs even when the plain-text copy-edit
        # projection has trimmed or lightly changed a line.
        draft.editor_state = lexical_editor_state(entries)
    if report.fixes:
        draft.notes.append(
            f"{len(report.fixes)} spacing/punctuation fix(es) applied during ingest."
        )
    if report.flags:
        draft.notes.append(
            f"{len(report.flags)} passage(s) need a human read: "
            + ", ".join(sorted({flag['kind'] for flag in report.flags}))
        )
    draft.variants = [
        VARIANT_RE.match(line).group(1).strip()
        for line in draft.paragraphs
        if VARIANT_RE.match(line)
    ]

    if had_wrapper:
        draft.notes.append("Model Letter wrapper removed; only the section's advice was kept.")
    if tracked:
        draft.notes.append(f"{tracked} tracked change(s) accepted during ingest.")
    if comments:
        draft.notes.append(f"{comments} reviewer comment(s) dropped during ingest.")

    draft.status = classify_status(draft)
    return draft


def classify_status(draft: SectionDraft) -> str:
    """Decide whether a section can be offered without a warning.

    Accepting an editor's marks is not the same as approving them. A section
    whose merge left a passage needing a human read stays out of the default
    picker; one whose accepted text copy-edits clean is treated as ready, with
    the acceptance recorded in its notes so the history is not lost.
    """
    if draft.word_count < STUB_WORD_COUNT:
        return STATUS_STUB
    if draft.comments:
        return STATUS_NEEDS_REVIEW
    if (draft.copyedit or {}).get("flags"):
        return STATUS_NEEDS_REVIEW
    return STATUS_READY


def extract_wrapper(path: Path) -> list[SectionDraft]:
    """Split the Model Letter into its opening and closing sections."""
    document, _touched = accept_tracked_changes(Document(path))
    _paragraph_lines(document)

    intro, closing = [], []
    intro_paragraphs, closing_paragraphs = [], []
    seen_closer = False
    for paragraph in document.paragraphs:
        text = paragraph.text
        normalized = _normalized(text)
        if any(normalized.startswith(closer) for closer in WRAPPER_CLOSERS):
            seen_closer = True
        if seen_closer:
            if not any(normalized.startswith(entry) for entry in WRAPPER_ADDRESS) and text.strip():
                closing.append(substitute_known_placeholders(text))
                closing_paragraphs.append(paragraph)
            continue
        if any(normalized.startswith(opener) for opener in WRAPPER_OPENERS):
            intro.append(substitute_known_placeholders(text))
            intro_paragraphs.append(paragraph)

    def with_spacing(selected):
        if not selected:
            return []
        selected_ids = {id(paragraph._p) for paragraph in selected}
        positions = [
            index
            for index, paragraph in enumerate(document.paragraphs)
            if id(paragraph._p) in selected_ids
        ]
        if not positions:
            return []
        start, end = min(positions), max(positions)
        return [
            (
                paragraph,
                substitute_known_placeholders(paragraph.text)
                if paragraph.text.strip()
                else "",
            )
            for paragraph in document.paragraphs[start : end + 1]
            if id(paragraph._p) in selected_ids or not paragraph.text.strip()
        ]

    drafts = []
    if intro:
        drafts.append(
            SectionDraft(
                slug="letter-opening",
                title="Opening",
                role=ROLE_INTRO,
                topic="Wrapper",
                paragraphs=intro,
                source_file=path.name,
                source_sha256=sha256_file(path),
                status=STATUS_READY,
                editor_state=lexical_editor_state(with_spacing(intro_paragraphs)),
            )
        )
    if closing:
        drafts.append(
            SectionDraft(
                slug="letter-closing",
                title="Closing",
                role=ROLE_CLOSING,
                topic="Wrapper",
                paragraphs=closing,
                source_file=path.name,
                source_sha256=sha256_file(path),
                status=STATUS_READY,
                editor_state=lexical_editor_state(with_spacing(closing_paragraphs)),
            )
        )
    return drafts


# ---------------------------------------------------------------- catalog sheet


def _column_index(reference):
    match = re.match(r"([A-Z]+)", reference or "")
    if not match:
        return 0
    index = 0
    for character in match.group(1):
        index = index * 26 + (ord(character) - 64)
    return index - 1


def read_catalog_sheet(path: Path) -> list[dict]:
    """Read the working group's spreadsheet into catalog rows."""
    with zipfile.ZipFile(path) as archive:
        shared = []
        if "xl/sharedStrings.xml" in archive.namelist():
            root = ET.fromstring(archive.read("xl/sharedStrings.xml"))
            shared = [
                "".join(node.text or "" for node in item.iter(f"{{{SHEET_NS}}}t"))
                for item in root.findall(f"{{{SHEET_NS}}}si")
            ]
        sheet_names = sorted(
            name for name in archive.namelist()
            if re.match(r"xl/worksheets/sheet\d+\.xml$", name)
        )
        if not sheet_names:
            return []
        root = ET.fromstring(archive.read(sheet_names[0]))

    rows = []
    letter_type = "brief_advice"
    topic = ""
    for row in root.iter(f"{{{SHEET_NS}}}row"):
        cells = {}
        for cell in row.findall(f"{{{SHEET_NS}}}c"):
            kind = cell.get("t")
            value_node = cell.find(f"{{{SHEET_NS}}}v")
            if kind == "s" and value_node is not None:
                value = shared[int(value_node.text)]
            elif value_node is not None:
                value = value_node.text
            else:
                value = ""
            if value and value.strip():
                cells[_column_index(cell.get("r"))] = " ".join(value.split())
        if not cells:
            continue
        first = cells.get(0, "")
        if first.lower().startswith("phase 2"):
            letter_type = "full_rep"
            continue
        if first.lower().startswith("phase 1") or first == "Type":
            continue
        if first:
            letter_type = "full_rep" if "full rep" in first.casefold() else "brief_advice"
        if cells.get(1):
            topic = cells[1]
        name = cells.get(2, "")
        if not name:
            continue
        rows.append(
            {
                "name": name,
                "topic": topic,
                "letter_type": letter_type,
                "cleveland_specific": bool(cells.get(3)),
                "cleveland_note": cells.get(3, ""),
                "region": (cells.get(4) or "").upper(),
            }
        )
    return rows


def match_catalog_row(section_title: str, rows: list[dict]) -> dict | None:
    """Match a sub-section file to its catalog row by name overlap."""
    def tokens(value):
        return {
            token
            for token in re.findall(r"[a-z0-9]+", (value or "").casefold())
            if token not in {"the", "a", "of", "for", "and", "to", "in", "cle", "neo", "draft"}
        }

    wanted = tokens(section_title)
    if not wanted:
        return None
    best, best_score = None, 0
    for row in rows:
        overlap = wanted & tokens(row["name"])
        score = len(overlap)
        if score > best_score:
            best, best_score = row, score
    return best if best_score >= 2 else None


# ---------------------------------------------------------------- packaging


def write_section_docx(draft: SectionDraft, path: Path):
    """Write a prepared section as a small DOCX for composition."""
    document = Document()
    state = draft.editor_state or {}
    root = state.get("root") if isinstance(state, dict) else None
    if isinstance(root, dict) and isinstance(root.get("children"), list):
        for node in root["children"]:
            if node.get("type") != "paragraph":
                continue
            text_nodes = [
                child for child in node.get("children", [])
                if child.get("type") == "text"
            ]
            first_text = text_nodes[0].get("text", "") if text_nodes else ""
            is_bullet = first_text.startswith("- ")
            paragraph = document.add_paragraph(style="List Bullet" if is_bullet else None)
            trim_prefix = 2 if is_bullet else 0
            for child in text_nodes:
                text = child.get("text", "")
                if trim_prefix:
                    removed = min(trim_prefix, len(text))
                    text = text[removed:]
                    trim_prefix -= removed
                if not text:
                    continue
                run = paragraph.add_run(text)
                format_value = int(child.get("format") or 0)
                run.bold = bool(format_value & LEXICAL_BOLD)
                run.italic = bool(format_value & LEXICAL_ITALIC)
                run.underline = bool(format_value & LEXICAL_UNDERLINE)
    else:
        for line in draft.paragraphs:
            if line.startswith("- "):
                document.add_paragraph(line[2:], style="List Bullet")
            else:
                document.add_paragraph(line)
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)
    return path


def section_to_manifest_row(draft: SectionDraft, docx_path: str) -> dict:
    return {
        "slug": draft.slug,
        "title": draft.title,
        "role": draft.role,
        "topic": draft.topic,
        "letter_type": draft.letter_type,
        "region": draft.region,
        "cleveland_specific": draft.cleveland_specific,
        "status": draft.status,
        "body": draft.body,
        "docx": docx_path,
        "fields": draft.fields,
        "slots": draft.slots,
        "variants": draft.variants,
        "notes": draft.notes,
        "copyedit": draft.copyedit,
        "editor_state": draft.editor_state,
        "word_count": draft.word_count,
        "source": {
            "file": draft.source_file,
            "sha256": draft.source_sha256,
            "tracked_changes": draft.tracked_changes,
            "comments": draft.comments,
        },
    }


def build_catalog(
    source_root: Path,
    output_root: Path,
    *,
    completions: dict | None = None,
    hints: dict | None = None,
    repairs: dict | None = None,
    derived: dict | None = None,
    retired: dict | None = None,
) -> Path:
    """Convert the maintained OneDrive folder into a prepared catalog package."""
    completions = completions or {}
    hints = hints or {}
    repairs = repairs or {}
    derived = derived or {}
    retired = retired or {}
    rows = []
    catalog_rows = []

    catalog_sheet = next(source_root.glob("*.xlsx"), None)
    if catalog_sheet:
        catalog_rows = read_catalog_sheet(catalog_sheet)

    model_letter = next(
        (path for path in source_root.glob("*.docx") if "model letter" in path.stem.casefold()),
        None,
    )
    drafts: list[SectionDraft] = []
    if model_letter:
        drafts.extend(extract_wrapper(model_letter))

    section_dir = source_root / "Letter Sub-Sections"
    for path in sorted(section_dir.rglob("*.docx")) if section_dir.exists() else []:
        if path.name.startswith("~$"):
            continue
        draft = extract_section(path)
        row = match_catalog_row(draft.title, catalog_rows)
        if row:
            draft.topic = row["topic"]
            draft.letter_type = row["letter_type"]
            draft.region = row["region"]
            draft.cleveland_specific = row["cleveland_specific"]
        drafts.append(draft)

    # Hand-completed replacements override the maintained text for sections the
    # working group had not finished.
    for draft in drafts:
        completion = completions.get(draft.slug)
        if not completion:
            continue
        draft.paragraphs = completion["paragraphs"]
        # The completion is a plain-text replacement, not the maintained DOCX
        # state. Do not let the old rich state silently resurrect the replaced
        # wording on export.
        draft.editor_state = {}
        draft.title = completion.get("title", draft.title)
        draft.status = completion.get("status", STATUS_AI_DRAFTED)
        draft.notes.append(completion.get("note", "Completed during ingest; attorney review required."))
        if completion.get("topic"):
            draft.topic = completion["topic"]

    # Sections authored here rather than lifted from a maintained file.
    for slug, spec in derived.items():
        drafts.append(
            SectionDraft(
                slug=slug,
                title=spec.get("title", slug.replace("-", " ").title()),
                topic=spec.get("topic", ""),
                region=spec.get("region", ""),
                status=spec.get("status", STATUS_AI_DRAFTED),
                paragraphs=list(spec["paragraphs"]),
                notes=[spec.get("note", "Drafted during ingest; attorney review required.")],
            )
        )

    # A composite whose reusable parts now live in their own sections.
    for draft in drafts:
        retirement = retired.get(draft.slug)
        if not retirement:
            continue
        draft.status = STATUS_STUB
        draft.notes.append(
            "Retired: "
            + retirement["reason"]
            + " Use instead: "
            + ", ".join(retirement.get("replaced_by", []))
            + "."
        )

    # Repair sentences the editor's own unfinished edit left ungrammatical.
    for draft in drafts:
        for repair in repairs.get(draft.slug, []):
            applied = False
            for index, line in enumerate(draft.paragraphs):
                if repair["broken"] in line:
                    draft.paragraphs[index] = line.replace(repair["broken"], repair["fixed"])
                    applied = True
            if applied:
                draft.notes.append(
                    f"Repaired a broken merge: \"{repair['broken']}\" -> "
                    f"\"{repair['fixed']}\". {repair['why']}"
                )
            else:
                draft.notes.append(
                    f"Merge repair no longer matches the source: \"{repair['broken']}\". "
                    "The maintained text may have been fixed upstream; re-check this section."
                )

    # Completions, repairs, and derived sections can be authored outside the
    # DOCX converter. Normalize them too, so every catalog path presents the
    # same Jinja field surface before it reaches the editor or exporter.
    for draft in drafts:
        normalized = []
        fields = set(draft.fields or [])
        for index, line in enumerate(draft.paragraphs, start=1):
            converted, conversion = convert_text(line, f"{draft.slug}_{index}")
            normalized.append(converted)
            fields.update(conversion.fields)
        draft.paragraphs = normalized
        draft.fields = sorted(fields)
        if draft.editor_state:
            draft.editor_state, conversion = convert_editor_state(
                draft.editor_state, draft.slug
            )
            draft.fields = sorted(set(draft.fields) | conversion.fields)

    # A readable record of what accepting the marks produced, so a reviewer can
    # open the resolved text in Word instead of re-deriving it from the
    # tracked-changes original.
    accepted_root = output_root / "accepted"
    for draft in drafts:
        if draft.tracked_changes or draft.comments:
            write_section_docx(draft, accepted_root / f"{draft.slug}.docx")

    sections_root = output_root / "sections"
    for draft in drafts:
        relative = f"sections/{draft.slug}/section.docx"
        write_section_docx(draft, sections_root / draft.slug / "section.docx")
        row = section_to_manifest_row(draft, relative)
        row.update(hints.get(draft.slug, {}))
        rows.append(row)

    for row in rows:
        report = check_readability(row["body"])
        row["readability"] = {
            "metrics": report.metrics,
            "warnings": [finding.as_dict() for finding in report.warnings],
            "estimatedPages": report.estimated_pages,
        }

    manifest = {
        "schema_version": CATALOG_VERSION,
        "slug": "client-advice-letters",
        "title": "Client advice letters",
        "description": (
            "Modular client advice letters: a fixed wrapper plus interchangeable "
            "sections chosen for the tenant's situation."
        ),
        "source": {
            "converted_at": datetime.now(timezone.utc).isoformat(),
            "converter": "apps.templates_app.advice_letters",
            "catalog_sheet": catalog_sheet.name if catalog_sheet else "",
            "model_letter": model_letter.name if model_letter else "",
        },
        "topics": sorted({row["topic"] for row in rows if row["topic"]}),
        "sections": rows,
    }
    output_root.mkdir(parents=True, exist_ok=True)
    manifest_path = output_root / "catalog.yaml"
    manifest_path.write_text(yaml.safe_dump(manifest, sort_keys=False, allow_unicode=True))
    return manifest_path
