import tempfile
import zipfile
from pathlib import Path

from django.test import TestCase
from docx import Document
from docx.oxml.ns import qn
from lxml import etree
from docxtpl import DocxTemplate

from apps.templates_app.jinja_filters import template_environment
from apps.templates_app.letterheads import letterhead_context, prepare_letterhead


def make_letterhead(path: Path, *, advocate="Julia Bertone", fax="440.352.0015", email="jbertone@example.org"):
    """A letterhead shaped like the maintained originals.

    The contact block sits in the header and the continuation header carries the
    blank the advocate fills in by hand.
    """
    document = Document()
    header = document.sections[0].header
    header.paragraphs[0].text = "Letter to ______, 6/1/2021, Page 1 of 2"
    header.add_paragraph("EXAMPLE LEGAL AID")
    header.add_paragraph(advocate)
    header.add_paragraph(f"Phone:  216.297.7957")
    if fax:
        header.add_paragraph(f"Fax:      {fax}")
    header.add_paragraph(email)
    document.core_properties.author = advocate
    document.core_properties.last_modified_by = advocate
    document.save(path)


class LetterheadPreparationTests(TestCase):
    def setUp(self):
        self.temp = tempfile.TemporaryDirectory()
        self.addCleanup(self.temp.cleanup)
        self.root = Path(self.temp.name)
        self.source = self.root / "advocate.docx"
        make_letterhead(self.source)

    def prepare(self, **kwargs):
        output = self.root / "letterhead.docx"
        report = prepare_letterhead(self.source, output)
        return output, report

    def test_contact_lines_become_variables(self):
        output, report = self.prepare()

        text = "\n".join(p.text for p in Document(output).sections[0].header.paragraphs)
        self.assertIn("{{ advocate_name }}", text)
        self.assertIn("{{ advocate_phone }}", text)
        self.assertIn("{{ advocate_email }}", text)
        self.assertIn("{{ letter_subject }}", text)
        self.assertNotIn("Julia Bertone", text)
        self.assertNotIn("jbertone@example.org", text)
        self.assertIn("advocate_name", report.variables)

    def test_masthead_text_is_left_alone(self):
        output, _report = self.prepare()

        text = "\n".join(p.text for p in Document(output).sections[0].header.paragraphs)
        self.assertIn("EXAMPLE LEGAL AID", text)

    def test_source_advocate_is_removed_from_document_properties(self):
        """The seed file is one person's copy; their name must not travel with it."""
        output, report = self.prepare()

        with zipfile.ZipFile(output) as archive:
            core = archive.read("docProps/core.xml").decode("utf-8")
        self.assertNotIn("Julia Bertone", core)
        self.assertTrue(any("document property" in entry for entry in report.replaced))

    def test_rendering_fills_the_contact_block(self):
        output, _report = self.prepare()
        template = DocxTemplate(output)
        template.render(
            letterhead_context(
                {"displayName": "Dana Ruiz", "phone": "216.555.0142", "fax": "440.352.0015", "email": "druiz@example.org"},
                subject="M. Alvarez",
                date="August 2, 2026",
            ),
            jinja_env=template_environment(),
        )
        rendered = self.root / "rendered.docx"
        template.save(rendered)

        text = "\n".join(p.text for p in Document(rendered).sections[0].header.paragraphs)
        self.assertIn("Dana Ruiz", text)
        self.assertIn("216.555.0142", text)
        self.assertIn("440.352.0015", text)
        self.assertIn("M. Alvarez, August 2, 2026", text)
        self.assertNotIn("{{", text)

    def test_fax_line_disappears_for_an_advocate_without_one(self):
        output, _report = self.prepare()
        template = DocxTemplate(output)
        template.render(
            letterhead_context(
                {"displayName": "Sam Okafor", "phone": "216.555.0199", "fax": "", "email": "sokafor@example.org"},
                subject="R. Chen",
                date="August 2, 2026",
            ),
            jinja_env=template_environment(),
        )
        rendered = self.root / "no_fax.docx"
        template.save(rendered)

        text = "\n".join(p.text for p in Document(rendered).sections[0].header.paragraphs)
        self.assertIn("Sam Okafor", text)
        self.assertNotIn("Fax", text)

    def test_letterhead_without_a_contact_block_reports_a_warning(self):
        plain = self.root / "plain.docx"
        document = Document()
        document.sections[0].header.paragraphs[0].text = "EXAMPLE LEGAL AID"
        document.save(plain)

        report = prepare_letterhead(plain, self.root / "out.docx")

        self.assertTrue(report.warnings)

    def test_no_relationship_reference_is_left_dangling(self):
        """Word reports an unresolvable r:id as unreadable content.

        Dropping the `attachedTemplate` relationship removed the authoring
        machine's path, but left `<w:attachedTemplate r:id="rId1"/>` in
        settings.xml pointing at nothing, so Word offered to repair every letter.
        """
        source = self.root / "attached.docx"
        make_letterhead(source)
        document = Document(source)
        settings_part = next(
            part
            for part in document.part.package.iter_parts()
            if str(part.partname).endswith("settings.xml")
        )
        settings = settings_part.element
        attached = settings.makeelement(qn("w:attachedTemplate"), {})
        rel_id = settings_part.relate_to(
            "file:///C:/Users/someone/Letterhead.dotx",
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/attachedTemplate",
            is_external=True,
        )
        attached.set(
            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id", rel_id
        )
        settings.insert(0, attached)
        document.save(source)

        output = self.root / "prepared.docx"
        prepare_letterhead(source, output)

        relationship_ns = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
        with zipfile.ZipFile(output) as archive:
            for name in archive.namelist():
                if not name.startswith("word/") or not name.endswith(".xml"):
                    continue
                rels_name = f"word/_rels/{name.split('/')[-1]}.rels"
                declared = set()
                if rels_name in archive.namelist():
                    declared = {
                        node.get("Id")
                        for node in etree.fromstring(archive.read(rels_name))
                    }
                used = {
                    node.get(relationship_ns)
                    for node in etree.fromstring(archive.read(name)).iter()
                    if node.get(relationship_ns)
                }
                self.assertEqual(used - declared, set(), f"{name} references a missing relationship")

    def test_mailto_hyperlink_is_dropped_without_orphaning_its_runs(self):
        """A `w:hyperlink` whose relationship is gone makes Word report damage."""
        source = self.root / "linked.docx"
        make_letterhead(source)
        document = Document(source)
        header = document.sections[0].header
        paragraph = next(p for p in header.paragraphs if "@" in p.text)
        rel_id = header.part.relate_to(
            "mailto:jbertone@example.org",
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = paragraph._p.makeelement(qn("w:hyperlink"), {})
        hyperlink.set(
            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id", rel_id
        )
        run = paragraph.runs[0]._r
        paragraph._p.replace(run, hyperlink)
        hyperlink.append(run)
        document.save(source)

        output = self.root / "unlinked.docx"
        report = prepare_letterhead(source, output)

        with zipfile.ZipFile(output) as archive:
            rels = [name for name in archive.namelist() if name.endswith(".rels")]
            joined = "".join(archive.read(name).decode("utf-8") for name in rels)
            document_xml = "".join(
                archive.read(name).decode("utf-8")
                for name in archive.namelist()
                if name.startswith("word/header")
            )
        self.assertNotIn("mailto:jbertone@example.org", joined)
        self.assertNotIn("<w:hyperlink", document_xml)
        self.assertIn("{{ advocate_email }}", document_xml)
        self.assertTrue(any("mailto" in entry for entry in report.replaced))
