import tempfile
import zipfile
from io import BytesIO
from pathlib import Path

import yaml
from django.db import connection
from django.test import TestCase
from django.test.utils import CaptureQueriesContext
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

from apps.core.content_library import organization_content_library_dir
from apps.drafting.models import DraftDocument, DraftingSession
from apps.exporting.services import _remove_editorial_caption_label, _render_docx_template, export_docx
from apps.matters.models import Matter
from apps.templates_app.content_library import sync_prepared_templates, sync_template_overrides
from apps.templates_app.ingestion import (
    LATITUDE_GENERATE,
    LATITUDE_GUIDED,
    LATITUDE_LOCKED,
    classify_latitude,
    discover_blocks,
    ingest_docx,
    is_heading,
)
from apps.templates_app.placeholders import convert_paragraph, convert_text
from apps.templates_app.jinja_filters import (
    as_noun,
    comma_and_list,
    did_verb,
    does_verb,
    pronoun_objective,
    pronoun_possessive,
    pronoun_subjective,
)
from apps.templates_app.models import DocumentTemplate
from apps.templates_app.template_variables import template_field_values


def make_source_docx(path: Path):
    document = Document()
    section = document.sections[0]
    section.header.paragraphs[0].text = "Legal Aid letterhead"
    document.add_heading("FACTS", level=1)
    fact = document.add_paragraph(style="List Number")
    fact.add_run("[Insert case specific facts]").bold = True
    document.add_paragraph("Further affiant sayeth naught.")
    document.add_heading("CERTIFICATE OF SERVICE", level=1)
    document.add_paragraph("I served this document on [DATE].")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Case No. [CASE NUMBER]"
    document.save(path)


def highlighted_run(paragraph, text):
    run = paragraph.add_run(text)
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return run


class PlaceholderConversionTests(TestCase):
    """The maintained wording survives; only marked fill-ins become bindings."""

    def test_conversion_keeps_run_formatting_on_surrounding_text(self):
        document = Document()
        paragraph = document.add_paragraph()
        paragraph.add_run("Now comes ")
        paragraph.add_run("Defendant ").bold = True
        highlighted_run(paragraph, "[DEFENDANT NAME]")
        paragraph.add_run(", by counsel.").italic = True

        convert_paragraph(paragraph, "body_1")

        self.assertIn("{{ defendant }}", paragraph.text)
        self.assertIn("Now comes", paragraph.text)
        self.assertIn(", by counsel.", paragraph.text)
        bold_runs = [run.text for run in paragraph.runs if run.bold]
        italic_runs = [run.text for run in paragraph.runs if run.italic]
        self.assertIn("Defendant ", bold_runs)
        self.assertIn(", by counsel.", italic_runs)

    def test_placeholder_split_across_runs_is_still_converted(self):
        # Word splits a bracketed placeholder across runs after a spell-check
        # pass, which a run-at-a-time substitution would miss entirely.
        document = Document()
        paragraph = document.add_paragraph()
        paragraph.add_run("Set for hearing on ")
        highlighted_run(paragraph, "[")
        highlighted_run(paragraph, "DATE] at [TIME]")
        highlighted_run(paragraph, ".")

        convert_paragraph(paragraph, "body_1")

        self.assertIn("{{ fields.hearing_date }}", paragraph.text)
        self.assertIn("{{ fields.hearing_time }}", paragraph.text)
        self.assertTrue(paragraph.text.endswith("."))

    def test_highlighted_sentence_keeps_its_wording_behind_a_toggle(self):
        document = Document()
        paragraph = document.add_paragraph()
        paragraph.add_run("Plaintiff will not be prejudiced as ")
        sentence = (
            "Defendant will pay monthly bond while the case is pending and "
            "Plaintiff retains rights in the second cause of action"
        )
        highlighted_run(paragraph, sentence)

        conversion = convert_paragraph(paragraph, "body_1")

        self.assertIn(sentence, paragraph.text)
        self.assertEqual(len(conversion.flags), 1)
        flag = next(iter(conversion.flags))
        self.assertIn("{%% if %s %%}" % flag, paragraph.text)
        self.assertIn("{% endif %}", paragraph.text)

    def test_conversion_clears_the_authors_highlighting(self):
        document = Document()
        paragraph = document.add_paragraph()
        highlighted_run(paragraph, "[DATE]")

        convert_paragraph(paragraph, "body_1")

        self.assertTrue(all(run.font.highlight_color is None for run in paragraph.runs))

    def test_editorial_case_caption_marker_stays_literal(self):
        document = Document()
        paragraph = document.add_paragraph()
        highlighted_run(paragraph, "Case Caption - Defendant's Motion for Bench Trial")

        conversion = convert_paragraph(paragraph, "body_1")

        self.assertTrue(paragraph.text.startswith("Case Caption"))
        self.assertEqual(conversion.fields, set())
        self.assertEqual(conversion.flags, set())

    def test_mixed_prepared_and_legacy_placeholders_are_both_supported(self):
        document = Document()
        paragraph = document.add_paragraph("Served on {{ fields.plaintiff_name }} [DATE].")

        convert_paragraph(paragraph, "body_1")

        self.assertEqual(
            paragraph.text,
            "Served on {{ fields.plaintiff_name }} {{ fields.filing_date }}.",
        )

    def test_brackets_inside_jinja_control_tags_are_not_rebound(self):
        converted, conversion = convert_text(
            '{%p for item in blocks["statement-case"]["items"] %} [Filing Date]',
            "body_1",
        )

        self.assertIn('blocks["statement-case"]["items"]', converted)
        self.assertIn("{{ fields.filing_date }}", converted)
        self.assertEqual(sorted(conversion.fields), ["fields.filing_date"])


class HeadingAndLatitudeTests(TestCase):
    def test_bold_normal_paragraph_is_recognized_as_a_heading(self):
        document = Document()
        paragraph = document.add_paragraph()
        paragraph.add_run("Complexity of Legal and Factual Issues").bold = True

        self.assertTrue(is_heading(paragraph))

    def test_bold_sentence_is_not_a_heading(self):
        document = Document()
        paragraph = document.add_paragraph()
        paragraph.add_run("The Court should grant this Motion because time is needed.").bold = True

        self.assertFalse(is_heading(paragraph))

    def test_quoted_authority_block_is_locked(self):
        document = Document()
        document.add_heading("STANDARD OF REVIEW", level=2)
        document.add_paragraph(
            "A court presented with a motion under Civ.R. 12(B)(1) must determine "
            "whether Plaintiff alleges any cause of action. See R.C. 5321.04 and "
            "24 C.F.R. 247.4(a)(2)."
        )
        block = discover_blocks(document)[0]

        self.assertEqual(classify_latitude(document, block), LATITUDE_LOCKED)

    def test_instruction_only_block_is_generated(self):
        document = Document()
        document.add_heading("FACTS", level=1)
        document.add_paragraph("[Insert case specific facts]")
        block = discover_blocks(document)[0]

        self.assertEqual(classify_latitude(document, block), LATITUDE_GENERATE)

    def test_block_with_fill_ins_is_guided(self):
        document = Document()
        document.add_heading("TIME NEEDED FOR DISCOVERY", level=2)
        document.add_paragraph(
            "Defendant anticipates propounding interrogatories to Plaintiff and "
            "issuing a subpoena for records to [PHA]. Defendant anticipates "
            "needing [X] days to complete discovery."
        )
        block = discover_blocks(document)[0]

        self.assertEqual(classify_latitude(document, block), LATITUDE_GUIDED)


class AlternativeClauseTests(TestCase):
    """An editorial "[OR]" becomes a choice, not two certificates and a marker."""

    def setUp(self):
        self.temp = tempfile.TemporaryDirectory()
        self.addCleanup(self.temp.cleanup)
        self.root = Path(self.temp.name)

    def build(self, first, second, *, heading="CERTIFICATE OF SERVICE", marker="[OR]"):
        source = self.root / "Alternatives.docx"
        document = Document()
        document.add_heading(heading, level=1)
        document.add_paragraph(first)
        document.add_paragraph(marker)
        document.add_paragraph(second)
        document.save(source)
        manifest_path = ingest_docx(
            source, self.root / "prepared", self.root / "snippets", force=True
        )
        return yaml.safe_load(manifest_path.read_text()), manifest_path.parent / "template.docx"

    def render(self, template_path, manifest, **context):
        values = {
            "fields": template_field_values({}),
            "blocks": {
                block["key"]: {"body": "", "paragraphs": [], "items": [], "revision": ""}
                for block in manifest["blocks"]
            },
            "defendant": "Jane Tenant",
            "plaintiff": "Acme LLC",
        }
        values.update({flag: True for flag in manifest.get("flags", [])})
        values.update(context)
        output = self.root / "rendered.docx"
        _render_docx_template(template_path, values, output)
        return [p.text.strip() for p in Document(output).paragraphs if p.text.strip()]

    def test_service_alternatives_become_a_named_choice(self):
        manifest, _path = self.build(
            "I certify that on [DATE] I served the motion on [PLAINTIFF] by email to [PLAINTIFF'S EMAIL].",
            "I certify that on [DATE] I served the motion on [PLAINTIFF] by United States mail to [PLAINTIFF'S ADDRESS].",
        )

        choice = manifest["choices"][0]
        self.assertEqual(choice["name"], "service_method")
        self.assertEqual(choice["options"], ["email", "mail"])
        self.assertEqual(choice["default"], "email")
        self.assertEqual(choice["block"], "certificate-of-service")

    def test_only_the_chosen_alternative_renders(self):
        manifest, path = self.build(
            "I certify that I served the motion by email to [PLAINTIFF'S EMAIL].",
            "I certify that I served the motion by United States mail to [PLAINTIFF'S ADDRESS].",
        )

        mailed = self.render(path, manifest, service_method="mail")

        self.assertTrue(any("United States mail" in line for line in mailed))
        self.assertFalse(any("by email" in line for line in mailed))
        self.assertNotIn("[OR]", mailed)

    def test_an_unanswered_choice_falls_back_to_the_first_alternative(self):
        """A certificate of service must never silently disappear."""
        manifest, path = self.build(
            "I certify that I served the motion by email to [PLAINTIFF'S EMAIL].",
            "I certify that I served the motion by United States mail to [PLAINTIFF'S ADDRESS].",
        )

        rendered = self.render(path, manifest)

        certificates = [line for line in rendered if "I certify" in line]
        self.assertEqual(len(certificates), 1)
        self.assertIn("by email", certificates[0])

    def test_three_alternatives_chain_into_one_choice(self):
        source = self.root / "Three.docx"
        document = Document()
        document.add_heading("CERTIFICATE OF SERVICE", level=1)
        document.add_paragraph("I served the motion by email to [PLAINTIFF'S EMAIL].")
        document.add_paragraph("[OR]")
        document.add_paragraph("I served the motion by United States mail to [PLAINTIFF'S ADDRESS].")
        document.add_paragraph("[OR]")
        document.add_paragraph("I served the motion by personal service on [PLAINTIFF].")
        document.save(source)
        manifest_path = ingest_docx(
            source, self.root / "prepared", self.root / "snippets", force=True
        )
        manifest = yaml.safe_load(manifest_path.read_text())

        choice = manifest["choices"][0]
        self.assertEqual(choice["options"], ["email", "mail", "personal"])

        rendered = self.render(
            manifest_path.parent / "template.docx", manifest, service_method="personal"
        )
        served = [line for line in rendered if "I served" in line]
        self.assertEqual(len(served), 1)
        self.assertIn("personal service", served[0])

    def test_a_non_certificate_block_gets_a_block_scoped_variable(self):
        manifest, _path = self.build(
            "Defendant requests a stay of execution pending appeal.",
            "Defendant requests dismissal of the complaint.",
            heading="PRAYER FOR RELIEF",
        )

        choice = manifest["choices"][0]
        self.assertEqual(choice["name"], "prayer_for_relief_option")
        self.assertEqual(choice["options"], ["option_1", "option_2"])

    def test_prose_containing_the_word_or_is_not_treated_as_a_marker(self):
        manifest, path = self.build(
            "Plaintiff served a notice by email.",
            "Plaintiff served a notice by mail.",
            marker="The notice was defective or untimely.",
        )

        self.assertEqual(manifest.get("choices"), [])
        rendered = self.render(path, manifest)
        self.assertTrue(any("defective or untimely" in line for line in rendered))


class TemplatePhrasingFilterTests(TestCase):
    def test_list_and_number_aware_language_matches_assemblyline_conventions(self):
        self.assertEqual(comma_and_list(["Ada", "Grace", "Katherine"]), "Ada, Grace, and Katherine")
        self.assertEqual(comma_and_list(["Ada", "Grace"], and_string="or"), "Ada or Grace")
        self.assertEqual(does_verb(["Ada"], "live"), "lives")
        self.assertEqual(does_verb(["Ada", "Grace"], "live"), "live")
        self.assertEqual(did_verb(["Ada"], "be"), "was")
        self.assertEqual(did_verb(["Ada", "Grace"], "be"), "were")
        self.assertEqual(as_noun(["Ada"], "occupant"), "occupant")
        self.assertEqual(as_noun(["Ada", "Grace"], "occupant"), "occupants")

    def test_explicit_and_custom_pronoun_sets_are_not_inferred_from_names(self):
        she = {"name": "Ada", "pronouns": "she/her/hers"}
        custom = {"name": "River", "pronouns": "ze/zir/zirs"}
        unknown = {"name": "Morgan", "pronouns": ""}

        self.assertEqual(pronoun_subjective(she), "she")
        self.assertEqual(pronoun_objective(custom), "zir")
        self.assertEqual(pronoun_possessive(custom, "home"), "zir home")
        self.assertEqual(pronoun_subjective(unknown), "Morgan")
        self.assertEqual(pronoun_possessive(unknown, "home"), "Morgan's home")

    def test_docx_renderer_registers_custom_filters(self):
        with tempfile.TemporaryDirectory() as directory:
            source = Path(directory) / "phrasing.docx"
            output = Path(directory) / "rendered.docx"
            document = Document()
            document.add_paragraph(
                '{{ client | pronoun_subjective(capitalize=True) }} lives with '
                '{{ household | comma_and_list }} in '
                '{{ client | pronoun_possessive("home") }}.'
            )
            document.save(source)

            _render_docx_template(
                source,
                {
                    "client": {"name": "Ada", "pronouns": "she/her/hers"},
                    "household": ["Ben", "Cam"],
                },
                output,
            )

            self.assertEqual(
                Document(output).paragraphs[0].text,
                "She lives with Ben and Cam in her home.",
            )


class TemplateIngestionTests(TestCase):
    def setUp(self):
        self.temp = tempfile.TemporaryDirectory()
        self.addCleanup(self.temp.cleanup)
        self.root = Path(self.temp.name)
        self.content = self.root / "content"
        self.source = self.content / "original_templates" / "Test Motion.docx"
        self.source.parent.mkdir(parents=True)
        make_source_docx(self.source)

    def ingest(self):
        return ingest_docx(
            self.source,
            self.content / "document-templates",
            self.content / "docx-snippets",
            force=True,
        )

    def content_library(self, public_root=None):
        """Point both providers at this test's fixtures.

        `ORGANIZATION_CONTENT_LIBRARY_DIR` defaults to the checkout's private
        submodule. Leaving it alone let a developer's real prepared templates
        into these assertions, so results depended on whose machine ran them.
        """
        return self.settings(
            CONTENT_LIBRARY_DIR=public_root or self.content,
            ORGANIZATION_CONTENT_LIBRARY_DIR=self.root / "private-content",
        )

    def test_ingestion_preserves_word_structure_and_adds_list_loop(self):
        manifest_path = self.ingest()
        manifest = yaml.safe_load(manifest_path.read_text())
        template_path = manifest_path.parent / "template.docx"

        self.assertEqual(manifest["render"]["strategy"], "full_document")
        self.assertEqual(manifest["description"], "Maintained Word template for Test Motion.")
        self.assertEqual(
            manifest["goal"],
            "Draft Test Motion with case-specific facts, legal grounds, and requested relief.",
        )
        facts = next(block for block in manifest["blocks"] if block["type"] == "facts")
        self.assertEqual(facts["input"]["type"], "array")
        self.assertEqual(facts["lexical"]["node"], "list")
        self.assertTrue((self.content / facts["docx"]).is_file())

        converted = Document(template_path)
        self.assertEqual(converted.sections[0].header.paragraphs[0].text, "Legal Aid letterhead")
        self.assertEqual(len(converted.tables), 1)
        self.assertIn("{{ case_number }}", converted.tables[0].cell(0, 0).text)
        with zipfile.ZipFile(template_path) as archive:
            xml = archive.read("word/document.xml").decode("utf-8")
        self.assertIn('{%p for item in blocks["facts"]["items"] %}', xml)
        self.assertIn("{{ item }}", xml)
        self.assertIn("{%p endfor %}", xml)
        self.assertIn("Further affiant sayeth naught.", xml)

    def test_ingestion_preserves_editorial_brackets_and_names_generic_inputs_from_context(self):
        converted, conversion = convert_text(
            "Equity [a]bhors forfeiture under note [216]. Hearing date: [INSERT]",
            "body_1",
        )

        self.assertIn("[a]bhors", converted)
        self.assertIn("[216]", converted)
        self.assertIn("{{ fields.hearing_date }}", converted)
        self.assertEqual(sorted(conversion.fields), ["fields.hearing_date"])

    def test_blank_placeholder_uses_surrounding_party_context(self):
        converted, conversion = convert_text(
            "Now comes Defendant [ ] by and through counsel.",
            "caption_1",
        )

        self.assertIn("{{ defendant }}", converted)
        self.assertEqual(sorted(conversion.fields), [])

    def test_underscore_blanks_become_fields(self):
        converted, conversion = convert_text(
            "Case No. ________ served on ____________________.",
            "body_1",
        )

        self.assertNotIn("___", converted)
        self.assertIn("{{ case_number }}", converted)
        self.assertIn("{{ fields.service_date }}", converted)
        self.assertEqual(sorted(conversion.fields), ["fields.service_date"])

    def test_styled_email_placeholder_is_not_discovered_as_a_section(self):
        document = Document()
        document.add_heading("SIGNATURE", level=1)
        document.add_paragraph("[Email]", style="Heading 2")

        blocks = discover_blocks(document)

        self.assertEqual(len(blocks), 1)
        self.assertEqual(blocks[0].label, "Signature")

    def test_export_cleanup_removes_editorial_case_caption_label(self):
        path = self.root / "caption.docx"
        document = Document()
        document.add_paragraph("Case Caption - Defendant's Motion to Dismiss")
        document.add_paragraph("Now comes Defendant Jane Tenant.")
        document.save(path)
        block = type(
            "Block",
            (),
            {"label": "Case Caption - Defendant's Motion to Dismiss"},
        )()

        _remove_editorial_caption_label(
            path,
            block,
            {
                "court": "Housing Court",
                "plaintiff": "Example Landlord",
                "defendant": "Jane Tenant",
                "case_number": "CASE-1",
                "document": {"title": "Motion to Dismiss"},
            },
        )

        rendered = Document(path)
        text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
        table_text = "\n".join(cell.text for cell in rendered.tables[0].rows[0].cells)
        self.assertNotIn("Case Caption -", text)
        self.assertIn("Now comes Defendant Jane Tenant.", text)
        self.assertIn("Example Landlord", table_text)
        self.assertIn("CASE-1", table_text)

    def test_first_block_keeps_leading_caption_tables(self):
        source = self.source.parent / "Table Caption Motion.docx"
        document = Document()
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Plaintiff"
        table.cell(0, 1).text = "Case No. [CASE NUMBER]"
        document.add_heading("INTRODUCTION", level=1)
        document.add_paragraph("Body text.")
        document.save(source)

        manifest_path = ingest_docx(
            source,
            self.content / "document-templates",
            self.content / "docx-snippets",
            force=True,
        )
        manifest = yaml.safe_load(manifest_path.read_text())
        first_block_path = self.content / manifest["blocks"][0]["docx"]
        first_block = Document(first_block_path)

        self.assertEqual(len(first_block.tables), 1)
        self.assertIn("{{ case_number }}", first_block.tables[0].cell(0, 1).text)

    def test_sync_indexes_manifests_and_preserves_admin_slug_conflicts(self):
        manifest_path = self.ingest()
        with self.content_library():
            results = sync_prepared_templates()
            template = DocumentTemplate.objects.get(slug="test-motion")
            self.assertEqual(template.source_kind, "content_library")
            self.assertEqual(template.content_path, "document-templates/test-motion/manifest.yaml")
            self.assertEqual(
                template.goal,
                "Draft Test Motion with case-specific facts, legal grounds, and requested relief.",
            )
            self.assertTrue(template.blocks.filter(block_type="facts", input_schema__type="array").exists())
            self.assertEqual(results[0]["status"], "created")

            template.source_kind = "database"
            template.title = "Admin title"
            template.save()
            results = sync_prepared_templates()
            template.refresh_from_db()
            self.assertEqual(template.title, "Admin title")
            self.assertEqual(results[0]["status"], "conflict")

    def test_sync_repairs_legacy_generic_template_description_without_forced_reingestion(self):
        manifest_path = self.ingest()
        manifest = yaml.safe_load(manifest_path.read_text())
        manifest.pop("goal")
        manifest["description"] = "Prepared from the maintained original Word template."
        manifest_path.write_text(yaml.safe_dump(manifest, sort_keys=False), encoding="utf-8")
        with self.content_library():
            sync_prepared_templates()
            template = DocumentTemplate.objects.get(slug="test-motion")
            template.goal = ""
            template.description = "Prepared from the maintained original Word template."
            template.save(update_fields=["goal", "description", "updated_at"])

            results = sync_prepared_templates()

        template.refresh_from_db()
        self.assertEqual(results[0]["status"], "updated")
        self.assertEqual(template.description, "Maintained Word template for Test Motion.")
        self.assertEqual(
            template.goal,
            "Draft Test Motion with case-specific facts, legal grounds, and requested relief.",
        )

    def test_unchanged_sync_does_not_write_to_database(self):
        self.ingest()
        with self.content_library():
            sync_prepared_templates()

            with CaptureQueriesContext(connection) as queries:
                results = sync_prepared_templates()

        self.assertEqual(results[0]["status"], "unchanged")
        write_statements = [
            query["sql"]
            for query in queries.captured_queries
            if query["sql"].lstrip().upper().startswith(("INSERT", "UPDATE", "DELETE"))
        ]
        self.assertEqual(write_statements, [])

    def test_missing_content_provider_does_not_deactivate_indexed_templates(self):
        template = DocumentTemplate.objects.create(
            slug="provider-temporarily-missing",
            title="Provider temporarily missing",
            kind="motion",
            source_kind="content_library",
            content_path="document-templates/provider-temporarily-missing/manifest.yaml",
        )
        missing_root = self.root / "not-mounted"

        with self.content_library(missing_root):
            results = sync_prepared_templates()

        template.refresh_from_db()
        self.assertEqual(results, [])
        self.assertTrue(template.is_active)

    def test_private_template_override_updates_metadata_and_blocks(self):
        template = DocumentTemplate.objects.create(
            slug="private-override-test",
            title="Original title",
            kind="motion",
            jurisdiction="Ohio",
        )
        block = template.blocks.create(
            key="facts",
            label="Facts",
            block_type="facts",
            body="Original body",
        )
        removable = template.blocks.create(
            key="editorial-note",
            label="Editorial note",
            block_type="optional_clause",
            body="Remove me",
        )
        private = self.root / "private"
        override_dir = private / "template-overrides"
        override_dir.mkdir(parents=True)
        (override_dir / "private-override-test.yaml").write_text(
            yaml.safe_dump(
                {
                    "schema_version": 1,
                    "slug": template.slug,
                    "template": {
                        "jurisdiction": "Housing Court",
                        "metadata": {
                            "fields": ["fields.client_pronouns"],
                            "fieldSchema": {"client_pronouns": {"type": "pronouns"}},
                        },
                    },
                    "blocks": [
                        {"key": block.key, "body": "{{ client | pronoun_subjective }} lives here."},
                        {"key": removable.key, "delete": True},
                        {
                            "key": "argument",
                            "create": True,
                            "label": "Argument",
                            "block_type": "argument",
                            "order": 30,
                            "body": "Argument body",
                            "ai_fill_mode": "constrained_generation",
                        },
                    ],
                }
            ),
            encoding="utf-8",
        )

        with self.settings(
            CONTENT_LIBRARY_DIR=self.content,
            ORGANIZATION_CONTENT_LIBRARY_DIR=private,
        ):
            results = sync_template_overrides()

        template.refresh_from_db()
        block.refresh_from_db()
        self.assertEqual(results[0]["status"], "updated")
        self.assertEqual(template.jurisdiction, "Housing Court")
        self.assertEqual(template.metadata["fieldSchema"]["client_pronouns"]["type"], "pronouns")
        self.assertIn("pronoun_subjective", block.body)
        self.assertFalse(template.blocks.filter(pk=removable.pk).exists())
        self.assertTrue(template.blocks.filter(key="argument", block_type="argument").exists())

    def test_full_template_export_uses_edited_lexical_block_values(self):
        manifest_path = self.ingest()
        template_path = manifest_path.parent / "template.docx"
        source_doc = Document(template_path)
        source_doc.add_paragraph("Numbered facts: {{ blocks.facts.numbered_items | join('; ') }}")
        source_doc.save(template_path)
        with self.content_library():
            sync_prepared_templates()
            template = DocumentTemplate.objects.prefetch_related("blocks").get(slug="test-motion")
            matter = Matter.objects.create(
                external_id="2026-CVG-1",
                client_name="Jane Tenant",
                matter_type="Eviction",
                jurisdiction="Housing Court",
            )
            session = DraftingSession.objects.create(
                mode="draft_from_template",
                matter=matter,
                template=template,
                template_data={"filing_date": "June 28, 2026"},
            )
            draft = DraftDocument.objects.create(
                session=session,
                title="Test Motion",
                sections=[
                    {"key": "facts", "label": "Facts", "body": "First edited fact.\nSecond edited fact.", "format": {"style": "numbered"}},
                    {"key": "certificate-of-service", "label": "Certificate of Service", "body": "Edited certificate text.\nAdded overflow paragraph."},
                ],
                plain_text="",
            )

            response = export_docx(draft)

        with zipfile.ZipFile(BytesIO(response.content)) as archive:
            xml = archive.read("word/document.xml").decode("utf-8")
        self.assertIn("First edited fact.", xml)
        self.assertIn("Second edited fact.", xml)
        self.assertIn("1. First edited fact.", xml)
        self.assertIn("2. Second edited fact.", xml)
        self.assertIn("Edited certificate text.", xml)
        self.assertIn("Added overflow paragraph.", xml)
        self.assertNotIn("{%p", xml)

    def test_unedited_block_renders_maintained_wording_not_its_plain_text_copy(self):
        """An untouched section must come from the DOCX, keeping its formatting."""
        manifest_path = self.ingest()
        with self.content_library():
            sync_prepared_templates()
            template = DocumentTemplate.objects.prefetch_related("blocks").get(slug="test-motion")
            certificate = template.blocks.get(key="certificate-of-service")
            matter = Matter.objects.create(
                external_id="2026-CVG-2",
                client_name="Jane Tenant",
                matter_type="Eviction",
                jurisdiction="Housing Court",
            )
            session = DraftingSession.objects.create(
                mode="draft_from_template",
                matter=matter,
                template=template,
                template_data={"service_date": "June 28, 2026"},
            )
            draft = DraftDocument.objects.create(
                session=session,
                title="Test Motion",
                # The editor round-trips the maintained body unchanged.
                sections=[
                    {
                        "key": "certificate-of-service",
                        "label": "Certificate of Service",
                        "body": certificate.body,
                    }
                ],
                plain_text="",
            )

            response = export_docx(draft)

        with zipfile.ZipFile(BytesIO(response.content)) as archive:
            xml = archive.read("word/document.xml").decode("utf-8")
        self.assertIn("I served this document on", xml)
        self.assertIn("June 28, 2026", xml)
        self.assertNotIn("{{", xml)
        self.assertNotIn("{%p", xml)

    def test_maintained_wording_is_not_a_model_written_slot(self):
        """The regression this replaced: prose rebound to blocks[...] paragraphs."""
        manifest_path = self.ingest()
        with zipfile.ZipFile(manifest_path.parent / "template.docx") as archive:
            xml = archive.read("word/document.xml").decode("utf-8")

        self.assertIn("Further affiant sayeth naught.", xml)
        self.assertIn("I served this document on", xml)
        self.assertNotIn('blocks["certificate-of-service"]["paragraphs"]', xml)


class RepositorySnippetRenderingTests(TestCase):
    def test_every_maintained_docx_snippet_renders(self):
        snippet_root = organization_content_library_dir() / "docx-snippets"
        paths = sorted(snippet_root.glob("*/blocks/*.docx"))
        if not paths:
            self.skipTest("No private organization snippets are mounted")
        block_keys = {path.stem for path in paths}
        context = {
            "fields": template_field_values(),
            "blocks": {key: {"body": "", "paragraphs": [], "items": []} for key in block_keys},
            "document": {},
            "section": {},
            "matter": {},
            "author": {},
            "selected_facts": [],
            "selected_curated_facts": [],
            "selected_sources": [],
            "instructions": "",
            "court": "[Court]",
            "plaintiff": "Plaintiff",
            "defendant": "[Defendant]",
            "case_number": "[Case Number]",
            "advocate_name": "[Advocate]",
            "advocate_signoff": "Respectfully submitted,",
            "advocate_salutation": "",
            "advocate_organization": "",
            "advocate_email": "[Email]",
            "advocate_phone": "[Phone]",
            "advocate_address": "[Address]",
            "advocate_contact": "[Advocate Contact]",
            "advocate_signature_image": "",
        }

        with tempfile.TemporaryDirectory() as directory:
            for index, path in enumerate(paths):
                output = Path(directory) / f"{index}.docx"
                _render_docx_template(path, context, output)
                Document(output)
