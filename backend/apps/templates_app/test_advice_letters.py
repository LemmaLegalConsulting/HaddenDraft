import tempfile
from pathlib import Path

import yaml
from django.test import TestCase
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from apps.matters.models import Matter
from apps.templates_app.advice_letter_hints import (
    load_selection_hints,
    seed_selection_hints,
)
from apps.templates_app.advice_letters import (
    STATUS_AI_DRAFTED,
    STATUS_NEEDS_REVIEW,
    STATUS_READY,
    accept_tracked_changes,
    build_catalog,
    extract_section,
    is_wrapper_line,
    match_catalog_row,
    strip_wrapper,
)
from apps.templates_app.models import AdviceLetterSection
from apps.templates_app.recommendations import recommend_advice_sections


def add_tracked_paragraph(document, kept: str, inserted: str, deleted: str):
    """A paragraph with an unaccepted insertion and deletion, as Word writes it."""
    paragraph = document.add_paragraph()
    paragraph.add_run(kept)

    insertion = OxmlElement("w:ins")
    insertion.set(qn("w:id"), "1")
    insertion.set(qn("w:author"), "Reviewer")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = inserted
    text.set(qn("xml:space"), "preserve")
    run.append(text)
    insertion.append(run)
    paragraph._p.append(insertion)

    deletion = OxmlElement("w:del")
    deletion.set(qn("w:id"), "2")
    deletion.set(qn("w:author"), "Reviewer")
    deleted_run = OxmlElement("w:r")
    deleted_text = OxmlElement("w:delText")
    deleted_text.text = deleted
    deleted_run.append(deleted_text)
    deletion.append(deleted_run)
    paragraph._p.append(deletion)
    return paragraph


class TrackedChangeTests(TestCase):
    def setUp(self):
        self.temp = tempfile.TemporaryDirectory()
        self.addCleanup(self.temp.cleanup)
        self.root = Path(self.temp.name)

    def test_insertions_are_kept_and_deletions_dropped(self):
        """Unresolved edits shredded the text: runs inside w:ins were skipped."""
        path = self.root / "tracked.docx"
        document = Document()
        add_tracked_paragraph(document, "The notice must be ", "conspicuous.", " obvious.")
        document.save(path)

        accepted, touched = accept_tracked_changes(Document(path))
        text = "\n".join(p.text for p in accepted.paragraphs)

        self.assertIn("The notice must be conspicuous.", text)
        self.assertNotIn("obvious", text)
        # The edited paragraph is reported so a copy-editor knows to read it.
        self.assertEqual(touched, {0})

    def test_a_section_with_tracked_changes_is_marked_for_review(self):
        path = self.root / "tracked.docx"
        document = Document()
        add_tracked_paragraph(
            document,
            "You can ask the court to seal this case. You can do this without a "
            "lawyer. It costs money to file. The court staff can help you fill "
            "out the forms. Call them if you have questions. It may take a few "
            "months for the court to decide. The court will mail you a copy of "
            "the decision. If you have moved, tell the court your new address. "
            "Keep a copy of everything you file. Bring your court papers with "
            "you. Ask for a receipt when you file. Read every paper the court "
            "sends you. Write down the date of your next hearing. Tell the "
            "court if your phone number changes. Ask the clerk if you are "
            "unsure where to go. Arrive early on the day of your hearing. "
            "Bring a photo of anything you want to show the judge. ",
            "Do it soon.",
            " Later.",
        )
        document.save(path)

        draft = extract_section(path)

        self.assertEqual(draft.status, STATUS_NEEDS_REVIEW)
        self.assertTrue(any("tracked change" in note for note in draft.notes))


class WrapperTests(TestCase):
    def test_model_letter_scaffolding_is_recognized(self):
        for line in (
            "July 14, 2026",
            "Client Name",
            "Dear Mx.,",
            "Thank you for asking Legal Aid for help with your eviction.",
            "Closing Your File. I hope this advice is helpful.",
            "Sincerely,",
        ):
            self.assertTrue(is_wrapper_line(line), line)

    def test_advice_text_is_not_mistaken_for_scaffolding(self):
        for line in (
            "You can ask the Court to seal the record of this case.",
            "The 3-Day Notice is from your landlord.",
        ):
            self.assertFalse(is_wrapper_line(line), line)

    def test_repeated_wrapper_is_removed_from_a_section(self):
        kept, changed = strip_wrapper(
            [
                "July 14, 2026",
                "Dear Mx.,",
                "Thank you for asking Legal Aid for help with your eviction.",
                "You can ask the Court to seal the record of this case.",
                "Sincerely,",
            ]
        )

        self.assertTrue(changed)
        self.assertEqual(kept, ["You can ask the Court to seal the record of this case."])


class CatalogTests(TestCase):
    def setUp(self):
        self.temp = tempfile.TemporaryDirectory()
        self.addCleanup(self.temp.cleanup)
        self.root = Path(self.temp.name)
        self.source = self.root / "source"
        (self.source / "Letter Sub-Sections").mkdir(parents=True)

        model = Document()
        model.add_paragraph("July 14, 2026")
        model.add_paragraph("Dear Mx.,")
        model.add_paragraph(
            "Thank you for asking Legal Aid for help with your eviction. "
            "We are sorry, but we cannot represent you."
        )
        model.add_paragraph("Closing Your File. I hope this advice is helpful.")
        model.add_paragraph("Sincerely,")
        model.save(self.source / "Model Letter.docx")

        section = Document()
        section.add_paragraph("Sealing the Record")
        section.add_paragraph(
            "You can ask the Court to seal the record of this case. "
            "You can do this without an attorney. It costs $25 to file. "
            "The Court's Housing Specialists can help you. Call them at 216-664-4295. "
            "It may take a few months for the Court to rule. "
            "The Court will mail you a copy of the decision. "
            "If you have moved, file a change of address with the Court. "
            "You must attach copies of some of your court papers. "
            "Keep a copy of everything you file. "
            "If you cannot afford the fee, ask the Court to waive it. "
            "The Housing Specialists can help you with that form too. "
            "Do not sign the affidavit until you are in front of a notary. "
            "Most banks have a notary you can use for free."
        )
        section.save(self.source / "Letter Sub-Sections" / "Motion to Seal.docx")

    def build(self, **kwargs):
        manifest = build_catalog(self.source, self.root / "advice-letters", **kwargs)
        return yaml.safe_load(manifest.read_text())

    def test_wrapper_becomes_opening_and_closing_sections(self):
        data = self.build()
        roles = {row["slug"]: row["role"] for row in data["sections"]}

        self.assertEqual(roles.get("letter-opening"), "intro")
        self.assertEqual(roles.get("letter-closing"), "closing")

    def test_section_is_catalogued_with_a_readability_score(self):
        data = self.build()
        section = next(row for row in data["sections"] if row["slug"] == "motion-to-seal")

        self.assertEqual(section["role"], "body")
        self.assertEqual(section["status"], STATUS_READY)
        self.assertIn("flesch_kincaid_grade", section["readability"]["metrics"])
        self.assertTrue((self.root / "advice-letters" / section["docx"]).is_file())

    def test_a_completion_replaces_an_unfinished_section(self):
        data = self.build(
            completions={
                "motion-to-seal": {
                    "title": "Sealing the Record",
                    "paragraphs": ["You can ask the Court to hide this case."],
                    "status": STATUS_AI_DRAFTED,
                    "note": "Drafted during ingest.",
                }
            }
        )
        section = next(row for row in data["sections"] if row["slug"] == "motion-to-seal")

        self.assertEqual(section["status"], STATUS_AI_DRAFTED)
        self.assertIn("hide this case", section["body"])
        self.assertTrue(any("Drafted during ingest" in note for note in section["notes"]))

    def test_catalog_row_matching_needs_real_overlap(self):
        rows = [{"name": "Motion to Seal (CLE)", "topic": "Pro se How-To"}]

        self.assertIsNotNone(match_catalog_row("Motion to Seal", rows))
        self.assertIsNone(match_catalog_row("Security Deposit", rows))


class SelectionHintTests(TestCase):
    def setUp(self):
        self.temp = tempfile.TemporaryDirectory()
        self.addCleanup(self.temp.cleanup)
        self.root = Path(self.temp.name)

    def test_seeded_hints_are_not_overwritten(self):
        """An advocate's edits must survive re-ingestion."""
        path = seed_selection_hints(self.root)
        edited = path.read_text().replace("decarlo:", "decarlo:  # reviewed\n", 1)
        path.write_text(edited)

        seed_selection_hints(self.root)

        self.assertIn("# reviewed", path.read_text())

    def test_force_reseeds(self):
        path = seed_selection_hints(self.root)
        path.write_text("{}\n")

        seed_selection_hints(self.root, force=True)

        self.assertIn("decarlo", load_selection_hints(self.root))


class AdviceSectionRecommendationTests(TestCase):
    def setUp(self):
        self.matter = Matter.objects.create(
            external_id="2026-CVG-77",
            client_name="Maria Alvarez",
            matter_type="Eviction",
            jurisdiction="Cleveland Municipal Court",
            summary=(
                "Nonpayment eviction. The 3-day notice names a different landlord "
                "than the complaint. Hearing is scheduled next week."
            ),
        )
        self.decarlo = AdviceLetterSection.objects.create(
            slug="decarlo",
            title="DeCarlo",
            region="NEO",
            status="ready",
            body="Issue with 3-Day Notice.",
            selection_hints={
                "triggers": ["3-day notice names a different landlord than the complaint"],
                "requires": ["has_3_day_notice"],
                "excludes": ["negotiate-move-out"],
                "summary": "Notice and complaint name different parties.",
            },
        )
        self.negotiate = AdviceLetterSection.objects.create(
            slug="negotiate-move-out",
            title="Negotiate move-out",
            status="ready",
            body="You must go to your hearing.",
            selection_hints={
                "triggers": ["3-day notice names a different landlord than the complaint"],
                "summary": "Settle without a judgment.",
            },
        )
        self.unreviewed = AdviceLetterSection.objects.create(
            slug="unreviewed",
            title="Unreviewed section",
            status="needs_review",
            needs_attorney_review=True,
            review_reason="4 tracked changes accepted here",
            body="Something.",
            selection_hints={
                "triggers": ["3-day notice names a different landlord than the complaint"]
            },
        )

    def sections(self):
        return list(AdviceLetterSection.objects.filter(is_active=True))

    def test_a_matching_trigger_ranks_the_section(self):
        results = recommend_advice_sections(
            self.sections(), self.matter, conditions={"has_3_day_notice": True}
        )

        self.assertEqual(results[0]["section"].slug, "decarlo")
        self.assertTrue(any("3-day notice names" in reason for reason in results[0]["reasons"]))

    def test_conflicting_sections_are_not_both_offered(self):
        """Advising both "you have a defense" and "settle" reads as indecision."""
        results = recommend_advice_sections(
            self.sections(), self.matter, conditions={"has_3_day_notice": True}
        )

        slugs = [entry["section"].slug for entry in results]
        self.assertIn("decarlo", slugs)
        self.assertNotIn("negotiate-move-out", slugs)

    def test_unreviewed_sections_are_offered_with_a_warning(self):
        """Withholding them hid the best match for a case behind a stale flag."""
        results = recommend_advice_sections(
            self.sections(), self.matter, conditions={"has_3_day_notice": True}
        )

        unreviewed = next(
            entry for entry in results if entry["section"].slug == "unreviewed"
        )
        self.assertTrue(unreviewed["needsReview"])
        self.assertIn("tracked changes", unreviewed["reviewReason"])

    def test_reviewed_only_excludes_flagged_sections(self):
        results = recommend_advice_sections(
            self.sections(),
            self.matter,
            conditions={"has_3_day_notice": True},
            reviewed_only=True,
        )

        self.assertNotIn("unreviewed", [entry["section"].slug for entry in results])

    def test_an_unmet_condition_is_reported_not_hidden(self):
        results = recommend_advice_sections(self.sections(), self.matter, conditions={})

        decarlo = next(entry for entry in results if entry["section"].slug == "decarlo")
        self.assertIn("has_3_day_notice", decarlo["unmetConditions"])

    def test_an_unrelated_case_matches_nothing(self):
        unrelated = Matter.objects.create(
            external_id="2026-CVG-78",
            client_name="Sam Okafor",
            matter_type="Benefits",
            jurisdiction="Cleveland Municipal Court",
            summary="Client asks about a utility shutoff.",
        )

        results = recommend_advice_sections(self.sections(), unrelated, conditions={})

        self.assertEqual(results, [])
