"""Loss-minimizing conversion of ordinary DOCX files into prepared templates.

The converter edits WordprocessingML in place. It does not round-trip through
HTML, Markdown, or plain text, so styles, numbering, tables, headers, footers,
images, section settings, and relationships remain in the package.

A prepared template keeps the maintained original's wording. Only language the
author marked as variable -- square brackets, underscore blanks, highlighting --
becomes a Jinja binding, so an export reads as the original document with the
case's details filled in rather than as a regenerated lookalike.

How much freedom the model gets is a per-block decision recorded as
`ai_latitude`:

``locked``
    Captions, certificates of service, signature blocks, quoted statutes, and
    citation-bearing legal standards. Rendered verbatim; only fill-ins vary.
``guided``
    The original prose is a strong starting draft that still needs adapting to
    the case. It renders literally, and an accepted rewrite from review replaces
    it through ``blocks[<key>]["revision"]``.
``generate``
    The original says "insert case specific facts" or similar. The instruction
    becomes the model's prompt instead of output text, and the model supplies
    the paragraphs.
"""

from __future__ import annotations

import hashlib
import re
import shutil
from collections import defaultdict
from copy import deepcopy
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path

import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from django.utils.text import slugify

from apps.templates_app.placeholders import (
    BRACKET_RE,
    PLACEHOLDER_ALIASES,
    convert_paragraph,
    convert_text,
    is_instruction,
)


MANIFEST_VERSION = 2
# Bump when conversion semantics change without changing the maintained source
# file, so startup can refresh already-prepared packages safely.
CONVERTER_VERSION = "2026-08-08-signature-rule-bindings-v1"

LATITUDE_LOCKED = "locked"
LATITUDE_GUIDED = "guided"
LATITUDE_GENERATE = "generate"

# Blocks whose wording is fixed by rule, court practice, or citation accuracy.
LOCKED_BLOCK_TYPES = {"caption", "certificate", "signature"}
# Blocks that carry the argument and are expected to be adapted per case.
GUIDED_BLOCK_TYPES = {"facts", "argument", "relief"}

AI_FILL_MODE_BY_LATITUDE = {
    LATITUDE_LOCKED: "none",
    LATITUDE_GUIDED: "revision_on_request",
    LATITUDE_GENERATE: "constrained_generation",
}

QUOTED_AUTHORITY_RE = re.compile(
    r"\b(?:R\.C\.|Civ\.R\.|Loc\.R\.|C\.F\.R\.|U\.S\.C\.|Cod\.Ord\.|Ohio App\.3d|N\.E\.2d|Ohio-\d{4})"
)
HEADING_WORDS = {
    "background",
    "caption",
    "certificate of service",
    "client goals",
    "conclusion",
    "deadlines",
    "defenses & strategy",
    "desired outcomes",
    "facts",
    "introduction",
    "law and argument",
    "memorandum in support",
    "prayer for relief",
    "relevant facts",
    "respectfully submitted",
    "signature",
    "statement of facts",
    "statement of relevant facts",
    "tasks (delete as you complete each task)",
}
ROMAN_HEADING_RE = re.compile(r"^(?:[IVXLCDM]+|[A-Z]|\d+)[.)]\s+", re.I)
LIST_PROMPT_RE = re.compile(
    r"(?:insert|add|describe|list|synopsis).*(?:fact|event|allegation|question|document|section)|case specific facts",
    re.I,
)
@dataclass
class BlockDefinition:
    key: str
    label: str
    block_type: str
    start: int
    end: int
    heading_index: int | None
    expects_list: bool = False

    @property
    def body_start(self):
        return self.start + 1 if self.heading_index == self.start else self.start


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def normalize_label(text: str) -> str:
    text = ROMAN_HEADING_RE.sub("", " ".join(text.split())).strip(" -–—:.")
    return text.title() if text.isupper() else text


def _visible_runs(paragraph):
    return [run for run in paragraph.runs if run.text.strip()]


def _is_emphasized(paragraph) -> bool:
    """True when every visible run is bold or underlined.

    The maintained originals rarely use Word's Heading styles. Section titles
    are Normal or List Paragraph text that has simply been bolded, so style name
    alone misses most of the document's structure.
    """
    runs = _visible_runs(paragraph)
    if not runs:
        return False
    return all(run.bold for run in runs) or all(run.underline for run in runs)


def _is_centered(paragraph) -> bool:
    return paragraph.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER


def is_heading(paragraph) -> bool:
    text = " ".join(paragraph.text.split()).strip()
    if not text or len(text) > 120:
        return False
    if BRACKET_RE.fullmatch(text) or text.strip("[] ").casefold() in PLACEHOLDER_ALIASES:
        return False
    style = (paragraph.style.name if paragraph.style else "").lower()
    normalized = ROMAN_HEADING_RE.sub("", text).strip(" -–—:.").lower()
    alpha = "".join(character for character in text if character.isalpha())
    if (
        style.startswith("heading")
        or normalized in HEADING_WORDS
        or (len(alpha) >= 4 and alpha.isupper() and len(text.split()) <= 12)
        or (bool(ROMAN_HEADING_RE.match(text)) and text.upper() == text)
        or text.lower().startswith("case caption")
    ):
        return True
    # A short, fully emphasized or centered line that does not end like prose is
    # a section title: "Complexity of Legal and Factual Issues".
    if len(text.split()) <= 12 and not text.endswith((".", ";", ",")):
        return _is_emphasized(paragraph) or _is_centered(paragraph)
    return False


def classify_block(label: str) -> str:
    lowered = label.lower()
    if "caption" in lowered:
        return "caption"
    if "certificate" in lowered or "service" in lowered:
        return "certificate"
    if "signature" in lowered or "respectfully submitted" in lowered:
        return "signature"
    if "fact" in lowered or "statement of the case" in lowered:
        return "facts"
    if "conclusion" in lowered or "relief" in lowered:
        return "relief"
    if "law" in lowered or "argument" in lowered or "standard" in lowered:
        return "argument"
    return "optional_clause"


def discover_blocks(document) -> list[BlockDefinition]:
    paragraphs = document.paragraphs
    nonempty = [index for index, paragraph in enumerate(paragraphs) if paragraph.text.strip()]
    if not nonempty:
        return [BlockDefinition("body", "Body", "optional_clause", 0, len(paragraphs), None)]

    boundaries = [index for index in nonempty if is_heading(paragraphs[index])]
    if not boundaries or boundaries[0] != nonempty[0]:
        boundaries.insert(0, nonempty[0])
    list_boundaries = set()
    for index in nonempty:
        if not LIST_PROMPT_RE.search(paragraphs[index].text):
            continue
        previous = max((boundary for boundary in boundaries if boundary < index), default=None)
        previous_type = classify_block(normalize_label(paragraphs[previous].text)) if previous is not None else "optional_clause"
        if previous_type != "facts":
            boundaries.append(index)
            list_boundaries.add(index)
    boundaries = sorted(set(boundaries))

    blocks = []
    keys = defaultdict(int)
    for position, start in enumerate(boundaries):
        end = boundaries[position + 1] if position + 1 < len(boundaries) else len(paragraphs)
        heading = is_heading(paragraphs[start]) and start not in list_boundaries
        if start in list_boundaries:
            raw_label = "Facts" if "fact" in paragraphs[start].text.lower() else "List items"
        else:
            raw_label = paragraphs[start].text if heading else ("Document body" if position == 0 else f"Section {position + 1}")
        label = normalize_label(raw_label) or f"Section {position + 1}"
        base_key = slugify(label)[:100] or f"section-{position + 1}"
        keys[base_key] += 1
        key = base_key if keys[base_key] == 1 else f"{base_key}-{keys[base_key]}"
        block_type = "facts" if start in list_boundaries else classify_block(label)
        sample = "\n".join(paragraph.text for paragraph in paragraphs[start:end])
        expects_list = bool(LIST_PROMPT_RE.search(sample))
        blocks.append(BlockDefinition(key, label, block_type, start, end, start if heading else None, expects_list))
    return blocks


def _set_paragraph_text_preserving_first_run(paragraph, text: str):
    runs = list(paragraph.runs)
    if runs:
        target = next((run for run in runs if run.text), runs[0])
        target.text = text
        for run in runs:
            if run is not target:
                run.text = ""
    else:
        paragraph.add_run(text)


def _marker_paragraph_like(paragraph, text: str):
    marker = deepcopy(paragraph._p)
    for child in list(marker):
        if child.tag != qn("w:pPr"):
            marker.remove(child)
    run = OxmlElement("w:r")
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    marker.append(run)
    return marker


def _replace_with_loop(paragraph, expression: str):
    paragraph._p.addprevious(_marker_paragraph_like(paragraph, "{%p for item in " + expression + " %}"))
    _set_paragraph_text_preserving_first_run(paragraph, "{{ item }}")
    paragraph._p.addnext(_marker_paragraph_like(paragraph, "{%p endfor %}"))


def _all_story_paragraphs(document):
    """Yield main, table, header, and footer paragraphs without duplicates."""
    seen = set()

    def emit(paragraphs):
        for paragraph in paragraphs:
            identity = paragraph._p
            if identity not in seen:
                seen.add(identity)
                yield paragraph

    yield from emit(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from emit(cell.paragraphs)
    for section in document.sections:
        yield from emit(section.header.paragraphs)
        yield from emit(section.footer.paragraphs)
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from emit(cell.paragraphs)
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from emit(cell.paragraphs)


def _converted_block_body(document, block: BlockDefinition) -> str:
    lines = []
    texts = [document.paragraphs[index].text.strip() for index in range(block.body_start, block.end)]
    for offset, text in enumerate(texts):
        if not text:
            continue
        index = block.body_start + offset
        converted, _conversion = convert_text(
            text, f"{block.key}_{index}", _nearby_text(texts, offset)
        )
        lines.append(converted)
    return "\n".join(lines)


def _nearby_text(texts, position) -> str:
    """The lines either side of one line, for a fill-in that stands alone."""
    previous = next((text for text in reversed(texts[:position]) if text.strip()), "")
    following = next((text for text in texts[position + 1 :] if text.strip()), "")
    return "\n".join(text for text in (previous, following) if text)


def block_instructions(document, block: BlockDefinition) -> list[str]:
    """Drafting directions the author wrote into a block, e.g. "[describe X]"."""
    instructions = []
    for index in range(block.body_start, block.end):
        text = document.paragraphs[index].text
        if not text.strip():
            continue
        for match in BRACKET_RE.finditer(text):
            label = " ".join(match.group(1).split())
            if label and is_instruction(label) and label not in instructions:
                instructions.append(label)
    return instructions


def _block_text(document, block: BlockDefinition) -> str:
    return "\n".join(
        document.paragraphs[index].text
        for index in range(block.body_start, block.end)
        if document.paragraphs[index].text.strip()
    )


def _has_fill_ins(document, block: BlockDefinition) -> bool:
    for index in range(block.body_start, block.end):
        paragraph = document.paragraphs[index]
        if not paragraph.text.strip():
            continue
        if BRACKET_RE.search(paragraph.text) or "___" in paragraph.text:
            return True
        if any(run.font.highlight_color is not None for run in paragraph.runs if run.text.strip()):
            return True
    return False


def classify_latitude(document, block: BlockDefinition) -> str:
    """Decide how much of a block the model may write.

    Wording that carries legal authority is locked because a paraphrase of a
    quoted statute or a regenerated citation is a correctness problem, not a
    style one. Blocks whose original text is an instruction to the drafter are
    the model's to write.
    """
    if block.expects_list:
        return LATITUDE_GENERATE
    if block.block_type in LOCKED_BLOCK_TYPES:
        return LATITUDE_LOCKED
    text = _block_text(document, block)
    if not text.strip():
        return LATITUDE_GENERATE
    instructions = block_instructions(document, block)
    words = len(text.split())
    if instructions and words < 60:
        # Little more than the instruction itself: nothing to preserve.
        return LATITUDE_GENERATE
    # Quoted authority and citation strings must survive verbatim.
    if len(QUOTED_AUTHORITY_RE.findall(text)) >= 3 and not instructions:
        return LATITUDE_LOCKED
    # Fill-ins are the author's own signal that the passage is case-specific, so
    # a sub-section like "Time Needed for Discovery" is adaptable even though its
    # heading carries no argument keyword.
    if block.block_type in GUIDED_BLOCK_TYPES or instructions or _has_fill_ins(document, block):
        return LATITUDE_GUIDED
    return LATITUDE_LOCKED


OR_MARKER_RE = re.compile(r"^\[?\s*or\s*\]?[.:]?$", re.I)

# Choice values are named for what distinguishes the alternatives, following
# Docassemble/AssemblyLine's snake_case variable and option naming.
ALTERNATIVE_OPTION_CUES = (
    ("email", ("by email", "electronic mail", "e-mail", "electronically")),
    ("mail", ("united states mail", "u.s. mail", "regular mail", "ordinary mail", "certified mail")),
    ("personal", ("personal service", "personally serv", "hand deliver", "in person")),
    ("courier", ("courier", "commercial carrier", "bonnie speed")),
    ("fax", ("facsimile", "by fax")),
)


def _option_key(text: str, position: int) -> str:
    lowered = " ".join(text.split()).casefold()
    for key, cues in ALTERNATIVE_OPTION_CUES:
        if any(cue in lowered for cue in cues):
            return key
    return f"option_{position}"


def _choice_variable(block: BlockDefinition | None) -> str:
    if block is not None and block.block_type == "certificate":
        return "service_method"
    base = (block.key if block else "clause").replace("-", "_")
    return f"{base}_option"


def _is_or_marker(paragraph) -> bool:
    return bool(OR_MARKER_RE.match(" ".join(paragraph.text.split())))


def _alternative_groups(document, blocks, original_paragraphs):
    """Find "A / [OR] / B" runs the author left for the drafter to choose between.

    Ownership is resolved against the paragraph list as it stood before any
    markers were inserted, because a block's bounds are indices into that list
    and every insertion shifts the current one.
    """
    owner = {}
    for block in blocks:
        for index in range(block.body_start, block.end):
            if 0 <= index < len(original_paragraphs):
                owner[id(original_paragraphs[index]._p)] = block

    live = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip()
        and paragraph._p.getparent() is not None
        and not paragraph.text.lstrip().startswith("{%")
    ]

    groups = []
    index = 0
    while index < len(live):
        if not _is_or_marker(live[index]) or index == 0 or index + 1 >= len(live):
            index += 1
            continue
        members = [live[index - 1]]
        markers = []
        cursor = index
        # Consecutive markers chain into one choice with three or more options.
        while cursor + 1 < len(live) and _is_or_marker(live[cursor]):
            markers.append(live[cursor])
            members.append(live[cursor + 1])
            cursor += 2
        groups.append((owner.get(id(members[0]._p)), members, markers))
        index = cursor
    return groups


def bind_alternatives(document, blocks, original_paragraphs) -> list[dict]:
    """Turn an editorial "[OR]" into a chooseable conditional.

    The maintained originals put both a certificate of service by email and one
    by ordinary mail in the document with "[OR]" between them, expecting the
    advocate to delete the inapplicable one. Left literal, a generated filing
    certifies service twice and prints "[OR]" in between.

    The first alternative is also the default, so an unanswered choice still
    produces a complete certificate rather than deleting the passage.
    """
    choices = []
    for block, members, markers in _alternative_groups(document, blocks, original_paragraphs):
        variable = _choice_variable(block)
        keys = []
        for position, paragraph in enumerate(members, start=1):
            key = _option_key(paragraph.text, position)
            if key in keys:
                key = f"{key}_{position}"
            keys.append(key)

        first, last = members[0], members[-1]
        first._p.addprevious(
            _marker_paragraph_like(
                first,
                '{%p if ' + variable + ' == "' + keys[0] + '" or not ' + variable + ' %}',
            )
        )
        for marker, key in zip(markers, keys[1:]):
            _set_paragraph_text_preserving_first_run(
                marker, '{%p elif ' + variable + ' == "' + key + '" %}'
            )
        last._p.addnext(_marker_paragraph_like(last, "{%p endif %}"))

        choices.append(
            {
                "name": variable,
                "label": re.sub(r"\s+", " ", variable.replace("_", " ")).strip().title(),
                "options": keys,
                "default": keys[0],
                "block": block.key if block else "",
            }
        )
    return choices


def _wrap_revisable_block(paragraphs, block_key: str):
    """Let an accepted revision stand in for the maintained wording.

    The original prose renders by default. When the advocate edits the section
    or accepts a proposed rewrite, `blocks[<key>]["revision"]` is set and the
    whole block is replaced.

    This wrapping is not what `ai_latitude` controls. Latitude governs whether
    the *model* may propose a rewrite; a human edit must always reach the export,
    including in a locked block, or the editor would silently discard it.
    """
    live = [paragraph for paragraph in paragraphs if paragraph._p.getparent() is not None]
    if not live:
        return
    first, last = live[0], live[-1]
    first._p.addprevious(
        _marker_paragraph_like(first, f'{{%p if blocks["{block_key}"]["revision"] %}}')
    )
    revision = _marker_paragraph_like(first, f'{{{{ blocks["{block_key}"]["revision"] }}}}')
    otherwise = _marker_paragraph_like(first, "{%p else %}")
    first._p.addprevious(revision)
    first._p.addprevious(otherwise)
    last._p.addnext(_marker_paragraph_like(last, "{%p endif %}"))


def annotate_document(document, blocks: list[BlockDefinition]) -> dict:
    """Bind fill-ins and AI slots while leaving the author's wording in place."""
    fields = set()
    flags = set()
    main_paragraphs = list(document.paragraphs)
    latitudes = {block.key: classify_latitude(document, block) for block in blocks}

    # Blocks the author left to the drafter become model-written paragraphs. The
    # instruction text is the prompt, so it must not survive into the export.
    for block in blocks:
        if latitudes[block.key] != LATITUDE_GENERATE:
            continue
        body_paragraphs = [
            main_paragraphs[index]
            for index in range(block.body_start, block.end)
            if main_paragraphs[index].text.strip()
        ]
        if not body_paragraphs:
            continue
        list_prompt = next((p for p in body_paragraphs if LIST_PROMPT_RE.search(p.text)), None)
        loop_target = list_prompt or body_paragraphs[0]
        collection = "items" if block.expects_list else "paragraphs"
        _replace_with_loop(loop_target, f'blocks["{block.key}"]["{collection}"]')
        for paragraph in body_paragraphs:
            if paragraph is loop_target or paragraph._p.getparent() is None:
                continue
            # Remaining instruction-only lines would otherwise print as prose.
            if BRACKET_RE.fullmatch(" ".join(paragraph.text.split())):
                paragraph._p.getparent().remove(paragraph._p)

    # Every surviving paragraph keeps its wording; only fill-ins are rebound.
    story = list(_all_story_paragraphs(document))
    story_texts = [paragraph.text for paragraph in story]
    for index, paragraph in enumerate(story):
        conversion = convert_paragraph(
            paragraph, f"placeholder_{index + 1}", _nearby_text(story_texts, index)
        )
        fields.update(conversion.fields)
        flags.update(conversion.flags)

    for block in blocks:
        if latitudes[block.key] == LATITUDE_GENERATE:
            # Already bound to a model-written collection.
            continue
        body_paragraphs = [
            main_paragraphs[index]
            for index in range(block.body_start, block.end)
            if main_paragraphs[index].text.strip()
        ]
        _wrap_revisable_block(body_paragraphs, block.key)

    # Runs last so the choice conditional nests inside the revision wrapper
    # rather than straddling it.
    choices = bind_alternatives(document, blocks, main_paragraphs)

    return {
        "fields": sorted(fields),
        "flags": sorted(flags),
        "latitudes": latitudes,
        "choices": choices,
    }


def _copy_block_document(source_path: Path, output_path: Path, block: BlockDefinition):
    document = Document(source_path)
    body = document._body._element
    children = list(body)
    paragraphs = document.paragraphs
    first_nonempty = next((index for index, paragraph in enumerate(paragraphs) if paragraph.text.strip()), 0)
    if block.start <= first_nonempty:
        # Captions and other leading layout are often tables or text boxes that
        # python-docx omits from document.paragraphs. Keep those OOXML elements
        # with the first block instead of silently dropping them.
        start_position = 0
    else:
        start_element = paragraphs[min(block.start, len(paragraphs) - 1)]._p
        start_position = children.index(start_element)
    if block.end < len(paragraphs):
        end_position = children.index(paragraphs[block.end]._p)
    else:
        end_position = next((index for index, child in enumerate(children) if child.tag == qn("w:sectPr")), len(children))
    keep = set(children[start_position:end_position])
    for child in children:
        if child.tag == qn("w:sectPr"):
            continue
        if child not in keep:
            body.remove(child)
    local_paragraphs = document.paragraphs
    local_block = BlockDefinition(
        key=block.key,
        label=block.label,
        block_type=block.block_type,
        start=0,
        end=len(local_paragraphs),
        heading_index=0 if block.heading_index is not None else None,
        expects_list=block.expects_list,
    )
    annotate_document(document, [local_block])
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def infer_kind(title: str) -> str:
    lowered = title.lower()
    if "motion" in lowered:
        return "motion"
    if "appeal" in lowered or "affidavit" in lowered or "notice" in lowered:
        return "brief"
    return "shell"


def infer_goal(title: str, kind: str) -> str:
    if kind == "motion":
        if "motion" in title.casefold():
            return f"Draft {title} with case-specific facts, legal grounds, and requested relief."
        return f"Draft the {title} motion with case-specific facts, legal grounds, and requested relief."
    if kind == "brief":
        return f"Draft the {title} filing with case-specific facts, legal grounds, and requested relief."
    return f"Draft the {title} document with case-specific facts and the requested relief or outcome."


def infer_description(title: str) -> str:
    return f"Maintained Word template for {title}."


def ingest_docx(source: Path, prepared_root: Path, snippets_root: Path, *, force=False) -> Path:
    source = source.resolve()
    slug = slugify(source.stem) or "document-template"
    package_dir = prepared_root / slug
    package_dir.mkdir(parents=True, exist_ok=True)
    manifest_path = package_dir / "manifest.yaml"
    source_checksum = sha256_file(source)
    if manifest_path.exists() and not force:
        existing = yaml.safe_load(manifest_path.read_text()) or {}
        if (
            existing.get("source", {}).get("sha256") == source_checksum
            and existing.get("source", {}).get("converter_version") == CONVERTER_VERSION
        ):
            return manifest_path

    original = Document(source)
    blocks = discover_blocks(original)
    block_bodies = {block.key: _converted_block_body(original, block) for block in blocks}

    annotated = Document(source)
    discovery = annotate_document(annotated, blocks)
    template_path = package_dir / "template.docx"
    annotated.save(template_path)

    latitudes = discovery["latitudes"]
    block_rows = []
    for order, block in enumerate(blocks, start=1):
        relative_block_path = Path("docx-snippets") / slug / "blocks" / f"{block.key}.docx"
        block_path = snippets_root / slug / "blocks" / f"{block.key}.docx"
        _copy_block_document(source, block_path, block)
        latitude = latitudes.get(block.key, LATITUDE_LOCKED)
        block_rows.append(
            {
                "key": block.key,
                "label": block.label,
                "type": block.block_type,
                "order": order * 10,
                "required": True,
                "editable": latitude != LATITUDE_LOCKED,
                "ai_latitude": latitude,
                "ai_fill_mode": AI_FILL_MODE_BY_LATITUDE[latitude],
                "instructions": block_instructions(original, block),
                "body": block_bodies[block.key],
                "docx": relative_block_path.as_posix(),
                "sha256": sha256_file(block_path),
                "input": {
                    "type": "array" if block.expects_list else "rich_text",
                    "items": {"type": "string"} if block.expects_list else None,
                },
                "lexical": {
                    "node": "list" if block.expects_list else "paragraphs",
                    "listType": "number" if block.expects_list else None,
                    "sourceParagraphRange": [block.start, block.end],
                },
            }
        )

    try:
        source_path = source.relative_to(prepared_root.parent.resolve()).as_posix()
    except ValueError:
        source_path = source.as_posix()
    kind = infer_kind(source.stem)
    manifest = {
        "schema_version": MANIFEST_VERSION,
        "slug": slug,
        "title": source.stem,
        "kind": kind,
        "description": infer_description(source.stem),
        "goal": infer_goal(source.stem, kind),
        "negative_goal": "",
        "aliases": [],
        "jurisdiction": "Ohio",
        "source_label": "Content library",
        "active": True,
        "render": {"strategy": "full_document", "docx": "template.docx"},
        "source": {
            "path": source_path,
            "sha256": source_checksum,
            "converted_at": datetime.now(timezone.utc).isoformat(),
            "converter": "apps.templates_app.ingestion",
            "converter_version": CONVERTER_VERSION,
            "format_preservation": "in_place_ooxml",
        },
        "fields": discovery["fields"],
        "flags": discovery["flags"],
        "choices": discovery["choices"],
        "blocks": block_rows,
    }
    manifest_path.write_text(yaml.safe_dump(manifest, sort_keys=False, allow_unicode=True))
    return manifest_path


def promote_shared_blocks(manifest_paths: list[Path], snippets_root: Path, *, force=False):
    candidates = defaultdict(list)
    for manifest_path in manifest_paths:
        manifest = yaml.safe_load(manifest_path.read_text()) or {}
        for block in manifest.get("blocks", []):
            if block.get("type") not in {"caption", "signature", "certificate"}:
                continue
            normalized = re.sub(r"\s+", " ", block.get("body", "")).strip().lower()
            if normalized:
                candidates[(block["type"], hashlib.sha256(normalized.encode()).hexdigest())].append(block)
    promoted = []
    for (block_type, _digest), rows in candidates.items():
        if len(rows) < 2:
            continue
        source = snippets_root.parent / rows[0]["docx"]
        destination = snippets_root / "_shared" / "blocks" / f"{block_type}.docx"
        destination.parent.mkdir(parents=True, exist_ok=True)
        if not destination.exists() or force:
            shutil.copy2(source, destination)
        promoted.append(destination)
    return promoted


def ingest_directory(source_root: Path, prepared_root: Path, snippets_root: Path, *, force=False):
    from apps.templates_app.spreadsheets import ingest_xlsx

    manifests = [
        ingest_docx(path, prepared_root, snippets_root, force=force)
        for path in sorted(source_root.rglob("*.docx"))
        if not path.name.startswith("~$")
    ]
    promote_shared_blocks(manifests, snippets_root, force=force)
    # Some maintained exhibits are workbooks rather than filings.
    manifests += [
        ingest_xlsx(path, prepared_root, force=force)
        for path in sorted(source_root.rglob("*.xlsx"))
        if not path.name.startswith("~$")
    ]
    return manifests
