"""Generate the neutral letterhead a fresh checkout draws letters on.

Real organization stationery is private, so the public content library ships a
plain, obviously-generic letterhead instead. It carries the same variables as a
real one, which means the letter workflow, its tests, and a first-run demo all
work before anyone uploads their own.
"""

from pathlib import Path

import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from django.core.management.base import BaseCommand

from apps.core.content_library import content_library_dir
from apps.templates_app.letterhead_library import LETTERHEAD_DIR, PLACEHOLDER_SLUG, sync_letterheads


class Command(BaseCommand):
    help = "Write the neutral placeholder letterhead into the public content library."

    def add_arguments(self, parser):
        parser.add_argument("--no-sync", action="store_true", help="Do not index the result afterwards.")

    def handle(self, *args, **options):
        package = content_library_dir() / LETTERHEAD_DIR / PLACEHOLDER_SLUG
        package.mkdir(parents=True, exist_ok=True)

        document = Document()
        section = document.sections[0]

        masthead = section.header.paragraphs[0]
        masthead.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title = masthead.add_run("EXAMPLE LEGAL AID SOCIETY")
        title.bold = True
        title.font.size = Pt(16)
        tagline = section.header.add_paragraph()
        tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note = tagline.add_run(
            "Placeholder letterhead - replace this in Django admin before sending mail"
        )
        note.italic = True
        note.font.size = Pt(8)

        contact = section.header.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for text in (
            "{{ advocate_name }}",
            "{{ advocate_title }}",
            "Phone:  {{ advocate_phone }}",
            "{%p if advocate_fax %}Fax:  {{ advocate_fax }}{%p endif %}",
            "{{ advocate_email }}",
        ):
            line = section.header.add_paragraph(text)
            line.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in line.runs:
                run.font.size = Pt(9)

        office = section.header.add_paragraph("{{ office_name }} - {{ office_address }}")
        office.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in office.runs:
            run.font.size = Pt(8)

        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.text = "Letter to {{ letter_subject }}, {{ letter_date }}"
        for run in footer.runs:
            run.font.size = Pt(8)

        # The letter body is composed into this document, so the page itself
        # stays empty apart from the stationery.
        document.add_paragraph("")
        document.save(package / "letterhead.docx")

        manifest = {
            "schema_version": 1,
            "slug": PLACEHOLDER_SLUG,
            "title": "Example Legal Aid Society (placeholder)",
            "organization": "Example Legal Aid Society",
            "description": (
                "Neutral letterhead shipped so a fresh install can draft letters. "
                "Replace it with your organization's stationery in Django admin."
            ),
            "docx": "letterhead.docx",
            "default": True,
            "placeholder": True,
            "active": True,
            "variables": [
                "advocate_name",
                "advocate_title",
                "advocate_phone",
                "advocate_fax",
                "advocate_email",
                "office_name",
                "office_address",
                "letter_subject",
                "letter_date",
            ],
        }
        (package / "manifest.yaml").write_text(
            yaml.safe_dump(manifest, sort_keys=False, allow_unicode=True)
        )

        if not options["no_sync"]:
            sync_letterheads()
        self.stdout.write(self.style.SUCCESS(f"Wrote placeholder letterhead to {package}"))
