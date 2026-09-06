#!/usr/bin/env python3
"""Generate realistic Ohio eviction defense briefs in DOCX format for Argument Gym validation.

Produces:
1. 01_good_motion_to_dismiss_notice_defect_franklin_county.docx (Well-supported, compliant Civ.R. 12(B)(1) motion on R.C. 1923.04 notice defect)
2. 02_good_trial_brief_retaliation_and_habitability_cleveland.docx (Well-supported tenant trial brief under R.C. 5321.02 and R.C. 5321.04 with exhibits)
3. 03_defective_brief_unsupported_retaliation_and_habitability_hamilton.docx (Substantively flawed: unmet statutory elements, missing causal links, unanchored habitability claims)
4. 04_defective_brief_procedural_and_form_errors_akron.docx (Form flaws: unfilled placeholders, paragraph numbering gaps, legal typos, missing cert of service, invalid authority)
"""

import os
from pathlib import Path
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn


def create_base_document(margins_in=1.0, default_font="Times New Roman", default_size_pt=12, line_spacing=2.0):
    """Create a new Document with standardized legal styling and margins."""
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(margins_in)
        section.bottom_margin = Inches(margins_in)
        section.left_margin = Inches(margins_in)
        section.right_margin = Inches(margins_in)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)

    # Set normal style font
    style = doc.styles['Normal']
    font = style.font
    font.name = default_font
    font.size = Pt(default_size_pt)
    font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.line_spacing = line_spacing
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    return doc


def add_court_caption(doc, court_name, plaintiff, defendant, case_no, judge_name, doc_title):
    """Add a standard Ohio trial court caption using a clean two-column table."""
    p_court = doc.add_paragraph()
    p_court.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_court.paragraph_format.line_spacing = 1.0
    p_court.paragraph_format.space_after = Pt(12)
    run_court = p_court.add_run(court_name.upper() + "\n")
    run_court.bold = True
    run_court.font.name = "Times New Roman"
    run_court.font.size = Pt(12)

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    table.columns[0].width = Inches(3.4)
    table.columns[1].width = Inches(3.1)

    # Left cell (Parties)
    cell_left = table.cell(0, 0)
    p_left = cell_left.paragraphs[0]
    p_left.paragraph_format.line_spacing = 1.15
    p_left.paragraph_format.space_after = Pt(0)

    r = p_left.add_run(f"{plaintiff}\n")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    
    r_role1 = p_left.add_run("Plaintiff / Landlord,\n\n")
    r_role1.font.name = "Times New Roman"
    r_role1.font.size = Pt(10)
    
    r_v = p_left.add_run("v.\n\n")
    r_v.font.name = "Times New Roman"
    r_v.font.size = Pt(11)

    r_def = p_left.add_run(f"{defendant}\n")
    r_def.bold = True
    r_def.font.name = "Times New Roman"
    r_def.font.size = Pt(11)

    r_role2 = p_left.add_run("Defendant / Tenant.")
    r_role2.font.name = "Times New Roman"
    r_role2.font.size = Pt(10)

    # Right cell (Case info and Title)
    cell_right = table.cell(0, 1)
    p_right = cell_right.paragraphs[0]
    p_right.paragraph_format.line_spacing = 1.15
    p_right.paragraph_format.space_after = Pt(0)

    r_case = p_right.add_run(f"CASE NO. {case_no}\n")
    r_case.bold = True
    r_case.font.name = "Times New Roman"
    r_case.font.size = Pt(11)

    if judge_name:
        r_judge = p_right.add_run(f"JUDGE {judge_name.upper()}\n\n")
        r_judge.font.name = "Times New Roman"
        r_judge.font.size = Pt(10)
    else:
        p_right.add_run("\n")

    r_title = p_right.add_run(f"{doc_title.upper()}")
    r_title.bold = True
    r_title.font.name = "Times New Roman"
    r_title.font.size = Pt(11)

    p_space = doc.add_paragraph()
    p_space.paragraph_format.line_spacing = 1.0
    p_space.paragraph_format.space_after = Pt(6)


def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return p


def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return p


def add_body_paragraph(doc, text, indent=0.5, line_spacing=2.0, font_name="Times New Roman", font_size_pt=12):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(indent)
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    return p


def add_signature_block(doc, name, bar_no, title, org, address, phone, email, rep_for):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Inches(0.0)

    p.add_run("Respectfully submitted,\n\n\n")
    
    r_sig = p.add_run(f"/s/ {name}\n")
    r_sig.bold = True
    p.add_run(f"{name} ({bar_no})\n")
    if title:
        p.add_run(f"{title}\n")
    if org:
        p.add_run(f"{org}\n")
    p.add_run(f"{address}\n")
    p.add_run(f"Telephone: {phone}\n")
    p.add_run(f"Email: {email}\n\n")
    p.add_run(f"Attorney for {rep_for}\n")


def add_certificate_of_service(doc, service_text, attorney_name):
    p_heading = doc.add_paragraph()
    p_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_heading.paragraph_format.line_spacing = 1.0
    p_heading.paragraph_format.space_before = Pt(18)
    p_heading.paragraph_format.space_after = Pt(6)
    r_head = p_heading.add_run("CERTIFICATE OF SERVICE")
    r_head.bold = True
    r_head.font.name = "Times New Roman"
    r_head.font.size = Pt(12)

    p_body = doc.add_paragraph()
    p_body.paragraph_format.first_line_indent = Inches(0.5)
    p_body.paragraph_format.line_spacing = 1.5
    p_body.paragraph_format.space_after = Pt(12)
    r_body = p_body.add_run(service_text)
    r_body.font.name = "Times New Roman"
    r_body.font.size = Pt(12)

    p_sig = doc.add_paragraph()
    p_sig.paragraph_format.line_spacing = 1.0
    p_sig.paragraph_format.space_before = Pt(6)
    r_sig = p_sig.add_run(f"/s/ {attorney_name}\n{attorney_name}")
    r_sig.font.name = "Times New Roman"
    r_sig.font.size = Pt(12)


def generate_brief_1(output_path):
    """Generate Brief 1 (Good): Motion to Dismiss for Defective Notice under R.C. 1923.04."""
    doc = create_base_document()

    add_court_caption(
        doc,
        court_name="IN THE FRANKLIN COUNTY MUNICIPAL COURT\nENVIRONMENTAL DIVISION\nCOLUMBUS, OHIO",
        plaintiff="OAKWOOD MANOR APARTMENTS, LLC",
        defendant="MARCUS VANCE",
        case_no="2024 CVG 014892",
        judge_name="STEPHANIE MINGO",
        doc_title="DEFENDANT'S MOTION TO DISMISS FIRST CAUSE OF ACTION (FORCIBLE ENTRY AND DETAINER) PURSUANT TO CIV.R. 12(B)(1) AND R.C. 1923.04"
    )

    add_body_paragraph(
        doc,
        "Now comes Defendant Marcus Vance, by and through undersigned counsel, and respectfully moves this Court pursuant to Ohio Rule of Civil Procedure 12(B)(1) and Ohio Revised Code § 1923.04 for an Order dismissing Plaintiff's First Cause of Action (Forcible Entry and Detainer) for lack of subject-matter jurisdiction. As established in the accompanying Memorandum in Support, Plaintiff failed to serve a statutorily compliant three-day notice to leave the premises, and prematurely commenced this action prior to the expiration of the full statutory waiting period. Consequently, this Court lacks jurisdiction over Plaintiff's eviction claim."
    )

    add_heading_1(doc, "MEMORANDUM IN SUPPORT")
    add_heading_1(doc, "I. STATEMENT OF THE FACTS")

    add_body_paragraph(
        doc,
        "Defendant Marcus Vance has resided at 1428 East Broad Street, Apartment 3B, Columbus, Ohio 43205 since March 1, 2022, pursuant to a written residential lease agreement with Plaintiff Oakwood Manor Apartments, LLC."
    )
    add_body_paragraph(
        doc,
        "On Friday, May 3, 2024, at approximately 5:30 p.m., Plaintiff's property manager taped a paper document entitled 'Notice of Termination' to the exterior door of Defendant's apartment unit. A true and accurate copy of the paper posted by Plaintiff is attached hereto as Exhibit A."
    )
    add_body_paragraph(
        doc,
        "The paper posted by Plaintiff ordered Defendant to vacate within three days, but entirely omitted the mandatory conspicuous statutory warning required by R.C. 1923.04(A). Specifically, the notice contained no warning advising Defendant that if he did not leave, an eviction action might be filed against him, nor did it contain the required statutory language regarding seeking legal assistance."
    )
    add_body_paragraph(
        doc,
        "Furthermore, Plaintiff filed the instant Forcible Entry and Detainer Complaint in this Court on Wednesday, May 8, 2024. Counting from the date of service pursuant to R.C. 1.14 and Civ.R. 6(A)—which excludes the day of service (May 3) and intermediate weekend days (Saturday, May 4, and Sunday, May 5)—only two full business days elapsed before Plaintiff filed its eviction complaint."
    )

    add_heading_1(doc, "II. LAW AND ARGUMENT")
    add_heading_2(doc, "A. Proper Service of a Statutorily Compliant Three-Day Notice under R.C. 1923.04 Is a Mandatory Jurisdictional Precondition to Eviction.")

    add_body_paragraph(
        doc,
        "In Ohio, compliance with R.C. 1923.04 is a mandatory prerequisite to the jurisdiction of a municipal court in a forcible entry and detainer action. See Bella Vista Apts. v. Herzner, 125 Ohio Misc.2d 1, 2003-Ohio-4872, 796 N.E.2d 593 (M.C.); Voyager Village Ltd. v. Lehman, 161 Ohio App.3d 192, 2005-Ohio-2451, 829 N.E.2d 755 (2d Dist.). A landlord's failure to strictly comply with the notice requirements of R.C. 1923.04 deprives the trial court of subject-matter jurisdiction and requires dismissal of the complaint pursuant to Civ.R. 12(B)(1)."
    )

    add_heading_2(doc, "B. Plaintiff's Notice Omitted the Mandatory Statutory Language Mandated by R.C. 1923.04(A).")

    add_body_paragraph(
        doc,
        "R.C. 1923.04(A) explicitly dictates that every notice to leave the premises served upon a residential tenant must contain the following language, printed in a conspicuous manner: 'YOU ARE BEING ASKED TO LEAVE THE PREMISES. IF YOU DO NOT LEAVE, AN EVICTION ACTION MAY BE SUBSTANTIALLY FILED AGAINST YOU. IF YOU ARE IN DOUBT REGARDING YOUR LEGAL RIGHTS AND OBLIGATIONS AS A TENANT, IT IS RECOMMENDED THAT YOU SEEK LEGAL ASSISTANCE.'"
    )
    add_body_paragraph(
        doc,
        "The Supreme Court of Ohio and appellate courts across this state have repeatedly held that the inclusion of this exact statutory language is mandatory. See Mastics v. McGrew, 8th Dist. Cuyahoga No. 39682, 1979 WL 210543 (Dec. 6, 1979); Cincinnati Metro. Hous. Auth. v. Morgan, 104 Ohio St.3d 445, 2004-Ohio-6554, 820 N.E.2d 315. In the case at bar, as demonstrated by Exhibit A, Plaintiff's notice contained no such statutory language whatsoever. Because the statutory warning language is absent, the notice is void ab initio, and cannot support a forcible entry and detainer action."
    )

    add_heading_2(doc, "C. Plaintiff Prematurely Commenced this Action Prior to the Expiration of the Statutory Three-Day Waiting Period.")

    add_body_paragraph(
        doc,
        "Even assuming arguendo that the notice was substantively sufficient, Plaintiff failed to wait the requisite statutory three-day period before filing its complaint. Under R.C. 1923.04(A), a landlord must serve the notice 'at least three or more days before commencing the action.' When computing time under R.C. 1923.04, Ohio courts apply R.C. 1.14 and Civ.R. 6(A). See Wintrow v. Smith, 32 Ohio Misc.2d 12, 513 N.E.2d 363 (M.C. 1987); Dennis v. Morgan, 89 Ohio St.3d 417, 732 N.E.2d 391 (2000)."
    )
    add_body_paragraph(
        doc,
        "Pursuant to Civ.R. 6(A), the day of service (Friday, May 3) is excluded. Furthermore, because the prescribed period is less than seven days, intermediate Saturdays and Sundays are excluded from the computation. Therefore, the first day counted was Monday, May 6, 2024; the second day was Tuesday, May 7; and the third full day was Wednesday, May 8. Plaintiff was prohibited from filing an eviction action until Thursday, May 9, 2024. By filing its complaint on May 8, Plaintiff filed prematurely. Under settled Ohio law, a premature filing under R.C. 1923.04 is a fatal defect that deprives the court of jurisdiction and warrants immediate dismissal. See Showe Mgt. Corp. v. Hazelbaker, 12th Dist. Fayette No. CA2005-11-031, 2006-Ohio-6356."
    )

    add_heading_1(doc, "III. CONCLUSION AND PRAYER FOR RELIEF")

    add_body_paragraph(
        doc,
        "WHEREFORE, Defendant Marcus Vance respectfully requests that this Court issue an Order dismissing Plaintiff's First Cause of Action (Forcible Entry and Detainer) with prejudice for lack of subject-matter jurisdiction, awarding Defendant his reasonable costs and statutory fees, and granting such other and further relief as the Court deems just and equitable."
    )

    add_signature_block(
        doc,
        name="Clara E. Jenkins",
        bar_no="0098412",
        title="Staff Attorney",
        org="The Legal Aid Society of Columbus",
        address="1108 South High Street, Columbus, OH 43206",
        phone="(614) 224-8374",
        email="cjenkins@columbuslegalaid.org",
        rep_for="Defendant Marcus Vance"
    )

    add_certificate_of_service(
        doc,
        service_text="I hereby certify that on this 14th day of May, 2024, a true and correct copy of the foregoing Motion to Dismiss was served via electronic mail and ordinary U.S. Mail, postage prepaid, upon Plaintiff's counsel: Gregory T. Stone, Esq., Stone & Associates Co., L.P.A., 250 Civic Center Drive, Suite 400, Columbus, OH 43215 (gstone@stonelawcolumbus.com).",
        attorney_name="Clara E. Jenkins"
    )

    doc.save(output_path)


def generate_brief_2(output_path):
    """Generate Brief 2 (Good): Tenant Trial Brief on Retaliation (R.C. 5321.02) and Habitability (R.C. 5321.04)."""
    doc = create_base_document()

    add_court_caption(
        doc,
        court_name="IN THE CLEVELAND MUNICIPAL COURT\nHOUSING DIVISION\nCUYAHOGA COUNTY, OHIO",
        plaintiff="SUPERIOR PROPERTY MANAGEMENT CO.",
        defendant="ELENA ROSTOVA",
        case_no="2024-CVG-008912",
        judge_name="W. MONÁ SCOTT",
        doc_title="DEFENDANT'S TRIAL BRIEF IN OPPOSITION TO RESTITUTION OF PREMISES AND IN SUPPORT OF DEFENSES UNDER R.C. 5321.02 AND R.C. 5321.04"
    )

    add_body_paragraph(
        doc,
        "Defendant Elena Rostova, by and through undersigned legal counsel, respectfully submits this Trial Brief prior to the eviction hearing scheduled for November 14, 2024. As set forth below, Plaintiff's action for possession is barred because it constitutes an unlawful retaliatory eviction under R.C. 5321.02 following Defendant's formal housing code complaint to the City of Cleveland. Furthermore, Plaintiff's continuous and willful breach of its statutory duty to maintain fit and habitable premises under R.C. 5321.04(A) bars recovery of possession and entitles Defendant to rent abatement."
    )

    add_heading_1(doc, "I. STATEMENT OF FACTS AND PROCEDURAL HISTORY")

    add_body_paragraph(
        doc,
        "Defendant Elena Rostova has resided in unit 4 of the multi-family residential building located at 3241 West 84th Street, Cleveland, Ohio 44102 since January 2021. Monthly rent under the oral month-to-month tenancy is $750.00."
    )
    add_body_paragraph(
        doc,
        "Beginning in August 2024, Defendant experienced severe maintenance issues, including raw sewage backup into her bathtub and a complete failure of the apartment's heating furnace. Defendant sent written notices describing these hazardous conditions to Plaintiff's property manager via certified mail and text messaging on August 15, 2024, and September 4, 2024. Copies of these written communications are attached hereto as Exhibit B."
    )
    add_body_paragraph(
        doc,
        "Despite having ample reasonable time and opportunity to remedy these conditions, Plaintiff failed and refused to make repairs. Consequently, on October 2, 2024, Defendant engaged in statutorily protected activity by filing a formal housing code complaint with the City of Cleveland Department of Building and Housing."
    )
    add_body_paragraph(
        doc,
        "On October 8, 2024, a Cleveland Housing Code Inspector conducted an on-site inspection and issued a Notice of Violation citing Plaintiff for multiple hazardous conditions under the Cleveland Codified Ordinances and R.C. 5321.04(A)(1), including lack of operational heating facilities (Cleveland Cod. Ord. 369.15) and unsanitary plumbing backups (Cleveland Cod. Ord. 369.17). A certified copy of the City of Cleveland Inspection Report and Notice of Violation is attached hereto as Exhibit A."
    )
    add_body_paragraph(
        doc,
        "The City served the Notice of Violation on Plaintiff on October 12, 2024. Exactly six days later, on October 18, 2024, Plaintiff served Defendant with a Notice to Leave the Premises, followed immediately by the filing of this eviction complaint on October 25, 2024. Prior to Defendant's report to the city, Plaintiff had never issued any notice of lease violation or termination."
    )

    add_heading_1(doc, "II. LAW AND ARGUMENT")
    add_heading_2(doc, "A. Plaintiff's Action for Possession Is Barred as Unlawful Retaliation under R.C. 5321.02.")

    add_body_paragraph(
        doc,
        "Under R.C. 5321.02(A)(1), a residential landlord is strictly prohibited from bringing or threatening to bring an action for possession of the premises, increasing rent, or decreasing services when a tenant has complained to a governmental agency charged with the enforcement of building, housing, health, or safety codes regarding a violation that materially affects health and safety."
    )
    add_body_paragraph(
        doc,
        "To establish a defense under R.C. 5321.02, the tenant must demonstrate: (1) that the tenant engaged in protected activity; (2) that the landlord had notice or knowledge of the protected activity; (3) that the landlord took an adverse retaliatory action against the tenant, such as an eviction; and (4) that a causal connection exists between the protected activity and the adverse action. See Building Realty & Holding Co. v. Jackson, 8th Dist. Cuyahoga No. 40312, 1980 WL 354394 (Feb. 28, 1980); Karas v. Floyd, 2 Ohio App.3d 4, 440 N.E.2d 563 (2d Dist. 1981)."
    )
    add_body_paragraph(
        doc,
        "Here, every element of R.C. 5321.02 is established by uncontroverted documentary evidence. First, Defendant engaged in protected activity on October 2, 2024, by lodging a formal complaint with the Cleveland Department of Building and Housing regarding non-functioning heat and raw sewage. Second, Plaintiff received actual written notice of the code violations from the city inspector on October 12, 2024 (Exhibit A). Third, Plaintiff instituted eviction proceedings against Defendant. Fourth, the temporal proximity between Plaintiff learning of the city code violations on October 12 and serving the eviction notice on October 18—a mere six days—creates an overwhelming inference of retaliatory causation. See Smith v. Wright, 65 Ohio App.2d 101, 416 N.E.2d 655 (8th Dist. 1979). Because Plaintiff cannot establish any legitimate non-retaliatory justification, judgment on the first cause of action must be entered in favor of Defendant."
    )

    add_heading_2(doc, "B. Plaintiff Breached Its Mandatory Statutory Duties under R.C. 5321.04(A) and Cleveland Codified Ordinances, Precluding Restitution.")

    add_body_paragraph(
        doc,
        "Under R.C. 5321.04(A)(1), (A)(2), and (A)(4), a residential landlord has mandatory statutory obligations to comply with all applicable building, housing, and health codes; make all repairs and do whatever is reasonably necessary to put and keep the premises in a fit and habitable condition; and maintain in good and safe working order all heating, plumbing, and electrical fixtures. These obligations are implied into every residential lease agreement in Ohio and cannot be waived. See Shroades v. Rental Homes, Inc., 68 Ohio St.2d 20, 427 N.E.2d 774 (1981)."
    )
    add_body_paragraph(
        doc,
        "As documented by the city inspection report (Exhibit A) and Defendant's written repair notices (Exhibit B), Plaintiff committed severe breaches of R.C. 5321.04(A) by leaving Defendant without operational heat during freezing October temperatures and failing to remediate toxic sewage backup. The landlord had actual notice of the defects and more than two months to perform repairs, yet willfully failed to remedy the violations. Under Ohio law, a landlord's material breach of R.C. 5321.04 constitutes a complete defense to an action for possession and entitles the tenant to an order of rent abatement and statutory damages under R.C. 5321.04(B). See Miller v. Ritchie, 45 Ohio St.3d 222, 543 N.E.2d 1265 (1989)."
    )

    add_heading_1(doc, "III. RELIEF REQUESTED")

    add_body_paragraph(
        doc,
        "WHEREFORE, Defendant Elena Rostova respectfully requests that this Court enter judgment denying Plaintiff's claim for restitution of the premises; find that Plaintiff engaged in unlawful retaliation under R.C. 5321.02; determine that Plaintiff breached its statutory duties under R.C. 5321.04; award Defendant appropriate rent abatement and damages; and grant such further relief as law and equity require."
    )

    add_signature_block(
        doc,
        name="Marcus D. Washington",
        bar_no="0087654",
        title="Senior Attorney, Housing Practice Group",
        org="The Legal Aid Society of Cleveland",
        address="1223 West 6th Street, Cleveland, OH 44113",
        phone="(216) 687-1900",
        email="mwashington@lasclev.org",
        rep_for="Defendant Elena Rostova"
    )

    add_certificate_of_service(
        doc,
        service_text="I hereby certify that a true and accurate copy of the foregoing Defendant's Trial Brief was served upon Plaintiff's attorney of record, Arthur P. Vance, Esq., Vance & Sterling Co., LPA, 1300 East 9th Street, Suite 1400, Cleveland, OH 44114, via the Court's electronic filing system and email (avance@vancesterlinglaw.com) this 7th day of November, 2024.",
        attorney_name="Marcus D. Washington"
    )

    doc.save(output_path)


def generate_brief_3(output_path):
    """Generate Brief 3 (Defective - Substantive & Rule Audit Support Defects)."""
    doc = create_base_document()

    add_court_caption(
        doc,
        court_name="IN THE HAMILTON COUNTY MUNICIPAL COURT\nCINCINNATI, OHIO",
        plaintiff="RIVERFRONT HOLDINGS LP",
        defendant="JORDAN BLAKE",
        case_no="24CV-19342",
        judge_name="BERNIE BOUCHARD",
        doc_title="DEFENDANT'S BRIEF IN OPPOSITION TO EVICTION AND COUNTER-MEMORANDUM"
    )

    add_body_paragraph(
        doc,
        "Defendant Jordan Blake submits this brief in opposition to the eviction complaint filed by Riverfront Holdings LP. The eviction is improper, unfair, and retaliatory under R.C. 5321.02, and Plaintiff breached R.C. 5321.04 because the apartment is in bad shape."
    )

    add_heading_1(doc, "STATEMENT OF THE CASE")

    add_body_paragraph(
        doc,
        "Defendant Jordan Blake has lived at 824 Elm Street, Cincinnati, Ohio for a couple of years. The landlord has not been treating tenants fairly. Defendant was vocal around the building and in the neighborhood about problems in the community. Soon after, the landlord decided to bring this eviction case."
    )
    add_body_paragraph(
        doc,
        "The apartment has experienced multiple uncomfortable conditions. The general quality of the living space has been unsatisfactory and unpleasant. Because the landlord was not doing a good job, Defendant decided to withhold rent payments for two months until things improved."
    )

    add_heading_1(doc, "ARGUMENT")
    add_heading_2(doc, "A. The Eviction is Retaliation under R.C. 5321.02.")

    add_body_paragraph(
        doc,
        "Under R.C. 5321.02, a landlord cannot retaliate against a tenant. Here, Defendant was exercising free speech rights by discussing housing issues with neighbors. The landlord was unhappy with Defendant's attitude and brought an eviction action for possession. Therefore, the eviction is retaliatory and should be denied."
    )

    add_heading_2(doc, "B. The Landlord Breached Duties under R.C. 5321.04.")

    add_body_paragraph(
        doc,
        "Under R.C. 5321.04, landlords are supposed to keep apartments in fit and habitable condition. The premises here were not maintained at a high standard. There were inspections done that showed issues. Because the landlord breached its duties, Defendant had the right to withhold rent and the eviction cannot proceed."
    )

    add_heading_1(doc, "CONCLUSION")

    add_body_paragraph(
        doc,
        "WHEREFORE, Defendant requests that the Court deny the eviction and dismiss the case."
    )

    add_signature_block(
        doc,
        name="Thomas R. Miller",
        bar_no="0076543",
        title="Attorney at Law",
        org="Miller Law Office",
        address="312 Walnut Street, Suite 1600, Cincinnati, OH 45202",
        phone="(513) 555-0199",
        email="tmiller@millerlawohio.com",
        rep_for="Defendant Jordan Blake"
    )

    add_certificate_of_service(
        doc,
        service_text="A copy of this brief was sent to Plaintiff's attorney on September 12, 2024 by ordinary mail.",
        attorney_name="Thomas R. Miller"
    )

    doc.save(output_path)


def generate_brief_4(output_path):
    """Generate Brief 4 (Defective - Form, Typos, Placeholders, Numbering Gaps, Missing Cert)."""
    doc = create_base_document(default_size_pt=10, line_spacing=1.0)

    add_court_caption(
        doc,
        court_name="IN THE AKRON MUNICIPAL COURT\nSUMMIT COUNTY, OHIO",
        plaintiff="APEX RESIDENTIAL LLC",
        defendant="KEISHA TAYLOR",
        case_no="[Case Number TBD]",
        judge_name="",
        doc_title="MOTION TO DISMISS EVICTION COMPLIANT AND VACATE JUDGEMENT"
    )

    add_body_paragraph(
        doc,
        "Now comes defendent Keisha Taylor and moves this Court to dismiss the eviction complaint filed by plaintif Apex Residential LLC under Ohio statue.",
        indent=0.0,
        line_spacing=1.0,
        font_size_pt=10
    )

    add_heading_1(doc, "FACTS AND ARGUEMENT")

    p1 = doc.add_paragraph()
    p1.paragraph_format.line_spacing = 1.0
    p1.paragraph_format.space_after = Pt(4)
    p1.add_run("1. Defendent is a tenent residing at the subject property pursuant to Exhibit A (Attached Copy of Lease). Defendent never received any notice whatsoever from the landlorde.")

    p2 = doc.add_paragraph()
    p2.paragraph_format.line_spacing = 1.0
    p2.paragraph_format.space_after = Pt(4)
    p2.add_run("2. The the landlord failed to provide proper service under Cal. Civ. Code § 1161.The plaintif committed a material breech of contract.")

    p5 = doc.add_paragraph()
    p5.paragraph_format.line_spacing = 1.0
    p5.paragraph_format.space_after = Pt(4)
    p5.add_run("5. On [Insert Date of Notice], plaintif served a 3-day notice attached as Exhibit B (Notice).Plaintif filed this action ____ days after service without giving full statutory time.")

    add_heading_1(doc, "PRAYER FOR RELIEF")

    add_body_paragraph(
        doc,
        "WHEREFORE, defendent prays for an order dismissing the compliant and vacating any judgement.",
        indent=0.0,
        line_spacing=1.0,
        font_size_pt=10
    )

    p_sig = doc.add_paragraph()
    p_sig.paragraph_format.line_spacing = 1.0
    p_sig.paragraph_format.space_before = Pt(12)
    p_sig.add_run("Respectfully submitted,\n\n/s/ [Attorney Name]\n[Attorney Name]\nAttorney for Defendent\nAddress: TBD\nPhone: XXX-XXXX")

    doc.save(output_path)


def main():
    target_dir = Path("/home/quinten/agentic_housing_drafting/sample_briefs")
    target_dir.mkdir(parents=True, exist_ok=True)

    brief1_path = target_dir / "01_good_motion_to_dismiss_notice_defect_franklin_county.docx"
    brief2_path = target_dir / "02_good_trial_brief_retaliation_and_habitability_cleveland.docx"
    brief3_path = target_dir / "03_defective_brief_unsupported_retaliation_and_habitability_hamilton.docx"
    brief4_path = target_dir / "04_defective_brief_procedural_and_form_errors_akron.docx"

    print("Generating Brief 1 (Good - Notice Defect)...")
    generate_brief_1(brief1_path)

    print("Generating Brief 2 (Good - Retaliation & Habitability)...")
    generate_brief_2(brief2_path)

    print("Generating Brief 3 (Defective - Substantive & Rule Support Defects)...")
    generate_brief_3(brief3_path)

    print("Generating Brief 4 (Defective - Form, Placeholders, Typos & Missing Cert)...")
    generate_brief_4(brief4_path)

    print("All 4 sample briefs successfully created in", target_dir)


if __name__ == "__main__":
    main()
